#!/usr/bin/env python3
"""Generate Estonian guide DOCX for Puur on lahti / Peegel process."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/opt/cursor/artifacts/puur-on-lahti-juhend.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    return p


def add_table(doc, headers, rows, header_color="E8EEF4"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # Title page
    title = doc.add_heading("Puur on lahti", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Samm-sammuline juhend avalikele inimestele, poliitikutele ja kodanikele")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    meta = doc.add_paragraph("Peegel · Ausus · Suhete tugevdamine")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    doc.add_page_break()

    # --- 1. SISSEJUHATUS ---
    add_heading(doc, "1. Sissejuhatus", 1)
    add_para(
        doc,
        "See juhend põhineb kahel põhimõttel. Esiteks: traumeeritud keha (või organisatsioon) "
        "ei jookse isegi siis, kui uks on lahti — ta on õppinud, et liikumine ei aita. Teiseks: "
        "avalik eeskuju peegeldub kordistatult — kui avalik inimene kiusab või õpetab alandavalt, "
        "jälgijad kopeerivad sama tooni oma kogukondades.",
    )
    add_para(
        doc,
        "Lahendus ei ole süüdistamine. Lahendus on samm-sammuline protsess, kus suhted tugevnevad, "
        "kui keegi lõpetab teiste õpetamise ja hakkab ise riskima aususega.",
    )

    add_heading(doc, "1.1 Põhimõisted", 2)
    add_table(
        doc,
        ["Mõiste", "Tähendus"],
        [
            ["Puur on lahti", "Võimalus on olemas, aga keha/organisatsioon ei usu seda veel"],
            ["Peegel", "Mida eeskuju teeb, seda kordistatakse kogukondades"],
            ["Elicitation", "Avaldised küsimuste asemel — inimene avaneb, kui ta ei tunne rünnakut"],
            ["Ühine keel", "Üks lause, mida kõik saavad öelda ilma omistamise ja keerutamiseta"],
        ],
    )

    # --- 2. PROBLEEM ---
    add_heading(doc, "2. Probleemi kirjeldus", 1)
    add_para(
        doc,
        "Poliitilises kontekstis (sh presidendikandidatuur) on kõik osapooled samas lõksus:",
    )
    add_table(
        doc,
        ["Tegevus", "Mida see peaks tähendama", "Mida see tegelikult tähendab"],
        [
            ["Avalik toetus", "Erakondlikult erapooletu kandidaat", "Omistamine — kandidaat näeb erapoolik"],
            ["Vaikus", "Vastutustundlik hoidumine", "Paarisrakend keerutajaga"],
            ["Keerutamine", "Poliitiline manööver", "Avalik teater, privaatne lugupidamine"],
            ["Teiste kritiseerimine", "Aususe nõudmine", "Enda positsiooni kõrgemal hoidmine"],
        ],
    )

    add_heading(doc, "2.1 Võrdlustabel: mida ütlevad, mida teavad, mida jälgijad näevad", 2)
    add_table(
        doc,
        ["Kiht", "Mida nad ütlevad", "Mida nad teavad", "Mida jälgijad näevad"],
        [
            ["Avalik positsioon", "Seisan aususe eest", "Ausus maksab poliitiliselt", "Nii räägitakse teistest"],
            ["Teiste kohta", "Keerutab, silmakirjalik", "Meil on sama teadmine", "Õpetus: halvusta, ole targem"],
            ["Eeskuju", "Keegi peab ütlema", "Mitte mina esimesena", "17 000 peeglit kogukondades"],
            ["Tulemus", "Aitan avalikkust mõista", "Mängin samas teatris", "Avalik alandamine on normaalne"],
        ],
    )

    doc.add_page_break()

    # --- 3. LAHENDUS ---
    add_heading(doc, "3. Lahenduse ülevaade", 1)
    add_para(doc, "Suhted tugevnevad, kui:", bold=True)
    for item in [
        "Keegi nimetab lõksu ühiselt — ilma süüdistuseta",
        "Leitakse üks ühine keel, mida kõik saavad öelda",
        "Keegi läheb esimesena — haavatavusega, mitte õpetamisega",
        'Teised kinnitavad: "Sama" - ilma võistluseta',
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_quote(
        doc,
        "Suhted tugevnevad siis, kui keegi lõpetab teiste süüdistamise "
        "ja ütleb esimesena ausalt, mida tema ise on valmis tegema.",
    )

    # --- 4. REEGLID ---
    add_heading(doc, "4. Protsessi reeglid (enne alustamist)", 1)
    for i, rule in enumerate(
        [
            "Ei süüdistata üksteist — räägime lõksust, mitte süüdlasest",
            'Ei omistata kandidaati - toetame rolli, mitte "oma inimest"',
            "Üks samm korraga — järgmine samm tuleb alles siis, kui eelmine on tehtud",
            "Keegi ei pea võitma — kui keegi kaotab näo, protsess kukub",
        ],
        1,
    ):
        add_para(doc, f"{i}. {rule}")

    doc.add_page_break()

    # --- 5. SAMMUD ---
    add_heading(doc, "5. Samm-sammuline protsess", 1)

    steps = [
        (
            "SAMM 0: Peata avalik skoorimine (24 tundi)",
            "Kõik, kes on viimati postitanud teiste kohta süüdistusi, ei postita 24 tundi uut süüdistust.",
            "Ma ei kommenteeri teisi täna. Ma mõtlen, kuidas me kõik saaksime sama asja ausamalt öelda.",
            "Kas keegi suutis 24h ilma õpetamiseta olla?",
        ),
        (
            "SAMM 1: Nimetame lõksu ühiselt",
            "Üks inimene nimetab olukorra lõksuna — mitte ühe erakonna veana.",
            "Me oleme kõik samas olukorras. Kui ma toetan avalikult, näeb see välja nagu omistamine. "
            "Kui ma vaikin, näeb see välja nagu mäng. See on lõks — mitte ühe erakonna viga.",
            "Kas vähemalt 2 erakonda on öelnud, et tunnevad sama lõksu?",
        ),
        (
            "SAMM 2: Üks küsimus kandidaadile",
            'Keegi küsib neutraalselt: "Mis on üks viis, kuidas poliitikud saaksid teid toetada, '
            'ilma et te näeks erapoolik?"',
            'Kandidaadi vastus annab kõigile loa ja reeglid. Näide: "Öelge, et usaldate mu '
            'konstitutsioonilist sõltumatust. Ärge öelge, et olen teie kandidaat."',
            "Kas kandidaat on andnud ühe selge rea toetuse kohta?",
        ),
        (
            "SAMM 3: Ühine keel — üks lause",
            "Erakonnad lepivad kokku ühe lause, mida kõik võivad sõna-sõnalt öelda.",
            "Me hindame [kandidaadi] konstitutsioonilist rolli ja sõltumatust. "
            "Me ei tee temast erakondlikku kandidaati. Meie parlamendiliikmed otsustavad vabalt.",
            "Kas vähemalt 2 erakonda on öelnud sama lauset sõna-sõnalt?",
        ),
        (
            "SAMM 4: Avalik samm — keegi läheb esimesena",
            "Üks avalik inimene ütleb haavatavusega, ilma teisi nimetamata.",
            "Ka mina olen selles lõksus olnud. Täna ütlen ühe lause ausalt, ilma teisi süüdistamata: "
            "[SAMM 3 LAUSE].",
            "Kas postitus on ilma teiste nimeta?",
        ),
        (
            "SAMM 5: Teised kinnitavad",
            'Teised erakonnad ütlevad lihtsalt "Sama" - ilma "lõpuks" või "nagu me alati ütlesime".',
            "Sama.",
            "Kas on kordus ilma võistluseta?",
        ),
        (
            "SAMM 6: Ajakirjandusele vastus",
            "Ühtne keel lõpetab keerutamise. Edasi on parlament, mitte Facebook.",
            "Me ei oma kandidaati. Me usaldame konstitutsioonilist protsessi. "
            "Meie liikmed hääletavad vabalt.",
            "Kas ajakirjandus saab vastuse ilma uue spinni ringita?",
        ),
    ]

    for title, action, phrase, check in steps:
        add_heading(doc, title, 2)
        add_para(doc, "Mida teha:", bold=True)
        add_para(doc, action)
        add_para(doc, "Üks lause avalikult:", bold=True)
        add_quote(doc, phrase)
        add_para(doc, "Kontroll:", bold=True)
        add_para(doc, check)
        doc.add_paragraph()

    doc.add_page_break()

    # --- 6. VISUAALNE ÜLEVAADE ---
    add_heading(doc, "6. Protsessi ülevaade", 1)
    flow = (
        "SAMM 0  ->  24h ilma süüdistuseta\n"
        "   |\n"
        "SAMM 1  ->  Me kõik oleme lõksus\n"
        "   |\n"
        "SAMM 2  ->  Küsimus kandidaadile: Kuidas toetada õigesti?\n"
        "   |\n"
        "SAMM 3  ->  Üks ühine lause (kõigile)\n"
        "   |\n"
        "SAMM 4  ->  Keegi läheb esimesena (haavatavus)\n"
        "   |\n"
        "SAMM 5  ->  Teised: Sama\n"
        "   |\n"
        "SAMM 6  ->  Ajakirjandusele: ühtne keel, lõpp keerutamisele"
    )
    p = doc.add_paragraph()
    run = p.add_run(flow)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    # --- 7. TAGASILANGUS ---
    add_heading(doc, "7. Kui keegi kukub tagasi", 1)
    add_para(
        doc,
        "Ära vasta süüdistusega. Üks lause:",
    )
    add_quote(doc, "Me olime samm 3 juures. Kas saame tagasi ühise lause juurde?")

    # --- 8. KODANIK ---
    add_heading(doc, "8. Kodaniku juhend — kommentaarid samm-sammult", 1)
    add_table(
        doc,
        ["Samm", "Mida kirjutad kommentaari"],
        [
            ["0", "Kas saame ühe päeva ilma teiste süüdistamata?"],
            ["1", "Kas see lõks kirjeldab ka teie olukorda?"],
            ["2", "Kas keegi on kandidaadilt küsinud, kuidas teda õigesti toetada?"],
            ["3", "Kas see üks lause sobib kõigile?"],
            ["4", "Kes läheb esimesena — ilma teisi nimetamata?"],
            ["5", "Kas teised saavad lihtsalt öelda: sama?"],
        ],
    )

    add_heading(doc, "8.1 Elicitation-stiilis avaldused (sõbralikud)", 2)
    add_table(
        doc,
        ["Tehnika", "Näide"],
        [
            ["Rekordi parandamine", "Mulle tundub, et erakonnad tõesti tahavad suruda läbi oma kandidaadi..."],
            ['"Ma paneks, et..."', "Ma paneks, et kõik juba teavad, aga keegi ei taha olla esimene"],
            ["Peegel", "Kas te mõtlete, mida teie jälgijad teie toonist õpivad?"],
            ["Siirus", "Üks lause ausust — ja kogu mäng muutub. Kes ütleb esimesena?"],
        ],
    )

    doc.add_page_break()

    # --- 9. SOTSMEEDIA ---
    add_heading(doc, "9. Sotsiaalmeedia tekstid", 1)

    add_heading(doc, "9.1 YouTube Short — kirjeldus", 2)
    add_quote(
        doc,
        "Puur on lahti. Aga keha ei usu seda veel.\n\n"
        "Kui avalik inimene kiusab oma lehel teisi, siis tuhanded jälgijad ei näe erandit. "
        "Nad näevad luba.\n\n"
        "See on peegel: mida sa teed üleval, seda tehakse all.\n\n"
        "Kommenteeri PEEGEL, kui nägid seda oma silmaga.",
    )

    add_heading(doc, "9.2 Kinnitatud kommentaar", 2)
    add_quote(
        doc,
        "Peegel ei süüdista. Peegel näitab.\n\n"
        "Kommenteeri PEEGEL, kui see kõlas tõena.",
    )

    add_heading(doc, "9.3 Protsessi käivitamine", 2)
    add_quote(
        doc,
        "Kas me saame minna samm-sammult: esmalt lõksu nimetamine, siis ühine lause, "
        "siis keegi esimesena — ilma süüdistuseta?",
    )

    # --- 10. LÕPETUS ---
    add_heading(doc, "10. Kokkuvõte", 1)
    add_table(
        doc,
        ["", ""],
        [
            ["Ütlevad", "Keegi peab aus olema"],
            ["Teadavad", "Aga mitte mina esimesena"],
            ["Jälgijad näevad", "Nii käitutakse, kui oled targem ja õigem"],
            ["Lahendus", "Keegi lõpetab õpetamise ja riskib esimesena"],
        ],
    )

    add_para(doc, "")
    add_para(
        doc,
        "— Lõpp —",
        bold=True,
    )
    footer = doc.add_paragraph("Puur on lahti · Peegel · Samm-sammuline ausus")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
