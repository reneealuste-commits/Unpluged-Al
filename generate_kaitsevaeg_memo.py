#!/usr/bin/env python3
"""Generate official 1-page Kaitseväe memo DOCX with references."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/workspace/kaitsevaeg-memo-raamat-rakendamine.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def build():
    doc = Document()

    # Header block
    for label, value in [
        ("Kellele:", "Kaitseväe juhtkond / personali- ja arenguvaldkonna vastutajad"),
        ("Kellelt:", "[Nimi, kontakt — täiendada enne saatmist]"),
        ("Kuupääev:", "4. august 2026"),
        ("Teema:", "Raamatu rakendamine Kaitseväes: eeltingimused, põhjused ja soovituslik faasiprotsess"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(label + " ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(value)
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # Title
    t = doc.add_paragraph("MEMO")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.bold = True
        r.font.size = Pt(14)

    sub = doc.add_paragraph(
        "Miks varasem levitamine ei toonud rakendustulemust — ja kuidas protsessi korraldada"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(11)

    doc.add_paragraph()

    sections = [
        (
            "1. Taust ja tagasiside",
            "Kaitseväes on levitatud teos, mille sisu toetab ausat enesehindamist, "
            "emotsionaalset regulatsiooni ja organisatsioonilist usaldust. Tagasiside ühelt "
            "kommenteerijalt: raamat „ei töötanud“ — inimesed ei teadnud, kuidas sisu "
            "rakendada. See ei viitab teose ebaõnnestumisele, vaid rakendamise "
            "oskuse ja eeltingimuste puudumisele.",
        ),
        (
            "2. Põhjuslik analüüs",
            "Hierarhilises ja surve all olevas keskkonnas — kus domineeriv juhtimisstiil, "
            "avalik alandamine või kiusamine on mõjutanud osa inimeste käitumist — ei saabu "
            "kõik tööle samal autonoomse nervisüsteemi baasjoone tasemel. Osad on "
            "hüperaktiivses (fight/flight), osad külmutatud (freeze/fawn) režiimis. "
            "Sellises seisundis ei saa inimene usaldusväärselt rakendada uut käitumist või "
            "organisatsioonilisi muutusi, isegi kui teoreetiline materjal on kättesaadav "
            "(van der Kolk, 2014; Porges, 2011).",
        ),
        (
            "3. Eesti ja Kaitseväe kontekst",
            "Eestis ja eriti sõjalises struktuuris võib domineerimine ja kiusamine "
            "kujuneda „tugeva juhi“ normiks. See õpetab alluvatel: ära riski, ära ava emotsioone, "
            "ära vaidle avalikult. Uuringud näitavad, et korduv surve ja hirm jätab jäljed "
            "nii individi kui rühma usaldusväärsusesse (Mitchell & Everly, 1997; Bourne et al., "
            "2022). Ilma turvalisuse ja regulatsiooni taastamiseta jääb raamat loetud "
            "materjaliks — mitte muutuse vahendiks.",
        ),
        (
            "4. Soovitus: faasiprotsess enne „rakendamist“",
            "Soovitame mitte skaleerida raamatu „sammude“ rakendamist enne järgmisi faase:\n"
            "• Faas 0 (4–8 nädalat): baasjoone reguleerimine — kehapõhised resetid, "
            "juhtide turvalisuse signaal, null avalik alandamine pilotrühmades.\n"
            "• Faas 1: valuliste emotsioonide turvaline nimetamine (mitte sunnitud avalikustamine).\n"
            "• Faas 2: üks raamatu kontsepts = üks kehaline kogemus rühmas.\n"
            "• Faas 3: igapäevane rakendamine — alles siis skaleerimine teistesse üksustesse.\n"
            "Pilot: 1 juhtide rühm + 1 väike üksus (8–12 inimest). Edukuse mõõdik: "
            "vähenenud avalik alandamine, võime resetida enne rasket teemat, üksaus lause ilma karistuseta.",
        ),
        (
            "5. Järeldus ja otsussoovitus",
            "Raamat saab Kaitseväes tööle siis, kui lõpetame ainult levitamise ja alustame "
            "juhendatud faasiprotsessi, kus juhid lähevad esimesena ja domineerimise "
            "vähendamine on eeltingimus, mitte järeltegevus. Soovitame kinnitada pilotprogramm "
            "ja määrata juhendaja, kellel on teadmine ja praktiline oskus regulatsiooni toetamiseks.",
        ),
    ]

    for title, body in sections:
        h = doc.add_paragraph()
        hr = h.add_run(title)
        hr.bold = True
        hr.font.size = Pt(10)
        for para in body.split("\n"):
            p = doc.add_paragraph(para)
            for r in p.runs:
                r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    ref_h = doc.add_paragraph()
    ref_r = ref_h.add_run("Viited")
    ref_r.bold = True
    ref_r.font.size = Pt(10)

    refs = [
        "Bourne, J. E., Mackintosh, J. E., Holmes, E. A., & Rennie, C. L. (2022). "
        "The impact of organizational trauma on workplace behavior. Frontiers in Psychology, 13, 809209.",
        "Maier, S. F., & Seligman, M. E. P. (2016). Learned helplessness at fifty: "
        "Insights from neuroscience. Psychological Review, 123(4), 349–367.",
        "Mitchell, J. T., & Everly, G. S. (1997). Critical Incident Stress Debriefing (CISD): "
        "An operations manual for the prevention of traumatic stress among emergency services "
        "and public safety workers. Chevron Publishing.",
        "Porges, S. W. (2011). The polyvagal theory: Neurophysiological foundations of "
        "emotions, attachment, communication, and self-regulation. Norton.",
        "van der Kolk, B. (2014). The body keeps the score: Brain, mind, and body in the "
        "healing of trauma. Viking.",
        "Riigikogu / ühiskondlik kontekst: domineeriva juhtimise ja avaliku alandamise "
        "mõju usaldusväärsele suhtlusele (vt ka: van der Kolk, 2014, peatükk 2 — "
        "õpitud abitus ja immobiliseerimine).",
        "Kaitseväe siseprotseduurid: personali arendamine ja juhtimiskultuur — "
        "siduda pilot Kaitseväe arenguprioriteetidega (täiendada vastava viitega sisekorraldusele).",
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"{i}. {ref}")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(9)

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Dokument on ettepanekulise iseloomuga memo. Täiendada saatja andmed enne edastamist."
    )
    for r in foot.runs:
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
