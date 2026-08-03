#!/usr/bin/env python3
"""Generate Riigikogu letter DOCX + email list for Riho Ühtegi campaign."""

import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DOCX = "/workspace/riho-uhtegi-riigikogu-kiri.docx"
OUTPUT_TXT = "/workspace/riho-uhtegi-emailid.txt"
OUTPUT_CSV = "/workspace/riho-uhtegi-emailid.csv"
KONTAKTID = "/home/ubuntu/.cursor/projects/workspace/agent-tools/f082ccda-5218-4e97-9354-2ad6f71b39f3.txt"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def load_members():
    text = open(KONTAKTID, encoding="utf-8").read()
    liikmed = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*Riigikogu liige\s*\|\s*[^|]+\|\s*E-post\s*([a-z0-9._+-]+@riigikogu\.ee)",
        text,
    )
    juhatus = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*Riigikogu (?:esimees|I aseesimees|II aseesimees)\s*\|\s*[^|]+\|\s*E-post\s*([a-z0-9._+-]+@riigikogu\.ee)",
        text,
    )
    seen = set()
    members = []
    for name, email in juhatus + liikmed:
        email = email.lower()
        if email not in seen:
            seen.add(email)
            members.append({"nimi": name.strip(), "email": email, "roll": ""})
    for m in members:
        if m["email"] == "lauri.hussar@riigikogu.ee":
            m["roll"] = "Riigikogu esimees"
        elif m["email"] == "toomas.kivimagi@riigikogu.ee":
            m["roll"] = "I aseesimees"
        elif m["email"] == "arvo.aller@riigikogu.ee":
            m["roll"] = "II aseesimees"
        else:
            m["roll"] = "Riigikogu liige"
    return sorted(members, key=lambda x: x["nimi"])


LETTER_BODY = """Lugupeetud {nimi}!

Kirjutan Teile kodanikuna, sest Eesti vajab praegu presidendi valikut, mis on päris — mitte järjekordne poliitiline teater.

21.–24. augustil 2026 saab Riigikogu esitada presidendikandidaate. Kandidaadi ülesseadmiseks on vaja vähemalt 21 Riigikogu liikme allkirja. Esimene hääletusvoor toimub 2. septembril 2026.

MIKS JUST NÜÜD?

Presidendivalimiste ümber on kujunenud olukord, kus kõik osapooled teavad, et üks erapoolik toetus kahjustab kandidaadi usaldusväärsust — ja siiski keerutatakse, vaikitakse või omistatakse. See on lõks, mitte ühe erakonna viga.

Ühiskonnategelased — Sven Grünberg, Hugo Osula, Kris Taska, Hendrik Toompere ja Raivo Vare — on pakkunud välja Riho Ühtegi kui kandidaadi, kellel on nõusolek kandideerimiseks. Tema ei tuleks presidendiks erakonna esindajana, vaid rahva ettepanekuna.

MIKS RIHO ÜHTEGI?

• Ta ei ole erakonna esindaja — ta tuleb ühiskonnast, mitte koalitsioonist
• Ta on öelnud: „Kui saan presidendiks, panen tooli Kadrioru lossi ette ja kuulan inimesi"
• Ta nimetab asju õigete nimedega — ilma keerutamiseta
• Tal on kogemus kriisijuhtimises, riigikaitse ja rahvusvahelises koostöös
• Ta teab, et riigikaitse algab rahva enesekindlusest — mitte ainult relvadest

MIDA PALUN?

Te võite järgmisel korral toetada kellegi teist. See on täiesti mõistlik.

Aga seekord — augustis ja septembris — palun tehke see päriselt.

1. Kaaluge Riho Ühtegi kandidatuuri toetamist (21+ allkirja)
2. Kui toetate, öelge seda nii: „Usaldame tema konstitutsioonilist sõltumatust — me ei omista teda erakonnale"
3. Lõpetage keerutamine — üks aus lause on parem kui kuu poliitilist teatrit

ÜHINE KEEL (mida kõik saavad öelda):

„Me hindame Riho Ühtegi konstitutsioonilist rolli ja sõltumatust. Me ei tee temast erakondlikku kandidaati. Meie parlamendiliikmed otsustavad vabalt."

Suhted tugevnevad siis, kui keegi lõpetab teiste süüdistamise ja ütleb esimesena ausalt, mida ta ise on valmis tegema.

Seekord vali päriselt. Seekord — Riho Ühtegi.

Lugupidamisega,
[Teie nimi]
[Teie kontakt]

---
Allikas: riigikogu.ee kontaktid (uuendatud enne dokumendi koostamist)
Presidendivalimised 2026: kandidaatide esitamine 21.–24.08, hääletus 02.09"""


SUBJECT = "Palun: seekord vali päriselt — Riho Ühtegi presidendikandidaadiks (21.–24.08)"


def build_docx(members):
    doc = Document()

    title = doc.add_heading("Kiri Riigikogu liikmetele", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Riho Ühtegi presidendikandidatuur · august–september 2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.italic = True
        r.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_heading("E-kirja teema (subject)", 1)
    p = doc.add_paragraph(SUBJECT)
    for r in p.runs:
        r.bold = True

    doc.add_heading("Kirja sisu (mall)", 1)
    doc.add_paragraph(
        "Asenda {nimi} saadiku nimega. Saada ükshaaval — isiklik kiri on tõhusam kui masspost."
    )

    for para in LETTER_BODY.split("\n\n"):
        if para.startswith("•"):
            for line in para.split("\n"):
                doc.add_paragraph(line.lstrip("• ").strip(), style="List Bullet")
        elif para.startswith("„") or para.startswith('"'):
            q = doc.add_paragraph(para)
            q.paragraph_format.left_indent = Cm(1)
            for r in q.runs:
                r.italic = True
        else:
            doc.add_paragraph(para)

    doc.add_page_break()

    doc.add_heading("Juhised saatmiseks", 1)
    for tip in [
        "Saada ükshaaval — mitte BCC masskirjana (spam risk + vähem isiklik).",
        "Lisa oma nimi ja kontakt lõppu.",
        "Kui saad vastuse, edasta teistele — see aitab protsessi.",
        "E-postid on Riigikogu ametlikult kodulehelt (riigikogu.ee/kontaktid).",
        f"Kokku {len(members)} saadikut (XV Riigikogu koosseis + juhatus).",
    ]:
        doc.add_paragraph(tip, style="List Bullet")

    doc.add_heading("Olulised kuupäevad", 1)
    add_table(
        doc,
        ["Kuupäev", "Sündmus"],
        [
            ["21.–24. august 2026", "Kandidaatide esitamine (vaja 21+ allkirja)"],
            ["2. september 2026", "Esimene hääletusvoor Riigikogus"],
            ["3. september 2026", "Vajadusel 2. ja 3. hääletusvoor"],
            ["10. oktoober 2026", "Ametis oleva presidendi ametiaja lõpp"],
        ],
    )

    doc.add_page_break()

    doc.add_heading(f"Riigikogu liikmete e-postid ({len(members)})", 1)
    doc.add_paragraph("Märgi ✓ peale saatmist.")

    table = doc.add_table(rows=1 + len(members), cols=4)
    table.style = "Table Grid"
    headers = ["#", "Nimi", "E-post", "Saadetud"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "E8EEF4")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for idx, m in enumerate(members, 1):
        row = table.rows[idx].cells
        row[0].text = str(idx)
        row[1].text = m["nimi"]
        row[2].text = m["email"]
        row[3].text = "☐"
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    footer = doc.add_paragraph("Seekord vali päriselt. Seekord — Riho Ühtegi.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    doc.save(OUTPUT_DOCX)
    print(f"DOCX: {OUTPUT_DOCX}")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "E8EEF4")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()


def export_lists(members):
    with open(OUTPUT_TXT, "w", encoding="utf-8") as f:
        f.write(f"# Riho Ühtegi — Riigikogu e-postid ({len(members)})\n")
        f.write(f"# Allikas: https://www.riigikogu.ee/kontaktid/\n\n")
        for i, m in enumerate(members, 1):
            f.write(f"{i:3}. {m['nimi']:<35} {m['email']}\n")
        f.write("\n# Ainult e-postid (kopeerimiseks):\n\n")
        for m in members:
            f.write(m["email"] + "\n")

    with open(OUTPUT_CSV, "w", encoding="utf-8") as f:
        f.write("nimi,email,roll\n")
        for m in members:
            f.write(f"\"{m['nimi']}\",{m['email']},\"{m['roll']}\"\n")

    print(f"TXT: {OUTPUT_TXT}")
    print(f"CSV: {OUTPUT_CSV}")


if __name__ == "__main__":
    members = load_members()
    build_docx(members)
    export_lists(members)
    print(f"Total: {len(members)} members")
