#!/usr/bin/env python3
"""Build pandoc reference.docx with OPORD-friendly styles."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "reference-opord.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x2A)
GRAY = RGBColor(0x44, 0x44, 0x44)


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        (1, 18, NAVY),
        (2, 14, NAVY),
        (3, 12, GRAY),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph("OPERATSIOON PEEGEL", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "Calibri"
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.color.rgb = NAVY

    sub = doc.add_paragraph("OPORD koos koigi lisadega", style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Naidistekst - see fail on pandoc vormistusmall.", style="Normal")
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
