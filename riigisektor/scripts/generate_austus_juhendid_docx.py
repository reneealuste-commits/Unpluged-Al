#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate public-sector respect guides DOCX - Techno TLN flagship + all roles."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

OUT = "/workspace/riigisektor/techno-tln-austus-sisekorra-eeskiri.docx"
OUT_ALL = "/workspace/riigisektor/riigisektor-rollipohine-austus-juhend.docx"


def h(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))
    if level == 1:
        r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    p.paragraph_format.space_before = Pt(12 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(6)


def p(doc, text, bold=False, italic=False, center=False):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    para.paragraph_format.space_after = Pt(5)


def b(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for r in para.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def n(doc, text):
    para = doc.add_paragraph(text, style="List Number")
    for r in para.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def role_block(doc, title, roll, ei_tee, tee, keeles, raport):
    h(doc, title, 3)
    p(doc, f"Roll: {roll}", italic=True)
    p(doc, "Ei ole lugupidav / ei ole lubatud:", bold=True)
    for x in ei_tee:
        b(doc, x)
    p(doc, "N\u00e4itan austust ja lugupidamist:", bold=True)
    for x in tee:
        b(doc, x)
    p(doc, "Keelekasutus:", bold=True)
    for x in keeles:
        b(doc, x)
    p(doc, f"Konflikt / turvalisus: {raport}")


def build_techno_tln(doc):
    h(doc, "TALLINNA TEHNOLOOGIAKOLLEDZ Techno TLN", 1)
    p(doc, "Austava ja lugupidava suhtluse sisekorra eeskiri", italic=True, center=True)
    p(doc, "Registrikood 70003767 | techno.ee | info@techno.ee", center=True)
    p(doc, "Kinnitatud: juhtkonna ettepanekul | Kehtib: kogu kooliperele", center=True)

    h(doc, "1. Eeskiri ja eesm\u00e4rk", 2)
    p(doc, "See eeskiri on Techno TLN sisekorra osa. Eesm\u00e4rk on selgus, turvalisus ja lugupidav keskkond neljas \u00f5ppelinnakus (Mustam\u00e4e, J\u00e4rve, Lasnam\u00e4e, Kesklinn) umbes 6000 \u00f5ppija ja kogu t\u00f6\u00f6tajaskonna jaoks.")
    p(doc, "Eeskiri p\u00f5hineb kooli p\u00f5hiv\u00e4\u00e4rtustel:")
    b(doc, "Kvaliteet - hea haridus ja professionaalne suhtlus")
    b(doc, "K\u00e4ttesaadavus - austus s\u00f5ltumata taustast, v\u00f5imetest v\u00f5i linnakust")
    b(doc, "Hoolivus - lugupidav suhtumine endasse, kaaslastesse ja keskkonda")

    h(doc, "2. P\u00f5him\u00f5tted (k\u00f5igile)", 2)
    n(doc, "Iga inimene on v\u00e4\u00e4rtus. Auaste, roll v\u00f5i vanus ei anna \u00f5igust alav\u00e4\u00e4ristada.")
    n(doc, "Austus on k\u00e4itumine: kuulan, vastan, olen kohal, t\u00e4idan kokkulepped.")
    n(doc, "Konflikt lahendatakse rahulikult, kirjalikult vajadusel, mitte avalikult h\u00e4bistades.")
    n(doc, "Turvalisus on esimene: f\u00fc\u00fcsiline, vaimne ja digitaalne turvalisus on k\u00f5igi \u00fchine vastutus.")
    n(doc, "Nelja kooli \u00fchendamine n\u00f5uab kannatlikkust. Uus nimi - sama v\u00e4\u00e4rikus.")

    h(doc, "3. \u00dcldised keelatud tegevused", 2)
    b(doc, "Solvamine, soim, alav\u00e4\u00e4ristamine, rassistlikud v\u00f5i seksistlikud v\u00e4ljendid")
    b(doc, "Karjumine \u00f5pilase, kolleegi v\u00f5i lapsevanema peale")
    b(doc, "Katkendamine, ignoreerimine, kui keegi p\u00f6\u00f6rdub ametliku k\u00fcsimusega")
    b(doc, "Avalik h\u00e4bistamine sotsiaalmeedias v\u00f5i kooli kanalites")
    b(doc, "Mobiil telefonis \u00f5ppetunnil v\u00f5i ametliku koosoleku ajal (v\u00e4lja arvatud kokkulepitud erandid)")
    b(doc, "Diskrimineerimine soo, p\u00e4ritolu, keele, erivajaduse v\u00f5i maailmavaate alusel")
    b(doc, "V\u00e4givald, \u00e4hvardus v\u00f5i selle \u00f5hutamine")

    h(doc, "4. Rollip\u00f5hised juhised Techno TLN-s", 2)

    role_block(doc, "4.1 Direktor ja juhtkond (Ott P\u00e4rna, n\u00f5ukogu, juhtkond)",
        "Organisatsiooni eeskuju; otsused ja kommunikatsioon",
        ["Ei tee otsuseid avaliku h\u00e4bistamise kaudu", "Ei ignoreeri t\u00f6\u00f6taja v\u00f5i \u00f5pilase p\u00f6\u00f6rdumist", "Ei kasuta ametipositsiooni isiklikuks surveteks"],
        ["Kuulab enne otsust", "Selgitab otsuse p\u00f5hjuse l\u00fchidalt", "V\u00f5tab vastutuse ja annab tagasisidet", "On kohal kriisis"],
        ['"Ait\u00e4h, et t\u00f5id selle mulle.", "Selgitan, miks nii otsustasime."', '"Sul on \u00f5igus k\u00fcsida."'],
        "personal@ / info@techno.ee; kriisi korral kohe turvat\u00f6\u00f6 ja juhtkond")

    role_block(doc, "4.2 \u00d5ppejuht ja \u00f5petaja (Birgit Vilgats, \u00f5ppej\u00f5ud, juhendajad)",
        "\u00d5pikeskkonna turvalisus ja professionaalsus",
        ["Ei naera \u00f5pilase k\u00fcsimuse \u00fcle", "Ei v\u00f5rdle \u00f5pilasi avalikult", "Ei j\u00e4ta kiusamist m\u00e4rkamata"],
        ["Tervitab iga tundi", "Selgitab ootused esimesel tunnil", "Annab tagasisidet v\u00e4\u00e4rikalt", "Suunab abi vajadusel (sotsiaalpedagoog, ps\u00fchholoog)"],
        ['"Head k\u00fcsimus."', '"Ma kuulan sind."', '"Proovime koos lahendust."'],
        "valdkonnajuht -> \u00f5ppejuht -> direktor; kiusamise korral kohe")

    role_block(doc, "4.3 \u00d5pilane ja \u00f5pilasesindus (Kevin Lusti jt)",
        "\u00d5ppija ja kogukonna liige",
        ["Ei katkesta teadlikult teist", "Ei levita kuulujutte v\u00f5i pilte ilma loata", "Ei solva \u00f5petajat ega kaaslast"],
        ["Tervitab ja t\u00e4nab", "K\u00fcsib abi, kui ei saa hakkama", "Hoiatab, kui n\u00e4eb kiusamist", "J\u00e4rgib t\u00f6\u00f6ohutust t\u00f6\u00f6koda"],
        ['"Vabandust."', '"Kas saaksin k\u00fcsida?"', '"Ma vajan abi."'],
        "klassijuhataja, sotsiaalpedagoog, \u00f5pilasesindus")

    role_block(doc, "4.4 Personal ja HR (Andra Piirsalu)",
        "Inimeste k\u00fcsimused, konfliktide ennetus",
        ["Ei r\u00e4\u00e4gi t\u00f6\u00f6taja isiklikust teemast kolleegidega", "Ei j\u00e4ta p\u00f6\u00f6rdumist vastuseta \u00fcle 5 t\u00f6\u00f6p\u00e4eva"],
        ["Kuulab konfidentsiaalselt", "Selgitab protsessi", "Toetab juhte \u00fchendamise perioodil"],
        ['"Sinu p\u00f6\u00f6rdumine on registreeritud."', '"J\u00e4rgmine samm on..."'],
        "direktor; vajadusel t\u00f6\u00f6inspektsioon, \u00f5igusabi")

    role_block(doc, "4.5 IT ja digiteenused (Toivo P\u00e4rnpuu)",
        "Digitaalne turvalisus ja ligip\u00e4\u00e4setavus",
        ["Ei naera kasutaja oskamatuse \u00fcle", "Ei jaga paroole ega isikuandmeid"],
        ["Vastab arusaadavalt", "Aitab ilma kiirustamata", "Kaitseb andmeid"],
        ['"Selgitan samm-sammult."', '"Sinu andmed on kaitstud."'],
        "juhtkond; andmekaitse n\u00f5uded")

    role_block(doc, "4.6 Valdkonnajuht ja meister (Ander Sile, valdkonnad)",
        "Praktika, t\u00f6\u00f6koda, t\u00f6\u00f6stuskoost\u00f6\u00f6",
        ["Ei j\u00e4ta \u00f5pilast t\u00f6\u00f6koda j\u00e4relevalveta", "Ei alav\u00e4\u00e4rista nais\u00f5pilasi tehnikas"],
        ["\u00d5petab turvalisust enne masinat", "Annab ausat tagasisidet", "\u00dchendab mentoritega"],
        ['"Turvalisus enne tulemust."', '"Viga on \u00f5ppimise osa."'],
        "\u00f5ppejuht, t\u00f6\u00f6ohutus")

    role_block(doc, "4.7 Lapsevanem ja partner",
        "Kodu ja kooli koost\u00f6\u00f6",
        ["Ei solva \u00f5petajat ega teist lapsevanemat", "Ei lahenda lapse konflikti teise lapse \u00fcle karjudes"],
        ["V\u00f5tab \u00fchendust ametliku kanali kaudu", "Kuulab enne s\u00fc\u00fcdistamist", "Tuleb koosolekule valmis"],
        ['"Tahan koost\u00f6\u00f6d."', '"Mis on teie vaade?"'],
        "klassijuhataja, sotsiaalpedagoog")

    role_block(doc, "4.8 Hooldus, turva, administratsioon",
        "Igap\u00e4evane kogemus koolimajas",
        ["Ei ignoreeri \u00f5pilast ukse taga", "Ei diskrimineeri"],
        ["Tervitab", "Suunab \u00f5ige inimese juurde", "Teatab ohtlikust olukorrast"],
        ['"Kuidas saan aidata?"'],
        "vahetu \u00fclem, turvat\u00f6\u00f6")

    role_block(doc, "4.9 Arendus ja partnerlus (T\u00f5nu Armulik, Kristel Bankier, Kristel Martis)",
        "Areng, finants, kogukond ja avalik s\u00f5num",
        ["Ei lubab partneril solvata t\u00f6\u00f6tajat", "Ei j\u00e4ta kogukonna tagasisidet kuulmata"],
        ["Selgitab otsuseid ausalt", "Toetab mentorlust ja praktikat", "Kaitseb \u00f5ppija v\u00e4\u00e4rikust avalikus s\u00f5numis"],
        ['"Koos loome v\u00f5imalusi."', '"Sinu tagasiside on oluline."'],
        "direktor, \u00f5ppejuht")

    role_block(doc, "4.10 Muudatuste projektijuht (Mari Vavulski)",
        "Nelja kooli \u00fchendamine - inimesed enne protsesse",
        ["Ei suru muudatust ilma selgituseta", "Ei j\u00e4ta muret vastuseta"],
        ["Koordineerib selgelt", "Kuulab muret", "Raporteerib juhtkonnale"],
        ['"Selgitan, mis muutub ja miks."', '"Sinu mure on registreeritud."'],
        "direktor, personalijuht")

    h(doc, "5. Konflikti lahendamise kord", 2)
    n(doc, "Rahulik vestlus kohapeal (kui turvaline)")
    n(doc, "Kirjalik p\u00f6\u00f6rdumine ametliku kanali kaudu")
    n(doc, "Vahendaja (personal, sotsiaalpedagoog)")
    n(doc, "Juhtkonna otsus dokumenteeritult")
    n(doc, "Valised kanalid ainult kui sisemine protsess ei t\u00f6\u00f6ta (t\u00f6\u00f6inspektsioon, Haridus- ja Noorteamet, politsei ohu korral)")

    h(doc, "6. Turvalisus", 2)
    b(doc, "T\u00f6\u00f6\u00f5nnetus v\u00f5i oht t\u00f6\u00f6kojas - peatada t\u00f6\u00f6, teatada kohe")
    b(doc, "Kiusamine - ei ole \"lapse asi\"; reageeritakse 24h jooksul")
    b(doc, "Agressioon - 112, turvat\u00f6\u00f6, juhtkond")
    b(doc, "Vaimne kriis - suunamine tugispetsialistile; 116 123")

    h(doc, "7. Kiire viide seinale", 2)
    p(doc, "Kuula -> Vasta -> Ole kohal -> Aita suunata -> Teata ohtu", bold=True, center=True)
    p(doc, "Techno TLN: Kvaliteet | K\u00e4ttesaadavus | Hoolivus", center=True)

    p(doc, "")
    p(doc, "Dokument on koostatud kooli p\u00f5hiv\u00e4\u00e4rtuste ja riigisektori lugupidava suhtluse raamistiku alusel.", italic=True, center=True)


def build_riigisektor_all(doc):
    h(doc, "RIIGISEKTORI ROLLIP\u00d5HINE AUSTUSE JA SUHTLUSE JUHEND", 1)
    p(doc, "Iga t\u00f6\u00f6koht, tasand ja roll | Eeskuju: Techno TLN", center=True)

    h(doc, "Sissejuhatus", 2)
    p(doc, "Riigisektoris on reeglid vajalikud, et oleks selgus ja turvalisus. See juhend kehtib haridusasutustes, ministeeriumides, ametites, haiglates, vallamajades ja k\u00f5ikjal, kus t\u00f6\u00f6tab avalik teenistus.")
    p(doc, "P\u00f5him\u00f5te: austus on k\u00e4itumine. Lugupidamine on protsess. Turvalisus on esimene.")

    h(doc, "Tasandid ja rollid", 2)

    roles = [
        ("JUHTKOND (direktor, sekret\u00e4r, n\u00f5ukogu esimees)",
         "Strateegia ja eeskuju",
         ["Avalik h\u00e4bistamine", "Otsuste p\u00f5hjendamata j\u00e4tmine", "Alla hinnata p\u00f6\u00f6rdumist"],
         ["Selgitab otsuseid", "V\u00f5tab vastutuse", "On k\u00e4ttesaadav"],
         "personal, j\u00e4relevalveasutus"),
        ("KESKTASE (osakonnajuht, projektijuht, Mari Vavulski t\u00fc\u00fcpi koordinaator)",
         "\u00dchendab inimesi ja protsesse",
         ["Kirja tegemine ilma tagasisideta", "\u00dclekoormamine ilma teatamiseta"],
         ["Koordineerib selgelt", "Kaitseb meeskonda", "Eskaleerib \u00f5igel ajal"],
         "\u00fclemus, personal"),
        ("ESILIIN (\u00f5petaja, ametnik, sotsiaalt\u00f6\u00f6taja, meditsiin)",
         "Teenuse ja \u00f5ppe kvaliteet",
         ["Ignorantsus kodaniku suhtes", "Kiusamise vaikimine"],
         ["Tervitab", "Kuulab", "Suunab edasi kui vaja"],
         "vahetu \u00fclemus"),
        ("TUGITEENUSED (IT, hooldus, turva, k\u00f6\u00f6k)",
         "Igap\u00e4evane kogemus",
         ["Naeratus oskamatuse \u00fcle", "\"Pole minu t\u00f6\u00f6\""],
         ["Suunab", "Teatab rikkest", "Aitab"],
         "vahetu \u00fclemus"),
        ("\u00d5PPIJA / TEENUSE SAAJA (\u00f5pilane, \u00fcli\u00f5pilane, kodanik)",
         "\u00d5igus lugupidamisele",
         ["Solvamine", "V\u00e4givald", "Salvestamine ilma loata"],
         ["K\u00fcsib abi", "J\u00e4rgib reegleid", "Teatab probleemist"],
         "esindus, \u00f5iguskantsler, AM"),
        ("PARTNER (ettev\u00f5te, MT\u00dc, lapsevanem)",
         "Koost\u00f6\u00f6 avaliku asutusega",
         ["Surve avaliku raha eest", "Ebakorrektne suhtlus t\u00f6\u00f6tajaga"],
         ["Lepib kokku", "Austab protsessi", "Kaebuse korral kirjalikult"],
         "lepinguj\u00e4rgne kontakt"),
        ("HARIDUS: Kaitsev\u00e4e Akadeemia (KVA)",
         "Rakendusk\u00f5rgkool - juhtimine ja distsipliin lugupidavalt",
         ["Avalik alandamine", "Kord ilma selgituseta"],
         ["Selgitab ootused", "Juhib eeskujuga", "Toetab kriitilist m\u00f5tlemist"],
         "rektor, \u00f5ppeosakond"),
        ("HARIDUS: Rocca al Mare Kool, Waldorf, Montessori",
         "Erinevad pedagoogilised mudelid - sama austus",
         ["Erandite kasutamine solvamiseks", "Lapse tausta \u00fcle naermine"],
         ["Kuulab lapsevanemat", "Kaitseb \u00f5ppija v\u00e4\u00e4rikust", "Teatab ohtudest"],
         "direktor, sotsiaalpedagoog"),
    ]

    for title, roll, ei, tee, raport in roles:
        role_block(doc, title, roll, ei, tee, ["Vt \u00fcldised p\u00f5him\u00f5tted"], raport)

    h(doc, "Techno TLN - t\u00e4ispakk", 2)
    p(doc, "Tallinna Tehnoloogiakolledzi detailne sisekorra eeskiri on eraldi dokumendis: techno-tln-austus-sisekorra-eeskiri.docx")
    p(doc, "Nelja kooli \u00fchendamine (Lasnam\u00e4e Mehaanikakool, Pol\u00fctehnikum, T\u00f6\u00f6stushariduskeskus, Ehituskool). \u00d5ppesuunad: inseneriharidus, ehitus, mehaanika, IT, loovtehnoloogia ja muu rakenduslik haridus.")

    h(doc, "\u00dchine kiirviide (k\u00f5ik rollid)", 2)
    p(doc, "1. Kas ma kuulan? 2. Kas ma alav\u00e4\u00e4ristan? 3. Kas on turvaline? 4. Kuhu raporteerin?", center=True)


def main():
    doc1 = Document()
    for s in doc1.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.5)
    build_techno_tln(doc1)
    doc1.save(OUT)

    doc2 = Document()
    for s in doc2.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.5)
    build_riigisektor_all(doc2)
    doc2.save(OUT_ALL)

    print(f"Saved: {OUT}")
    print(f"Saved: {OUT_ALL}")


if __name__ == "__main__":
    main()
