#!/usr/bin/env python3
"""Generate WARNO Operatsioon VABADUS DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = "/opt/cursor/artifacts/WARNO-Operatsioon-VABADUS.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # Title
    title = doc.add_heading("WARNO — Operatsioon „VABADUS“", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = [
        ("Klassifikatsioon:", "Isiklik / sisemine"),
        ("Kuupäev:", "Teisipäev"),
        ("Operatsioonipäev (D-day):", "Neljapäev"),
        ("Raamistik:", "Unplugged Alpha — lahutusejärgne vabadus, seksuaalne eelis, null läbirääkimised kohustusest või emotsionaalsest sõltuvusest"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value)

    doc.add_paragraph()

    # 1. OLUKORD
    add_heading(doc, "1. OLUKORD", 1)
    bullets_situation = [
        "Lahutusejärgne faas. Eelmine suhe (kõrge konfliktiga) on suletud; lapse (8-a) stabiilsus on mittemüüdav kapital.",
        "Uus dünaamika (Alisha): esimene suudlus ja tugev romantiline/seksuaalne avamine on tehtud — algne tõmbenõudlus on lahendatud.",
        "Eks-partner: raha surve, piiride test, lein — ei tohi langetada sinu väärtust.",
        "Kogutud väärtus: elutöö, kompetents, ressursid, distsipliin, otsustusvõime — põhivara. See ei ole lõpmatu ega automaatselt ülekantav.",
    ]
    for b in bullets_situation:
        doc.add_paragraph(b, style="List Bullet")

    add_para(doc, "Raamistik (The Value of Others — Orion Taraban):", bold=True)
    doc.add_paragraph(
        "Inimeste vahel vahetatakse väärtust — staatust, ressursse, seksi, tähelepanu, emotsionaalset regulatsiooni, "
        "tulevikupotentsiaali, sotsiaalset tõestust. Mitte „armastus võidab kõik“. Mitte „hingeline paar“."
    )
    doc.add_paragraph(
        "Pärast kõrge konfliktiga lahutust on kiusatus üle-investeerida järgmisse intensiivsesse ühendusse — "
        "see on klassikaline väärtuse valehindamine."
    )

    # 2. ÜLESANNE
    add_heading(doc, "2. ÜLESANNE", 1)
    add_para(doc, "Peamine:", bold=True)
    for b in [
        "Säilitada töö ja taastumine läbiräägitamatuna.",
        "Nautida uut dünaamikat ülejäägi, mitte nälja positsioonist.",
        "Mitte lasta eksi kaosul end alla tõmmata.",
        "Lapse stabiilsus — mitte vaieldav.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_para(doc, "Neljapäev (D-day):", bold=True)
    doc.add_paragraph(
        "Logistika ja olekuhaldus võidavad „mängu“ ööl endas. Jää lõdvestunuks, juhi logistikat, hoia füüsiline "
        "eskalatsioon varakult ja kalibreeritult. Ära üle-räägi. Ära üle-seleta. Ära tee sellest erakordse eriliigse esitluse ülevaatust."
    )

    add_para(doc, "Väljumisvõimalus (Exit optionality):", bold=True)
    doc.add_paragraph(
        "Naudi ööd, sest saad puhtalt lahkuda. Ei tulevikufäkse. Ei „ehitame midagi koos“. Ei emotsionaalset võlga. "
        "Kui energia nihkub — lahkud olekuga puutumata. See on tegelik eelis."
    )

    # 3. TÄITMINE
    add_heading(doc, "3. TÄITMINE", 1)

    add_heading(doc, "3.1 Täna (teisipäev) — toitumine (jõudlusfookus)", 2)
    doc.add_paragraph(
        "Eesmärk: puhas energia, vereringe, stabiilne testosterooni tugi, null puhitus, null krahh. "
        "Kerge enough, et keha jääks reageerivaks."
    )
    add_table(
        doc,
        ["Söögikord", "Sisu"],
        [
            (
                "Hommikusöök",
                "3–4 tervet muna (võis või oliiviõlis); suur peotäis spinatit või rukolit; pool avokaadot; "
                "väike peotäis kreeka pähkleid või kõrvitsaseemneid",
            ),
            (
                "Lõunasöök",
                "Grillitud või pannil küpsetatud lõhe (või kana); suur segasalat + peet või kurk; oliiviõli + sidrun; "
                "valikuline: väike portsjon maguskartulit või kinoa",
            ),
            (
                "Pärastlõuna (valikuline)",
                "Kreeka jogurt (täiskreem) + mõned mustikad või ruut 85%+ tumedat šokolaadi; VÕI peotäis mandleid",
            ),
            (
                "Õhtusöök (varakult, kergelt)",
                "Lahja valk (steik, kana, kala); palju rohelist köögivilja; oliiviõli või natuke avokaadot. "
                "EI: rasked süsivesikud, alkohol, praetud toit",
            ),
        ],
    )
    doc.add_paragraph("Vedelikud: vesi + näpuotsatäis soola. Kofeiin lõpetada pärastlõunaks, kui mõjub und.")
    p = doc.add_paragraph()
    r = p.add_run(
        "See ei ole maagia. See eemaldab asjad, mis tapavad erektsiooni kvaliteeti ja energiat "
        "(suhkru tipud, rasked toidud, alkohol), ja annab süsteemile materjale, mida see tegelikult kasutab."
    )
    r.italic = True

    add_heading(doc, "3.2 Täna + homme — keha ja taastumine", 2)
    add_table(
        doc,
        ["Tegevus", "Korraldus"],
        [
            ("Uni", "Prioriteet"),
            ("Toitumine", "Ülaltoodud plaan"),
            ("Liikumine", "Jalutuskäik või kerged raskused"),
            ("Porn", "Vältida raskeid sessioone, mis tõstavad dopamiini ja jätab tasaseks. Üks kontrollitud sessioon OK; edge-and-crash tsüklid EI"),
        ],
    )

    add_heading(doc, "3.3 Neljapäev (D-day) — keskkond", 2)
    add_table(
        doc,
        ["Element", "Standard"],
        [
            ("Puhastus", "Korralikult. Vannituba, voodipesu, pinnad"),
            ("Lõhn", "Neutraalne või kergelt meeldiv"),
            ("Varustus", "Kondoomid ja libesti juba käeulatuses — mitte viimasel hetkel otsida"),
            ("Eesmärk", "Mitte „tema mugavuse teater“ — et sina jääksid kehasse ega murraks olekut"),
        ],
    )

    add_heading(doc, "3.4 Neljapäev — raamistik ööl", 2)
    for b in [
        "Puudutus varakult ja tihti, loomulikult",
        "Ära üle-räägi / ära üle-seleta / ära tee esitlusest",
        "Juhi logistikat, hoia lõdvestust",
        "Ära aja segi intensiivsust kaldist (leverage). Intensiivsus on lahutuse järel odav ja tavaline. Puhas teostus, hea taastumine ja null vajadus — tegelik serv.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "3.5 Pärast ööd (D+1)", 2)
    add_table(
        doc,
        ["Tegevus", "Korraldus"],
        [
            ("Järeltekstid", "EI ärevaid järelkirju"),
            ("Üle-investeerimine", "EI"),
            ("Eks", "Ära loo uut kaost uues dünaamikas"),
            ("Laps", "Stabiilsus jääb prioriteediks"),
        ],
    )

    # 4. TAGALA
    add_heading(doc, "4. TAGALA / LOGISTIKA", 1)
    add_para(doc, "Mittevaieldavad ressursid:", bold=True)
    for b in ["Töö ja taastumine", "Lapse stabiilsus", "Põhivara (kompetents, ressursid, distsipliin)"]:
        doc.add_paragraph(b, style="List Number")

    add_para(doc, "Sekundaarne (vahetatav):", bold=True)
    doc.add_paragraph("Kõik muu, sh uus dünaamika — nautida ülejäägist, mitte nälja positsioonist.")

    add_para(doc, "Väärtusloogika (ilma moraalitsemata):", bold=True)
    for b in [
        "Sinu väärtus ≠ lõpmatu, ≠ automaatselt ülekantav",
        "Teiste väärtus kõigub — sõltub valikutest, vajadusest (neediness), kohtlemisest",
        "Seksuaalne vabadus ja eelis jäävad eeliseks ainult siis, kui ei vaheta põhivara ajutise intensiivsuse, valideerimise või põgenemise eest eksi kaosest",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 5. JUHTIMINE
    add_heading(doc, "5. JUHTIMINE JA KORD", 1)
    add_para(doc, "Operatiivprintsiibid (Extreme Ownership):", bold=True)
    for b in [
        "Töö ja taastumine — läbiräägitamatud",
        "Alisha — nautida ülejäägist, mitte nälja positsioonist",
        "Eks — ära lase kaosul väärtust alla tõmmata",
        "Laps — mittemüüdav kapital; kõik muu sekundaarne",
    ]:
        doc.add_paragraph(b, style="List Number")

    p = doc.add_paragraph()
    p.add_run("Vabadus on reaalne ainult seni, kuni sa ei loo uues dünaamikas uut kaost.").bold = True

    # 6. JÄRGMISED SAMMUD
    add_heading(doc, "6. JÄRGMISED SAMMUD (valikuline)", 1)
    doc.add_paragraph("Kui vaja täpsustust:")
    doc.add_paragraph(
        "• The Value of Others peatükkide/principle'ide lahtivõtt sinu olukorra vastu "
        "(uus naine, lahutus, hooldus/medatsioon, seksuaalstrateegia)"
    )
    doc.add_paragraph("• VÕI puhtalt praktiline: söögiajad, olekuhaldus, neljapäeva logistika")
    add_para(doc, "Sinu otsus.", bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WARNO lõpp.\nSaad seda nautida. Ära aja segi intensiivsust kaldist.")
    r.bold = True
    r.italic = True

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
