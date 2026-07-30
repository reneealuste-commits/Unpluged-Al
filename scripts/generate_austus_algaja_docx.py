#!/usr/bin/env python3
"""Generate beginner respect guide DOCX — self and partner, based on military regulations."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        run.font.color.rgb = RGBColor(0x3D, 0x6B, 0x4F)
    p.paragraph_format.space_before = Pt(14 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(8)


def body(doc, text, bold=False, italic=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    heading(doc, "Austamise juhised algajale", 1)
    body(doc, "Iseendale ja paarilisele - rivieskirjade jargi", italic=True, center=True)
    body(doc, "Kohandatud Kaitsevaee sisemaarustikust ja eetikakoodeksist", center=True)

    body(doc, "")
    body(doc, "OLULINE - loe enne alustamist", bold=True)
    bullet(doc, "See juhend toob sõjalised rivieskirjad igapäeva keelde - kodus, paarisuhtes, tööl.")
    bullet(doc, "Austus ei ole nõrkus. Rivieskirjade järgi on ausus ja distsipliin tugevus.")
    bullet(doc, "Ma nõuan, et minuga käitutakse austusega ja õiglaselt - ja annan seda tagasi.")
    bullet(doc, "Kui suhtes on vägivald, kontroll või hirm - see juhend ei piisa. Otsi abi.")

    heading(doc, "Mis on austus rivieskirjade järgi?", 2)
    body(doc, "Kaitseväe sisemäärustik ja eetikakoodeks ütlevad: kaitseväelane on korrektne, aus, distsiplineeritud ja austab teisi - sõltumata auastmest.")
    body(doc, "Sama kehtib kodus: austus on käitumine, mitte emotsioon. Sa võid vihastada ja ikka olla austav. Sa võid öelda ei ja ikka olla õiglane.")

    heading(doc, "Allikad (rivieskirjad)", 2)
    bullet(doc, "Kaitseväe sisemäärustik - käitumisreeglid, ausus, distsipliin")
    bullet(doc, "Kaitseväe eetikakoodeks - kuus põhiväärtust")
    bullet(doc, "Kaitseväelase tõotus - kohandatud kodaniku keelde")
    bullet(doc, "Vastastikmõju (tit-for-tat) - alusta heana, sea piir kurja korral")

    heading(doc, "1. Kuus väärtust - igapäevane distsipliin", 2)
    body(doc, "Kaitseväe eetikakoodeks. Kasuta neid nagu rivieskirja punkte.")

    heading(doc, "A. Ausus", 3)
    bullet(doc, "Räägi tõtt - endale ja teistele.")
    bullet(doc, "Ära naera kaaslase arvelt. Ära valet partnerile.")
    bullet(doc, "Kui eksid - ütle ausalt. Rivieskirjad: ausus on julgus.")

    heading(doc, "B. Vaprus", 3)
    bullet(doc, "Räägi välja, kui midagi on valesti - ilma kärata.")
    bullet(doc, "Vabanda, kui vaja. See on tugevus, mitte nõrkus.")
    bullet(doc, "Seisa oma sõna eest.")

    heading(doc, "C. Asjatundlikkus", 3)
    bullet(doc, "Tee oma osa hästi - kodus, tööl, suhtes.")
    bullet(doc, "Õpi. Paranda. Ära lase teistel kanda sinu kohustusi.")
    bullet(doc, "Ole kohal, kui lubasid.")

    heading(doc, "D. Ustavus", 3)
    bullet(doc, "Ole truu neile, kes sinust sõltuvad.")
    bullet(doc, "Ära räägi partneri selja taga halvustavalt.")
    bullet(doc, "Kaitse peret - ka siis, kui keegi teine ei vaata.")

    heading(doc, "E. Koostöövalmidus", 3)
    bullet(doc, "Aita, kui saad. Jaga koormat.")
    bullet(doc, "Küsi: mida saan teha, mitte mida sina peaksid tegema.")
    bullet(doc, "Üks meeskond - ka kodus.")

    heading(doc, "F. Avatus", 3)
    bullet(doc, "Kuula enne vastamist.")
    bullet(doc, "Tunnista, kui eksid.")
    bullet(doc, "Jäta ruumi teise arvamusele.")

    heading(doc, "2. Valmistumine (mõlemale)", 2)
    numbered(doc, "Vali rahulik hetk. Mitte tülise ajal.")
    numbered(doc, "Lepi kokku: me räägime austusest, mitte süüdistame.")
    numbered(doc, "Igaüks kirjutab ühe lause: mida ma täna austuse all mõtlen?")
    numbered(doc, "Hinda skaalal 0-10: kui austatud ma end praegu tunnen? (0 = üldse mitte, 10 = täielikult)")
    numbered(doc, "Sea aeg: 15-20 minutit.")

    heading(doc, "3. Iseendale - samm-sammult", 2)

    heading(doc, "Samm 1: Eneseaustus (3 min)", 3)
    body(doc, 'Küsi endalt: "Kas ma kohtlen ennast nii, nagu rivieskirjad käsivad kohelda kaaslasi?"')
    bullet(doc, "Kas ma luban endale puhata?")
    bullet(doc, "Kas ma räägin endaga ausalt?")
    bullet(doc, "Kas ma täidan oma lubadusi iseendale?")

    heading(doc, "Samm 2: Üks austusotsus (2 min)", 3)
    body(doc, "Vali üks tegu täna, mis näitab austust:")
    bullet(doc, "Helistan, kui lubasin.")
    bullet(doc, "Kuulan lõpuni, ilma telefonita.")
    bullet(doc, "Ütlen täpselt, mida tunnen - ilma solvamata.")
    bullet(doc, "Teen ühe asja, mida teine palus.")

    heading(doc, "Samm 3: Piir, kui vaja (5 min)", 3)
    body(doc, "Rivieskirjade järgi: austus töötab mõlemapoolselt.")
    numbered(doc, 'Kui keegi ei austa sind - ütle selgelt: "Minuga tuleb käituda austusega."')
    numbered(doc, "Ära vihasta ega põgene. Üks lause. Siis vaikust.")
    numbered(doc, "Kui kurjus kordub - sea piir. Vastastikmõju: üks hoiatus, siis kaitse.")
    numbered(doc, "Mida külvad, seda lõikad - ära külvata halba, aga ära lase end kurjaks teha.")

    heading(doc, "Samm 4: Kontroll (2 min)", 3)
    body(doc, "Hinda uuesti: austatus 0-10? Kirjuta üks lause: mis muutus?")

    heading(doc, "4. Paarilisele - austuse protokoll", 2)
    body(doc, "Sa ei pea olema terapeut. Sa oled kaaslane, kes järgib samu rivieskirju.")

    heading(doc, "Abilise / partneri roll", 3)
    bullet(doc, "Kuula ilma katkestamata - vähemalt 60 sekundit.")
    bullet(doc, 'Korda tagasi üks lause: "Ma kuulsin, et..." - mitte "aga sina..."')
    bullet(doc, "Küsi: mida sa vajad minult täna?")
    bullet(doc, "Anna aus vastus - mitte mugav vastus.")

    heading(doc, "Paariline protokoll (15-20 min)", 3)
    numbered(doc, "Istuge näost näkku. Telefonid ära.")
    numbered(doc, 'A: "Mida ma täna austuse all mõtlen on..." (1 lause)')
    numbered(doc, 'B: kordab tagasi. Siis B sama.')
    numbered(doc, 'A: "Üks asi, mida ma vajan austust..." B vastab: "Ma saan teha..."')
    numbered(doc, "Mõlemad: austatus 0-10 enne ja pärast - kirjutage üles.")
    numbered(doc, "Lõpetage: üks konkreetne tegu 24 tunni jooksul.")

    heading(doc, "Mida partner EI tee", 3)
    bullet(doc, 'Ei solva, ei naera, ei alaväärista ("sa oled liiga tundlik").')
    bullet(doc, "Ei võta austuse nõuet solvanguna.")
    bullet(doc, "Ei jäta vahele - kui teine palub austust, see on õigus, mitte rünnak.")
    bullet(doc, "Ei kasuta vaikimist karistusena.")

    heading(doc, "5. Austuse keel - mida öelda", 2)
    body(doc, "Positiivne:", bold=True)
    bullet(doc, '"Ma kuulen sind."')
    bullet(doc, '"Sa oled mulle oluline."')
    bullet(doc, '"Ma vabandan."')
    bullet(doc, '"Aitäh, et ütlesid."')
    bullet(doc, '"Kuidas saan sind toetada?"')

    body(doc, "Piir:", bold=True)
    bullet(doc, '"Minuga tuleb käituda austusega."')
    bullet(doc, '"Ma ei räägi nii, kui sa karjud."')
    bullet(doc, '"Ma vajan pausi. Tulen tagasi 30 minuti pärast."')
    bullet(doc, '"See on minu piir."')

    heading(doc, "6. Millal STOP", 2)
    bullet(doc, "Karjumine, solvamine, ähvardamine")
    bullet(doc, "Füüsiline vägivald või selle oht")
    bullet(doc, "Kontroll (telefon, raha, suhted)")
    bullet(doc, "Austatus jääb alla 3/10 pidevalt")
    body(doc, "STOP korral: lahku ruumist. Kõne lähivõrgustikuga. Kriisiabi: 116 123 (Eluliin). Kui oht - 112.")

    heading(doc, "7. Kiire viide", 2)
    body(doc, "Valmista → Eneseaustus → Üks tegu → Piir kui vaja → Kontroll 0-10", bold=True, center=True)
    body(doc, "Paar: kuula → korda tagasi → küsi vajadust → üks tegu 24h", bold=True, center=True)
    body(doc, "Kuus väärtust: ausus · vaprus · asjatundlikkus · ustavus · koostöö · avatus", center=True)

    heading(doc, "8. Näited algajale", 2)
    body(doc, "Ise - töö:", bold=True)
    body(doc, "Kolleeg saadab pikalt sõnumi. Sa vastad sama ausalt ja lühidalt. Ausus. Asjatundlikkus.")
    body(doc, "Paar - tüli:", bold=True)
    body(doc, 'Üks ütleb: "Ma tundsin, et sa mind ei kuulanud." Teine: "Ma kuulsin, et sa tundsid end eiranuna." Pause. Austatus 4 -> 7.')
    body(doc, "Ise - piir:", bold=True)
    body(doc, 'Sõber naerab sinu arvelt. Sa ütled: "Minuga tuleb käituda austusega." Ei vihasta. Piir seatud.')
    body(doc, "Paar - laps:", bold=True)
    body(doc, "Laps tahab tähelepanu. Sa paned telefoni ära. 10 minutit silma silma. Ustavus ja koostöö.")

    heading(doc, "9. Kodaniku lause (rivieskirjast)", 2)
    body(doc, 'Ma nõuan, et minuga käitutakse austusega ja õiglaselt - ja käitun teistega samamoodi. Mida külvad, seda lõikad. Alustan heana, kuid ei lase kurja minna vastutuseta.', italic=True, center=True)

    body(doc, "")
    body(doc, "Kohandatud Kaitseväe sisemäärustikust ja eetikakoodeksist. Hariduslik juhend.", italic=True, center=True)
    body(doc, "Unpluged-Al · Austamine algajale", center=True)

    out = "/workspace/austus-algaja-juhised.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
