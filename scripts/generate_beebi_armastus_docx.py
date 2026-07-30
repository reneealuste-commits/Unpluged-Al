#!/usr/bin/env python3
"""Generate baby love guide DOCX — rivieskirjade stiilis, iseendale ja paarilisele."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        run.font.color.rgb = RGBColor(0x8B, 0x45, 0x6B)
    p.paragraph_format.space_before = Pt(14 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(8)


def body(doc, text, bold=False, italic=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    heading(doc, "Beebide armastamise juhised", 1)
    body(doc, "Iseendale ja paarilisele - rivieskirjade jargi", italic=True, center=True)
    body(doc, "Kohandatud austuse ja kuulamise pohimotetest - beebi keeles", center=True)

    body(doc, "")
    body(doc, "OLULINE - loe enne alustamist", bold=True)
    bullet(doc, "Beebi ei saa armastust sonadest. Beebi saab armastust sinu kehast, hoolest ja kohalolekust.")
    bullet(doc, "Armastus beebi suhtes on distsipliin: iga paev, isegi kui oled vasinud.")
    bullet(doc, "Mõlemad vanemad/partnerid on meeskond. Üks ei pea kandma kõike.")
    bullet(doc, "Kui tunned end uputatuna, depressiivsena või vägivaldsena - otsi abi. See on tugevus.")

    heading(doc, "Mis on beebi armastamine?", 2)
    body(doc, "Rivieskirjade loogika: austus on käitumine. Beebi puhul on armastus samuti käitumine - mitte emotsioon, mida tunned alati, vaid tegu, mida teed.")
    body(doc, "Beebi armastamine tähendab: kuulad, vastad, oled kohal, kaitstad, hoiad rahulikult.")

    heading(doc, "1. Millised tegevused on MITTE-armastavad?", 2)
    body(doc, "Vastus nagu rivieskirjas: mida sa EI tohiks teha.", bold=True)
    bullet(doc, "Ei kuula, kui beebi nutab või otsib silmakontakti")
    bullet(doc, "Ei vasta vajadusele (nälg, mähkmed, uni, kaisus)")
    bullet(doc, "Telefon käes, kui beebi on lähedal")
    bullet(doc, "Karjud beebi peale (beebi ei tea, mida ta tegi valesti)")
    bullet(doc, "Raputad tugevalt - KUNAGI")
    bullet(doc, "Jättad beebi pikaks ajaks üksi nutma \"harjutamiseks\"")
    bullet(doc, "Naerad või alavääristad teise vanema muresid")
    bullet(doc, "Oled kohal kehaga, aga mitte vaimuga")

    heading(doc, "2. Kuidas ma naitan armastust beebile?", 2)
    body(doc, "Vastus nagu rivieskirjas: AUSTA TEISI - kohandatud beebile.", bold=True)

    heading(doc, "A. Austav suhtlus (verbaalne)", 3)
    bullet(doc, "Räägi rahuliku, soe häälega - isegi kui beebi ei saa sõnadest aru")
    bullet(doc, "Nimeta beebi nimega. Ütle: \"Ma olen siin.\" \"Ma kuulan.\"")
    bullet(doc, "Laula, sosista, jutusta - hääl on turvalisus")

    heading(doc, "B. Hea taktitunne, toon, edastus ja ajastus", 3)
    bullet(doc, "Vaata, mis beebil vaja on - mitte mis sul endal parasjagu")
    bullet(doc, "Kui beebi on üle väsinud - vaikust, mitte mängu")
    bullet(doc, "Kui beebi on ärkvel ja aktiivne - suhtle, naerata, mängi")

    heading(doc, "C. Visuaalsed signaalid", 3)
    bullet(doc, "Silmakontakt - beebi näeb su silmi")
    bullet(doc, "Naeratus - beebi peegeldab")
    bullet(doc, "Avatud kehakeel - pööra beebi poole, mitte eemale")
    bullet(doc, "Kanna beebi turvaliselt - toetus kaelale ja peale")

    heading(doc, "D. Kohanda oma suhtlust", 3)
    bullet(doc, "Iga beebi on erinev - üks tahab rohkem kaisu, teine rohkem ruumi")
    bullet(doc, "Jälgi signaale: keerab ära, käed lahti, nutab = paus")
    bullet(doc, "Jälgi signaale: vaatab sind, rahuneb su häälel = jätka")

    heading(doc, "3. Kuulamine - beebi keel", 2)
    body(doc, "Millised tegevused naitavad, et sa EI kuula beebi?", bold=True)
    bullet(doc, "Nutab ja sa jätkad telefoniga")
    bullet(doc, "Keerab pea ära ja sa sunnid edasi")
    bullet(doc, "Rahuneb kaisus ja sa paned kohe maha")
    bullet(doc, "Ignorad öösel - \"ta lihtsalt manipuleerib\"")

    body(doc, "Kuidas kuulata beebi:", bold=True)
    numbered(doc, "Peatu. Vaata beebi keha.")
    numbered(doc, "Küsi: nälg? mähkmed? uni? liiga palju? liiga vähe?")
    numbered(doc, "Vasta ühe asjaga korraga.")
    numbered(doc, "Kui ei tea - kaisus, rütm, hääl. See on OK.")

    heading(doc, "4. Valmistumine (mõlemale vanemale)", 2)
    numbered(doc, "Lepi kokku: me oleme meeskond, mitte võistlevad pooled.")
    numbered(doc, "Jagage ööd ja päeva - konkreetne plaan, mitte \"sa peaksid teadma\".")
    numbered(doc, "Igaüks kirjutab: üks asi, mida ma täna beebile annan.")
    numbered(doc, "Hinda 0-10: kui kohal ma täna olen? (füüsiliselt ja vaimselt)")

    heading(doc, "5. Iseendale - igapäevane protokoll", 2)

    heading(doc, "Hommik (5 min)", 3)
    bullet(doc, "Tervita beebi nimega. Silmakontakt.")
    bullet(doc, "Kontrolli: nälg, mähkmed, uni.")
    bullet(doc, "Üks puudutus: pait, musi, kaisus.")

    heading(doc, "Päeva jooksul", 3)
    bullet(doc, "Telefon ära, kui söötmine või mängimine.")
    bullet(doc, "Räägi, mida teed: \"Nüüd vahetame mähkmeid.\"")
    bullet(doc, "Vähemalt 10 minutit täielikku tähelepanu - ilma ekraanita.")

    heading(doc, "Õhtu (10 min)", 3)
    bullet(doc, "Rahulik rutiin: vann / puhastus / lugu / laul / pime.")
    bullet(doc, "Sama järjekord iga õhtu - beebi tunneb turvalisust.")
    bullet(doc, "Ütle: \"Ma armastan sind. Sa oled turvaliselt.\"")

    heading(doc, "Kui oled väsinud (rivieskirja distsipliin)", 3)
    body(doc, "Väsimus ei ole vabandus halbadele tegudele. See on signaal küsida abi.")
    bullet(doc, "Pane beebi turvaliselt maha. Võta 5 sügavat hingetõmmet.")
    bullet(doc, "Kutsu partner appi. Üks lause piisab.")
    bullet(doc, "Kui tunned viha beebi suhtes - pane ta turvaliselt maha ja lahku ruumist.")

    heading(doc, "6. Paarilisele - meeskonnaprotokoll", 2)
    body(doc, "Teie kaks olete beebi esimene maailm. Mõlemad peavad olema meeskond.")

    heading(doc, "Partneri roll", 3)
    bullet(doc, "Kui teine on väsinud - võta beebi ilma süüdistuseta.")
    bullet(doc, "Tee üks asi iga päev beebiga iseseisvalt (mähkmed, õhtu, jalutuskäik).")
    bullet(doc, "Kiida teist vanemat: \"Sa teed hästi.\"")
    bullet(doc, "Ära korda teise vanema ees. Eraviisiliselt, kui vaja.")

    heading(doc, "Paariline protokoll (15 min, kord nädalas)", 3)
    numbered(doc, "Beebi magab või on teisega. Istuge rahulikult.")
    numbered(doc, 'A: "Üks asi, mida ma vajan sinult beebi juures on..."')
    numbered(doc, 'B: vastab: "Ma saan teha..." - üks konkreetne tegu.')
    numbered(doc, "Mõlemad: mis läheb hästi? Üks lause.")
    numbered(doc, "Mõlemad: mis on raske? Üks lause. Ilma lahendamata kohe.")

    heading(doc, "Mida partner EI tee", 3)
    bullet(doc, "Ei võrdle: \"Mina ei nutaks nii\"")
    bullet(doc, "Ei jätka telefoniga, kui teine vajab abi")
    bullet(doc, "Ei jäta teist vanemat üksi öösel ilma kokkuleppeta")
    bullet(doc, "Ei naera beebi või teise vanema üle")

    heading(doc, "7. Armastuse keel - mida öelda ja teha", 2)
    body(doc, "Sõnad (isegi kui beebi ei saa aru):", bold=True)
    bullet(doc, '"Ma olen siin."')
    bullet(doc, '"Sa oled turvaliselt."')
    bullet(doc, '"Ma kuulan sind."')
    bullet(doc, '"Ma armastan sind."')
    bullet(doc, '"Sa oled mulle oluline."')

    body(doc, "Teod:", bold=True)
    bullet(doc, "Kaisus ja nahakontakt")
    bullet(doc, "Rinnaga või pudeliga toitmine rahulikult")
    bullet(doc, "Rütmiline kõikumine")
    bullet(doc, "Peopesaga pea toetamine")
    bullet(doc, "Öösel tulemine, kui beebi kutsub")

    heading(doc, "8. Millal STOP ja otsi abi", 2)
    bullet(doc, "Tunned viha beebi suhtes - pane turvaliselt maha, lahku, helista kellelegi")
    bullet(doc, "Raputad või tahad raputada - STOP KOHE. Oht beebi elule.")
    bullet(doc, "Depressioon, lootusetus, enesetapumõtted")
    bullet(doc, "Partner ei aita, süüdistab, on vägivaldne")
    body(doc, "Abi: Perearst · 116 123 (Eluliin) · 112 (hädaolukord) · Naiste tugikeskus 1492")

    heading(doc, "9. Kiire viide", 2)
    body(doc, "Kuula → Vasta vajadusele → Ole kohal → Rutiin → Meeskond", bold=True, center=True)
    body(doc, "Beebi armastus = tähelepanu + turvalisus + rahu + järjepidevus", center=True)
    body(doc, "Telefon ära · Silmad beebile · Üks tegu korraga", center=True)

    heading(doc, "10. Näited", 2)
    body(doc, "Väsinud öö:", bold=True)
    body(doc, "Beebi nutab 3. kord. Sa paned telefoni ära. Võtad kaisu. Räägid rahulikult. 10 minuti pärast rahuneb. See on armastus.")
    body(doc, "Partner appi:", bold=True)
    body(doc, "Ema on läbi. Isa ütleb: \"Ma võtan järgmised 2 tundi.\" Ilma \"miks sa ei saa hakkama\". Meeskond.")
    body(doc, "Päevane 10 min:", bold=True)
    body(doc, "Beebi põrandal. Sa lamad kõrval. Jälgi. Naeratad. Räägid. Telefon teises toas. Kvaliteetaeg.")
    body(doc, "Vale vs õige:", bold=True)
    body(doc, 'Vale: "Ta manipuleerib mind." Õige: "Tal on vajadus. Ma uurin, mis see on."')

    heading(doc, "11. Rivieskirja lause beebile", 2)
    body(doc, "Ma olen siin. Ma kuulan. Ma kaitstan. Ma armastan sind tegudega - iga päev, isegi kui olen väsinud. Sa oled turvaliselt.", italic=True, center=True)

    body(doc, "")
    body(doc, "Kohandatud austuse rivieskirjade loogikast (Combat Ready / inimesekeskne juhtimine). Hariduslik juhend.", italic=True, center=True)
    body(doc, "Unpluged-Al · Beebide armastamine", center=True)

    out = "/workspace/beebi-armastamise-juhised.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
