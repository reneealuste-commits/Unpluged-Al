#!/usr/bin/env python3
"""Generate beginner EMDR instructions DOCX — self and partner."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    p.paragraph_format.space_before = Pt(14 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(8)
    return p


def body(doc, text, bold=False, italic=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    heading(doc, "EMDR juhised algajale", 1)
    body(doc, "Iseendale ja paarilisele kodus", italic=True, center=True)
    body(doc, "Lihtsad sammud bilateral stimulatsiooniga — 15–20 minutit", center=True)

    heading(doc, "Lugu — kuidas EMDR sündis", 2)
    body(doc, "Mai 1987. Üks jalutuskäik pargis.", bold=True)
    body(doc, "Francine Shapiro oli psühholoogia doktorant, kes otsis dissertatsiooni teemat ja katsetas oma kehaga nagu laboratooriumis. Ühel päeval jalutaski ta pargis ja märkas midagi imelikku: häiriv mõte tuli peas — ja siis kadus. Ilma igasuguse pingutuseta. Ta pole maha maganud, ta ei unustanud. See lihtsalt kaotas oma haava.")
    body(doc, "Ta hakkas tähele panna. Iga kord, kui selline mõte tuli, liikusid silmad kergelt, kiiresti — vasakule, paremale, ja tagasi. Ja mõte nõrgene. Kadus.")
    body(doc, "Siis ta proovis teadlikult. Tõi meelde midagi, mis teda häiris. Liigutas silmi sama viisil. Ja see toimus uuesti. Emotsioon lahjenes. Mõte oli ikka seal — aga enam ei haavanud.")
    body(doc, 'Francine ei leiutanud midagi uut. Ta märkas, mis kehas juba toimub — ja hakkas seda tahtlikult kasutama. Ta kutsus sõpru: "Too midagi, mis sind häirib." Juhendas neid liigutama silmi. See töötas ka nemal.')
    body(doc, "Kuue kuuga katsetas ta umbes 70 inimesega. 1989. aastal avaldas ta esimese kontrollitud uuringu. Täna on EMDR üks maailma enim uuritud traumateraapia meetodeid — ja see algas ühest jalutuskäikust.")

    heading(doc, "Miks see sind puudutab?", 2)
    body(doc, "Sa ei pea olema katki, et seda proovida. Sa ei pea aru saama, kuidas see töötab. Piisab 15 minutist.")
    body(doc, 'Sinu aju töötleb igal ööl unes mõtteid — silmad liiguvad REM une ajal vasakule ja paremale. See on looduslik. Aga mõnikord jääb mõte kinni: homme esitlus, tüliline lause, pingutus kaelas. See tuleb tagasi ja tagasi, isegi kui sa "ei taha sellele mõelda."')
    body(doc, "EMDR kodune versioon aitab ajul seda kinniolevat mõtet liikuma lükata — nagu Shapiro silmad pargis jalutades tegid. Mitte sellepärast, et sa oleksid haige. Sellepärast, et su keha juba teab, kuidas tulla rahule. Sa annad sellele 15 minutit ja natuke abi.")
    body(doc, "Sa ei pea uskuma mind. Proovi ühe korra. 15 minutit. Üks kerge mõte. Ja vaata, kas su keha vastab.", italic=True, center=True)

    heading(doc, "Mis on EMDR?", 2)
    body(doc, "EMDR (Eye Movement Desensitization and Reprocessing) kasutab bilateral stimulatsiooni — vahelduvat signaali mõlemale poole keha (silma liigutused, koputused, heli). See aitab ajul töödelda häirivaid mõtteid ja emotsioone loomulikul viisil — mitte surudes, mitte unustades, vaid läbi töötades.")
    body(doc, "Kodus saad kasutada lihtsustatud versiooni. See ei asenda terapeuti raske trauma korral, aga võib aidata kerge pinge, ärevuse ja uneprobleemide juures.")

    body(doc, "")
    p = body(doc, "OLULINE — loe enne alustamist", bold=True)
    bullet(doc, "See juhend on kerge stressi, mure ja igapäevase pingete jaoks — mitte trauma või PTSD raviks.")
    bullet(doc, "Kui sul on raske trauma, vägistamine, peksmine, sõjalaagrid, enesetapumõtted või dissotsiatsioon — pöördu EMDR-terapeudi poole.")
    bullet(doc, "Kui distress tõuseb üle 8/10, lõpeta kohe ja kasuta maandamist (lõpus).")
    bullet(doc, "Paariline töö nõuab usaldust. Te ei pea jagama kogu lugu — piisab ühest kergest teemast.")

    heading(doc, "1. Valmistumine (mõlemale)", 2)
    numbered(doc, "Vali vaikne koht. Istu mugavalt. Vesi lähedal.")
    numbered(doc, "Vali meetod: silmaliigutused, koputused või audio (kõrvaklapid).")
    numbered(doc, "Hinda distressi skaalal 0–10 (0 = rahulik, 10 = kõige hullem). Kirjuta number üles.")
    numbered(doc, "Vali turvakoht — kujuteldav koht, kus tunned end turvaliselt (rannas, metsas, kodus).")
    numbered(doc, "Sea taimer 15–20 minutile.")

    heading(doc, "2. Meetodid — vali üks", 2)

    heading(doc, "A. Silmaliigutused", 3)
    bullet(doc, "Tõsta kaks sõrme silmade kõrgusele, õlalaiuselt.")
    bullet(doc, "Liiguta silmi sujuvalt vasakult paremale ja tagasi.")
    bullet(doc, "Pea paigal. Tempo: umbes 1 sekund kummalegi poole.")
    bullet(doc, "Üks komplekt = 24–30 liigutust (12–15 korda edasi-tagasi).")

    heading(doc, "B. Koputused", 3)
    bullet(doc, "Istudes: koputa vaheldumisi vasakut ja paremat põlve.")
    bullet(doc, "Või risti käed ja koputa õlgu: vasak, parem, vasak…")
    bullet(doc, "Aeglane rütm: umbes 1 koputus sekundis.")
    bullet(doc, "Üks komplekt = 24–30 koputust.")

    heading(doc, "C. Audio (kõrvaklapid)", 3)
    bullet(doc, "Kasuta bilateral heli/rakendust, mis vaheldab heli vasakul ja paremal.")
    bullet(doc, "Madal helitugevus. Üks komplekt = 1–2 minutit.")

    heading(doc, "3. Iseendale — samm-sammult", 2)

    heading(doc, "Samm 1: Turvakoht (2–3 min)", 3)
    body(doc, "Sule silmad. Kujutle turvalist kohta. Mida sa näed, kuuled, tunned? Püsi seal, kuni kehas on väike rahutunne.")

    heading(doc, "Samm 2: Vali üks sihtmärk (1 min)", 3)
    body(doc, 'Üks asi korraga: üks mõte, üks pilt või üks keha tunne. Näiteks: "homme töökoosolek", "pinge rinnus", "tüliline lause".')
    body(doc, "Ära vali kõige raskemat mälestust. Algaja jaoks: kerge ärevus või tüli, mitte trauma.")

    heading(doc, "Samm 3: Bilateral komplektid (10–15 min)", 3)
    numbered(doc, "Too sihtmärk meelde.")
    numbered(doc, "Alusta bilateral stimulatsiooni (24–30 liigutust/koputust).")
    numbered(doc, "Peatu. Märka, mis tuleb — mõte, pilt, tunne. Ära sunni.")
    numbered(doc, "Korda 3–6 komplekti.")
    numbered(doc, "Iga 2 komplekti järel küsi: mis on distress nüüd (0–10)?")

    heading(doc, "Samm 4: Positiivne uskumus (valikuline)", 3)
    body(doc, 'Vali lause, mida tahaksid uskuda: "Ma saan hakkama", "Ma olen turvaliselt", "See möödub".')
    body(doc, "Hoia seda meeles ja tee veel 2 bilateral komplekti.")

    heading(doc, "Samm 5: Lõpeta", 3)
    bullet(doc, "3 aeglast hingetõmmet.")
    bullet(doc, "Joo vett.")
    bullet(doc, "Kirjuta üks lause: mis muutus?")

    heading(doc, "4. Paarilisele — kuidas aidata", 2)
    body(doc, "Sa ei pea olema terapeut. Sa oled abiline, kes juhib rütmi ja hoiab ruumi turvalisena.")

    heading(doc, "Abilise roll", 3)
    bullet(doc, "Juhi bilateral stimulatsiooni (kui partner valib koputused — sina koputad tema õlgu vaheldumisi).")
    bullet(doc, 'Ütle vaikselt: "Järgmine komplekt", "Peatu", "Kuidas on nüüd?"')
    bullet(doc, 'Ära küsi "miks" ega "mis juhtus". Küsi: "Mis tuleb nüüd?" või "Mis on distress 0-10?"')
    bullet(doc, "Kui partner vaikib — see on OK. Ära täida vaikust.")

    heading(doc, "Paariline protokoll", 3)
    numbered(doc, "Lepi kokku: üks teema, 15–20 min, mõlemad rahulikud.")
    numbered(doc, "Partner A teeb tööd. Partner B juhib koputusi või liigutab sõrme silmade ees.")
    numbered(doc, 'B: "Too meelde oma sihtmärk. Alustame." - 24-30 koputust/liigutust.')
    numbered(doc, 'B: "Peatu. Mis tuli?" - A vastab lühidalt või ütleb "ei tea" - mõlemad OK.')
    numbered(doc, 'B: "Distress 0-10?" - kirjuta üles.')
    numbered(doc, "Korda 3–6 korda. Lõpus: vesi, 3 hingetõmmet.")
    numbered(doc, "Vaheta rollid järgmisel korral, kui mõlemad tahavad.")

    heading(doc, "Mida abiline EI tee", 3)
    bullet(doc, 'Ei analüüsi ega anna nõu ("sa peaksid lihtsalt...").')
    bullet(doc, "Ei suru rääkima.")
    bullet(doc, "Ei jätka, kui partner ütleb stop.")
    bullet(doc, "Ei tee kerge teemaga alustades rasketeks.")

    heading(doc, "5. Millal STOP", 2)
    bullet(doc, "Distress üle 8/10")
    bullet(doc, "Pearinglus, iiveldus, paanika")
    bullet(doc, 'Tunne, et "lahutud" kehast (dissotsiatsioon)')
    bullet(doc, "Tugevad flashback'id")
    body(doc, "STOP korral: lõpeta bilateral. Nimi 5 asja, mida näed. 4 mida tunned. 3 mida kuuled. 2 mida lõhnad. 1 mida maitsed. Külma vett. Jalad põrandale.")

    heading(doc, "6. Kiire viide", 2)
    body(doc, "Turvakoht → Sihtmärk → Bilateral 24–30 → Peatu → Distress? → Korda → Positiivne lause → Lõpp", bold=True, center=True)
    body(doc, "Aeg: 15–20 min | Stop kui: distress > 8 | Meetodid: silmad · koputused · audio", center=True)

    heading(doc, "7. Näited algajale", 2)
    body(doc, "Ise — uneärevus:", bold=True)
    body(doc, 'Sihtmärk: "homme esitlus". Distress 6. 4 komplekti koputusi. Distress 3. Lõpp.')
    body(doc, "Paar — kerge tüli:", bold=True)
    body(doc, "Abiline koputab õlgu. Töötegija toob meelde ühe lause tülis. 3 komplekti. Distress langeb. Vabandust ei pea — ainult töötlus.")
    body(doc, "Ise — kehatunne:", bold=True)
    body(doc, "Sihtmärk: pingutus kaelas. Silmaliigutused. 5 komplekti. Kael lõdveneb natuke.")

    body(doc, "")
    body(doc, "See juhend on hariduslik. Ei asenda psühholoogi või EMDR-terapeudi abi.", italic=True, center=True)
    body(doc, "Unpluged-Al · EMDR algajale", center=True)

    out = "/workspace/emdr-algaja-juhised.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
