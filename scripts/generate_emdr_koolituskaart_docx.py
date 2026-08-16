#!/usr/bin/env python3
"""Generate EMDR algaja koolituskaart DOCX — KK1 format."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "opord" / "kursid" / "EMDR-algaja-koolituskaart.docx"

ASCII_CARD = """╔══════════════════════════════════════════════════════════════════╗
║  EMDR1 — ALGAJA JUHISED                        15–20 min         ║
║  Unpluged-Al · iseendale ja paarilisele kodus                   ║
╠══════════════════════════════════════════════════════════════════╣
║  OLULINE — loe enne alustamist                                   ║
║  Kerge stress, mure, pinged — MITTE trauma/PTSD.                 ║
║  STOP kohe kui distress > 8/10.                                  ║
╠══════════════════════════════════════════════════════════════════╣
║  VAHENDID          │  MEETODID (vali üks)                        ║
║  vaikne koht, vesi │  A. Silmaliigutused — 24–30 liigutust      ║
║  taimer 15–20 min  │  B. Koputused (põlved/õlad) — 24–30         ║
║  paber distressile │  C. Audio kõrvaklapid — 1–2 min komplekt    ║
╠══════════════════════════════════════════════════════════════════╣
║  PROTSESS (kiire viide)                                          ║
║  Turvakoht → Sihtmärk → Bilateral 24–30 → Peatu → Distress?    ║
║  → Korda 3–6× → Positiivne lause → Lõpp (3 hingetõmmet, vesi)  ║
╠══════════════════════════════════════════════════════════════════╣
║  STOP KORRAL                                                     ║
║  Lõpeta bilateral. 5-4-3-2-1 maandamine. Külma vett.             ║
║  Jalad põrandale. Pöördu terapeudi poole kui vaja.               ║
╚══════════════════════════════════════════════════════════════════╝"""

WORKSHEET = """MINU EMDR SEANSS — ________________ (kuupäev)

Distress ENNE: ___/10          Distress PÄRAST: ___/10

Sihtmärk (üks lause): _________________________________________

Meetod:  □ Silmaliigutused   □ Koputused   □ Audio

Komplekte tehtud: ___

Mis muutus (üks lause): _______________________________________

□ Jõin vett    □ 3 hingetõmmet    □ Kirjutasin üles"""

STORY_SECTIONS = [
    ("Lugu — kuidas EMDR sündis", [
        "Mai 1987. Üks jalutuskäik pargis.",
        "Francine Shapiro oli psühholoogia doktorant, kes otsis dissertatsiooni teemat ja katsetas oma kehaga nagu laboratooriumis. Ühel päeval jalutaski ta pargis ja märkas midagi imelikku: häiriv mõte tuli peas — ja siis kadus. Ilma igasuguse pingutuseta. Ta pole maha maganud, ta ei unustanud. See lihtsalt kaotas oma haava.",
        "Ta hakkas tähele panna. Iga kord, kui selline mõte tuli, liikusid silmad kergelt, kiiresti — vasakule, paremale, ja tagasi. Ja mõte nõrgene. Kadus.",
        "Siis ta proovis teadlikult. Tõi meelde midagi, mis teda häiris. Liigutas silmi sama viisil. Ja see toimus uuesti. Emotsioon lahjenes. Mõte oli ikka seal — aga enam ei haavanud.",
        'Francine ei leiutanud midagi uut. Ta märkas, mis kehas juba toimub — ja hakkas seda tahtlikult kasutama. Ta kutsus sõpru: "Too midagi, mis sind häirib." Juhendas neid liigutama silmi. See töötas ka nemal.',
        "Kuue kuuga katsetas ta umbes 70 inimesega. 1989. aastal avaldas ta esimese kontrollitud uuringu. Täna on EMDR üks maailma enim uuritud traumateraapia meetodeid — ja see algas ühest jalutuskäikust.",
    ]),
    ("Miks see sind puudutab?", [
        "Sa ei pea olema katki, et seda proovida. Sa ei pea aru saama, kuidas see töötab. Piisab 15 minutist.",
        'Sinu aju töötleb igal ööl unes mõtteid — silmad liiguvad REM une ajal vasakule ja paremale. See on looduslik. Aga mõnikord jääb mõte kinni: homme esitlus, tüliline lause, pingutus kaelas. See tuleb tagasi ja tagasi, isegi kui sa "ei taha sellele mõelda."',
        "EMDR kodune versioon aitab ajul seda kinniolevat mõtet liikuma lükata — nagu Shapiro silmad pargis jalutades tegid. Mitte sellepärast, et sa oleksid haige. Sellepärast, et su keha juba teab, kuidas tulla rahule. Sa annad sellele 15 minutit ja natuke abi.",
    ]),
    ("Mis on EMDR?", [
        "EMDR (Eye Movement Desensitization and Reprocessing) kasutab bilateral stimulatsiooni — vahelduvat signaali mõlemale poole keha: silmaliigutused, koputused või heli vasakul ja paremal. See aitab ajul töödelda häirivaid mõtteid ja emotsioone loomulikul viisil — mitte surudes, mitte unustades, vaid läbi töötades.",
        "Kodus saad kasutada lihtsustatud versiooni. See ei asenda terapeuti raske trauma korral, aga võib aidata kerge pinge, ärevuse ja uneprobleemide juures — esimene samm tagasi rahusse.",
    ]),
]

STORY_QUOTE = "Sa ei pea uskuma mind. Proovi ühe korra. 15 minutit. Üks kerge mõte. Ja vaata, kas su keha vastab."


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    run = p.add_run(text)
    set_run_font(run, 18, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    run = p.add_run(text)
    set_run_font(run, 12, italic=True)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, {1: 14, 2: 12, 3: 11}.get(level, 11), bold=True)


def add_normal(doc, text, italic=False, center=False, bold=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 11, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, 10, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, 10)
    return table


def add_meta_table(doc):
    rows = [
        ("Kood", "EMDR1"),
        ("Teema", "Bilateral stimulatsioon kodus — iseendale ja paarilisele"),
        ("Versioon", "1.1 · 9. august 2026"),
        ("Kasutaja", "Ise / paar kodus"),
        ("Sihtgrupp", "Täiskasvanud — kerge stress, mure, igapäevane pinge"),
        ("Eeldused", "Mitte trauma, PTSD, raske kriis — siis EMDR-terapeut"),
        ("Seotud", "emdr-algaja-juhised.docx, emdr-home-guide.html"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, 10, bold=(cell == table.rows[i].cells[0]))


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title(doc, "EMDR1 — Algaja juhised")
    add_subtitle(doc, "Koolituskaart + protokoll (15–20 min)")
    add_heading(doc, "KOOLITUSKAART EMDR1 — Algaja juhised", 1)
    add_normal(doc, "Unpluged-Al | Kodune juhend | Kestus: 15–20 min")
    add_meta_table(doc)
    add_normal(doc, "See juhend on hariduslik. Ei asenda psühholoogi või EMDR-terapeudi abi.", italic=True)

    for title, paragraphs in STORY_SECTIONS:
        add_heading(doc, title, 2)
        for i, para in enumerate(paragraphs):
            add_normal(doc, para, bold=(i == 0 and title.startswith("Lugu")))
    add_normal(doc, STORY_QUOTE, italic=True)

    add_heading(doc, "KOOLITUSKAART (prindi see leht — taskus või külmkapis)", 2)
    add_code_block(doc, ASCII_CARD)

    add_heading(doc, "PROTSESS — täielik sisu (tagakülg või järgmine leht)", 2)
    add_table(
        doc,
        ["Samm", "Faas", "Mida teed", "Märkused"],
        [
            ("0", "Valmistumine", "Vaikne koht. Vesi lähedal. Vali meetod (A/B/C). Hinda distress 0–10, kirjuta üles. Vali turvakoht. Sea taimer 15–20 min.", "Mõlemale"),
            ("1", "Turvakoht (2–3 min)", "Sule silmad. Kujutle turvalist kohta. Mida näed, kuuled, tunned? Püsi seal, kuni kehas on väike rahutunne.", "Ise"),
            ("2", "Sihtmärk (1 min)", "Üks asi korraga: üks mõte, pilt või kehatunne. Ära vali kõige raskemat — kerge ärevus või tüli, mitte trauma.", "Ise"),
            ("3", "Bilateral (10–15 min)", "Too sihtmärk meelde → bilateral 24–30 → peatu → märka mis tuleb → korda 3–6 komplekti. Iga 2 komplekti järel: distress 0–10?", "Ise"),
            ("4", "Positiivne uskumus", 'Vali lause: "Ma saan hakkama", "Ma olen turvaliselt", "See möödub". Hoia meeles + 2 bilateral komplekti.', "Valikuline"),
            ("5", "Lõpetus", "3 aeglast hingetõmmet. Joo vett. Kirjuta üks lause: mis muutus?", "Ise"),
        ],
    )

    add_heading(doc, "MEETODID — vali üks", 2)
    add_table(
        doc,
        ["Meetod", "Kuidas", "Komplekt"],
        [
            ("A. Silmaliigutused", "Kaks sõrme silmade kõrgusele, õlalaiuselt. Silmad sujuvalt vasakult paremale ja tagasi. Pea paigal. Tempo ~1 sek/pool.", "24–30 liigutust (12–15× edasi-tagasi)"),
            ("B. Koputused", "Istudes: vaheldumisi vasak/parem põlv. Või risti käed, koputa õlgu. Aeglane rütm ~1 koputus/sek.", "24–30 koputust"),
            ("C. Audio", "Bilateral heli/rakendus, mis vaheldab heli vasakul ja paremal. Madal helitugevus.", "1–2 minutit"),
        ],
    )

    add_heading(doc, "PAARILINE PROTOKOLL", 2)
    add_table(
        doc,
        ["Samm", "Abiline (B) teeb", "Töötegija (A) teeb"],
        [
            ("1", "Lepi kokku: üks teema, 15–20 min, mõlemad rahulikud.", "Valib kerge sihtmärgi."),
            ("2", "Juhi bilateral stimulatsiooni (koputused õlul või sõrm silmade ees).", "Too sihtmärk meelde."),
            ("3", '"Too meelde oma sihtmärk. Alustame." → 24–30 koputust/liigutust.', "Bilateral ajal keskendub sihtmärgile."),
            ("4", '"Peatu. Mis tuli?" — lase vastata või öelda "ei tea".', "Vastab lühidalt või vaikib."),
            ("5", '"Distress 0–10?" — kirjuta üles.', "Annab numbri."),
            ("6", "Korda 3–6 korda. Lõpus: vesi, 3 hingetõmmet. Vaheta rollid järgmisel korral.", "Osaleb protokollis."),
        ],
    )

    add_heading(doc, "Mida abiline EI tee", 3)
    add_table(
        doc,
        ["❌", "✓ Selle asemel"],
        [
            ('Ei analüüsi ega anna nõu ("sa peaksid lihtsalt...")', 'Küsi: "Mis tuleb nüüd?"'),
            ("Ei suru rääkima", "Luba vaikust"),
            ("Ei jätka, kui partner ütleb stop", "Peata kohe"),
            ("Ei tee kerge teemast rasket", "Hoia teema kerge"),
        ],
    )

    add_heading(doc, "MILLAL STOP", 2)
    add_table(
        doc,
        ["Olukord", "Tegevus"],
        [
            ("Distress üle 8/10", "Lõpeta bilateral kohe"),
            ("Pearinglus, iiveldus, paanika", "5-4-3-2-1 maandamine"),
            ('Tunne "lahutatud" kehast (dissotsiatsioon)', "Külma vett, jalad põrandale"),
            ("Tugevad flashback'id", "Pöördu EMDR-terapeudi poole"),
        ],
    )
    add_normal(doc, "5-4-3-2-1 maandamine: 5 asja mida näed · 4 mida tunned · 3 mida kuuled · 2 mida lõhnad · 1 mida maitsed.")

    add_heading(doc, "OSALEJA TÖÖLEHT (üks seanss)", 2)
    add_code_block(doc, WORKSHEET)

    add_heading(doc, "NÄITED ALGAJALE", 2)
    add_table(
        doc,
        ["Olukord", "Sihtmärk", "Meetod", "Tulemus"],
        [
            ("Ise — uneärevus", '"homme esitlus"', "Koputused, 4 komplekti", "Distress 6 → 3"),
            ("Paar — kerge tüli", "Üks lause tülis", "Abiline koputab õlgu, 3 komplekti", "Distress langeb, vabandust ei pea"),
            ("Ise — kehatunne", "Pingutus kaelas", "Silmaliigutused, 5 komplekti", "Kael lõdveneb natuke"),
        ],
    )

    add_heading(doc, "ENESEREFLEKSIOON (ei ole eksam)", 2)
    add_table(
        doc,
        ["Küsimus", "☐ jah ☐ osaliselt ☐ ei"],
        [
            ("Valisin turvakoha enne alustamist?", ""),
            ("Hoidsan sihtmärgi kerge (mitte trauma)?", ""),
            ("Kontrollisin distressi iga 2 komplekti järel?", ""),
            ("Lõpetasin, kui distress tõusis üle 8?", ""),
            ("Tean, millal pöörduda terapeudi poole?", ""),
        ],
    )

    add_heading(doc, "MATERJALID", 2)
    add_table(
        doc,
        ["Fail", "Otstarve"],
        [
            ("opord/kursid/EMDR-algaja-koolituskaart.md", "See kaart"),
            ("opord/kursid/EMDR-algaja-koolituskaart.docx", "Google Docs / printimine"),
            ("emdr-algaja-juhised.docx", "Täielik juhend (pikem formaat)"),
            ("emdr-home-guide.html", "Veebiversioon"),
        ],
    )

    add_normal(doc, "EMDR1 — Algaja juhised. Unpluged-Al. Avalik — isiklikuks ja paariliseks kasutamiseks kodus.", italic=True, center=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
