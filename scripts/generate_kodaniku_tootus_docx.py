#!/usr/bin/env python3
"""Generate Kodaniku tõotus Word document in Estonian."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_centered(doc, text, bold=False, size=12, italic=False, space_after=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Georgia"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_centered(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else 14)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    return p


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    add_heading_centered(doc, "Kodaniku tõotus", level=1)
    add_centered(doc, "Identiteet, mida võid kohe omaks võtta — mõtlemist pole vaja", italic=True, size=11)
    add_centered(doc, "Kes ma olen?", bold=True, size=16, space_after=14)

    add_centered(
        doc,
        "Sa ei pea ennast nullist leiutama.\nVõta see tõotus. See on nüüd sinu.",
        italic=True,
        size=12,
        space_after=14,
    )

    add_centered(
        doc,
        "See tekst on kohandatud ametlikust Eesti kaitseväelase tõotusest "
        "(kaitseväelase tõotus) — ümber kujundatud tavaliseks eluks: "
        "naabritele, vanematele, töötajatele, õpilastele ja kõigile, "
        "kes küsivad: kes ma olen?",
        size=11,
        space_after=16,
    )

    # Creed box
    add_centered(doc, "KODANIKU TÕOTUS", bold=True, size=10, space_after=8)
    creed = (
        "Mina, inimene, tõotan jääda ustavaks tõele, väärikusele ja inimestele, "
        "kes minust sõltuvad; kaitsta head endas ja teistes kogu oma mõistuse ja hoole eest; "
        "olla valmis ohverdama mugavust ja ego selle eest, mis on õige; "
        "hoida kinni sisemisest distsipliinist; ning täita oma kohustusi ausalt ja järjepidevalt, "
        "teades, et vastasel korral maksan hinda ise; "
        "nõuan, et minuga käitutakse austusega ja õiglaselt - ja käitun teistega samamoodi; "
        "tean, et mida külvad, seda lõikad: iga tegu paneb seemne, mis tuleb ringiga tagasi; "
        "alustan heana, kuid ei lase kurja minna vastutuseta - see on minu vastastikmõju strateegia."
    )
    p = add_centered(doc, creed, size=13, space_after=16)
    for run in p.runs:
        run.bold = False

    add_centered(
        doc,
        "Allikas: Kohandatud Eesti kaitseväelase ametlikust tõotusest "
        "(Kaitseväeteenistuse seadus § 10): "
        "„Mina, (ees- ja perekonnanimi), tõotan jääda ustavaks demokraatlikule "
        "Eesti Vabariigile ja tema põhiseaduslikule korrale, kaitsta Eesti Vabariiki "
        "vaenlase vastu kogu oma mõistuse ja jõuga… "
        "Sõjaväeline keel on tõlgitud igapäevase kodaniku keelde.",
        italic=True,
        size=10,
        space_after=18,
    )

    add_heading_centered(doc, "Kes ma olen? — Ütle seda valjusti", level=2)
    add_centered(doc, "Sul pole vaja täiuslikku vastust. Sul on vaja kindlat vastust. Kasuta neid:", space_after=10)

    identity = [
        "Olen inimene, kes peab oma sõna.",
        "Olen inimene, kes kaitseb head — alustades iseendast.",
        "Olen inimene, kes räägib tõtt, isegi kui see maksab.",
        "Olen inimene, kes ilmub kohale nende jaoks, kes minust loodavad.",
        "Olen inimene, kes valib distsipliini impulsi asemel.",
        "Olen inimene, kes kannab oma valikute hinna.",
        "Olen kodanik - mitte piiri, vaid väärikuse mõttes.",
        "Nõuan, et minuga käitutakse austusega ja õiglaselt - ja annan seda tagasi.",
        "Külvin head, sest tean: mida külvad, seda lõikad.",
        "Alustan heana, kuid ei lase kurja minna vastutuseta.",
    ]
    for line in identity:
        add_centered(doc, f"- {line}", space_after=6)

    add_heading_centered(doc, "Loodusseadus ja vastastikmõju", level=2)

    add_heading_centered(doc, "Konversioonilugu - kuidas ma sinna jõudsin", level=2)
    add_centered(doc, "Ma ei teadnud pikka aega, mida ma tegelikult teen.", italic=True, size=12, space_after=10)
    add_centered(
        doc,
        "Elasin, reageerisin, põrkasin inimestega kokku - aga mul polnud strateegiat. "
        "Ma ei teadnud, miks mõned suhted lagunevad ja teised püsivad. "
        "Ma ei teadnud, miks mõnikord tundus, et annan liiga palju, ja mõnikord liiga vähe. "
        "Lihtsalt elasin edasi, ilma raamita.",
        size=11,
        space_after=10,
    )
    add_centered(
        doc,
        "Siis sattusin YouTube'i peale. Üks video rääkis strateegiatest - "
        "kuidas arvutiprogrammid võistlevad üksteise vastu, korduvalt, tuhandeid kordi. "
        "Kokku oli 42 erinevat strateegiat. Igaüks püüdis võita. "
        "Mõni oli alati kuri. Mõni alati hell. Mõni püüdis petta. Mõni andis alla.",
        size=11,
        space_after=10,
    )
    add_centered(
        doc,
        "Võitja oli tit-for-tat - vastastikmõju. "
        "Alusta heana. Vasta samaga, mida saad. Kui teine on hea - ole hea. "
        "Kui teine on kuri - sea piir. Ära löö kunagi esimesena.",
        size=11,
        space_after=10,
    )
    add_centered(
        doc,
        "See ei olnud lihtsalt arvutimäng. See oli päriselu. "
        "Sellest sai mulle selgeks, miks mida külvad, seda lõikad on loodusseadus - "
        "mitte ainult vana õpetus, vaid reegel, mis töötab inimeste, sõprade, "
        "laste ja sinu enda vahel.",
        size=11,
        space_after=12,
    )
    add_centered(
        doc,
        "Tit-for-tat on ainus strateegia, kus keegi ei kaota.\n"
        "Inimesed ei kaota. Sõbrad ei kaota.\n"
        "Suhted lastega ei kaota. Ja sina ise ei kaota.",
        bold=True,
        size=12,
        space_after=12,
    )
    add_centered(
        doc,
        "Sa ei pea olema nõrk ega kange. Sa ei pea andma ennast ära ega sõdima iga päev. "
        "Sa alustad heana - sest see on see, mida sa tahad lõigata. "
        "Ja kui keegi ei vasta heaga, sa ei jää vaikides kannatama. "
        "Sa sead piiri. See on ausus. See on väärikus. See on loodusseadus.",
        size=11,
        space_after=10,
    )
    add_centered(
        doc,
        "Nii jõudsin ma siia. Mitte teooria kaudu. Elu kaudu. Ja ühe video kaudu, mis muutis kõik.",
        italic=True,
        size=11,
        space_after=14,
    )

    add_centered(
        doc,
        "Mida külvad, seda lõikad - see pole ainult vanasõna. See on loodusseadus.",
        italic=True,
        size=12,
        space_after=10,
    )
    add_centered(
        doc,
        "Iga sõna, iga tegu, iga valik paneb seemne. Head seemned kasvavad head saagi. "
        "Halvad seemned tulevad ringiga tagasi - mitte alati kohe, aga alati. "
        "Sa ei kontrolli saaki, aga sa kontrollid külvikut.",
        size=11,
        space_after=12,
    )

    natural_law = [
        ("Külvamine", "Enne kui teed, küsi: mida ma praegu külvin? Kas see on see, mida tahan lõigata?"),
        ("Lõikamine", "Kui elu toob sulle rasket, vaata ausalt: kas ma ise külvatasin selle seemne?"),
        (
            "Vastastikmõju (tit-for-tat)",
            "Alusta alati heana. Kui teine on hea - ole hea. Kui teine on kuri - sea piir. "
            "Anna teine võimalus, kui keegi eksib. Ära ole lollukindel. Ära ole ka nõrk.",
        ),
    ]
    for title, desc in natural_law:
        add_centered(doc, title, bold=True, size=12, space_after=4)
        add_centered(doc, desc, size=11, space_after=10)

    add_centered(
        doc,
        "Vastastikmõju strateegia lihtsalt: "
        "1. Alusta koostööga · 2. Vasta samaga, mida saad · 3. Anna andeks üks kord · "
        "4. Kui kurjus kordub - kaitse end · 5. Ära kunagi löö esimesena",
        italic=True,
        size=10,
        space_after=16,
    )

    add_heading_centered(doc, "Kuus väärtust", level=2)
    add_centered(
        doc,
        "Võetud Kaitseväe eetikakoodeksist — tõlgitud igapäevaseks eluks:",
        space_after=12,
    )

    values = [
        ("Ausus", "Räägin endale ja teistele tõtt. Ei peida end vabanduste taha."),
        ("Vaprus", "Seisan silmitsi raskega. Ei põgu vajalikest vestlustest ega tegudest."),
        ("Asjatundlikkus", "Teen oma tööd hästi. Õpin pidevalt. Võtan kohustusi tõsiselt."),
        ("Ustavus", "Jään truuks sellele, mida usun, ja inimestele, kes minusse usaldavad."),
        ("Koostöövalmidus", "Aitan, kui saan. Ei jäta teisi üksi kandma seda, mis on meie ühine."),
        ("Avatus", "Kuulan. Tunnistan, kui eksin. Jään valmis kasvama."),
    ]
    for title, desc in values:
        add_centered(doc, title, bold=True, size=12, space_after=4)
        add_centered(doc, desc, size=11, space_after=10)

    add_heading_centered(doc, "Igapäevane praktika — kolm minutit", level=2)
    add_centered(doc, "Iga hommik loe tõotus üks kord läbi. Siis küsi endalt üks küsimus:", space_after=8)
    add_centered(doc, "Mis on üks asi täna, mida minu-sugune inimene teeks?", bold=True, space_after=10)
    add_centered(doc, "See on piisav. Identiteet pole meeleolu. See on valikute muster.", space_after=16)

    add_heading_centered(doc, "Kuidas see päriselus välja näeb", level=2)

    examples = [
        ("Kodus", "Oled väsinud, aga su laps vajab tähelepanu. Sa ilmud ikkagi kohale. See on tõotus."),
        ("Tööl", "Võiksid nurgast lõigata. Sa teed ausa asja. See on sisemine distsipliin."),
        ("Sõbraga", "Talle on vaja karmi tõde, mitte mugavust. Sa räägid hoolikalt. See on vaprus."),
        ("Üksi", "Keegi ei vaata. Sa pead ikkagi oma sõna. See oled sina."),
        ("Kui sind ei austa", "Sa ei vihastu ega põgene. Sa ütled selgelt: minuga tuleb käituda austusega. See on sinu õigus."),
        ("Kui keegi on hea", "Sa vastad heaga. Külvid head. See on loodusseadus töös."),
        ("Kui keegi kurjust kordab", "Sa ei tasu kurja kurjaga - aga sa sead piiri. Vastastikmõju: üks hoiatus, siis kaitse."),
        ("Kui eksid", "Sa ei teeskle. Võtad vastutuse, parandad, mis saad, ja alustad uuesti. See on hinna maksmine."),
    ]
    for title, desc in examples:
        add_centered(doc, title, bold=True, size=12, space_after=4)
        add_centered(doc, desc, size=11, space_after=10)

    add_heading_centered(doc, "Võta see omaks — kohe praegu", level=2)
    add_centered(doc, "Sa ei pea seda täna täiuslikult uskuma.", space_after=6)
    add_centered(doc, "Sa pead seda ainult valima.", bold=True, space_after=8)
    add_centered(doc, "Loe tõotus valjusti. Allkirjasta all — paberil või mõttes.", space_after=20)
    add_centered(doc, "_________________________________", space_after=6)
    add_centered(doc, "Sinu nimi · Kuupäev", size=10, space_after=14)
    add_centered(
        doc,
        "Sellest hetkest, kui keegi küsib 'kes ma olen?' - on sul vastus.",
        italic=True,
        space_after=20,
    )

    # Comparison table
    add_heading_centered(doc, "Sõduri tõotus ja kodaniku tõotus", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Sõduri tõotus"
    hdr[1].text = "Kodaniku tõotus"
    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    rows = [
        ("Ustavus vabariigile", "Ustavus tõele ja väärikusele"),
        ("Kaitseb vaenlase eest", "Kaitseb head endas ja teistes"),
        ("Ohverdab elu isamaa eest", "Ohverdab mugavust ja ego õige eest"),
        ("Kaitseväe distsipliin", "Sisemine distsipliin"),
        ("Kohustused täpselt", "Ausad ja järjepidevad kohustused"),
        ("Seadus karistab", "Maksan hinda ise"),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        for cell in row:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_centered(
        doc,
        "Kohandatud Eesti kaitseväelase tõotusest ja eetikakoodeksist.\n"
        "Vabalt kõigile, kes küsivad: Kes ma olen?",
        size=10,
        italic=True,
    )
    add_centered(doc, "Unpluged-Al · Kodaniku tõotus", size=9, space_after=0)

    output = "/workspace/kodaniku-tootus.docx"
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    main()
