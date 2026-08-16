#!/usr/bin/env python3
"""Generate Renee Aluste human/citizen portrait with SWOT analysis as DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "docx" / "Renee-Aluste-inimlik-portree-SWOT.docx"


def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3D, 0x6B, 0x4F)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_color: str = "E8F2EC") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_color)
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, value in enumerate(row):
            row_cells[col_idx].text = value
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title = doc.add_heading("Renee Aluste — inimlik ja kodaniklik portree", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Täielik kirjeldus inimesena ja kodanikuna · SWOT-analüüs")
    run.italic = True
    run.font.size = Pt(12)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Põhineb jagatud materjalidel ja vestlustel · 8. august 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    add_heading(doc, "Oluline märkus", 2)
    add_para(
        doc,
        "See portree põhineb Renee Aluste enda jagatud materjalidel — Operatsioon „Peegel“ OPORD, "
        "Lisa R, renee-aluste-profiil, Kodaniku tõotus ja varasemad vestlused. See ei ole kliiniline "
        "diagnoos ega väljastpoolt tulev iseloomuanalüüs.",
    )

    add_heading(doc, "I. Kes sa inimesena oled", 1)

    add_heading(doc, "1.1 Identiteedi tuum", 2)
    add_para(
        doc,
        "Sa ei ole ühe rolli inimene. Sa oled mitmekihiline identiteet, mis on kujunenud aastakümnete "
        "kaupa ja mis praegu ühineb üheks selgeks teljeks: taastumine → distsipliin → teenimine → "
        "kodaniklik mõju.",
    )
    add_quote(doc, "„Paranemine algab sellel hetkel, kui sa mõistad, et sa pole katki olnudki.“")
    add_para(
        doc,
        "Sa oled inimene, kes on nii kaevikus kui ka üleval. Oled olnud süsteemi sees — Kaitsevägi, "
        "eriväed, Kaitseliit, — ja näed, kuidas paljud inimesed tegelevad sümptomite kustutamisega, "
        "mitte juure parandamisega.",
    )

    add_heading(doc, "1.2 Elukaared", 2)
    add_table(
        doc,
        ["Eluperiood", "Mida see sinasse jättis"],
        [
            ["Lapsepõlv ja noorus", "37 aastat domineeriva juhtimise ja vägivalla all. Oled ise olnud kiusaja — ausalt tunnistatud."],
            ["2008–2018: Kaitsevägi", "EOG operaator, sniprigrupi ülem. Vahipataljon. Estpla-21 UNIFIL Liibanonis (2015)."],
            ["2022: Ranger School", "USA Army Ranger School Class 07-22. Pärast ajukasvajaga võitlemist."],
            ["2018–…: Kaitseliit", "Kompaniiülem, instruktor. ~6 SOK kursust, 240+ tundi igaüks."],
            ["2023–2026:", "Juhtimisinstruktor ettevõtetele. Lõpetatud 2026. Siin põlesid läbi."],
            ["Tsiviil", "Arborist. Füüsiline töö, distsipliin käte kaudu."],
            ["Isiklik kriis", "Depressioon, infosõja mõju. Abi: narko.ee, ravimid, teraapia, EMDR."],
            ["Praegu", "Operatsioon „Peegel“ eestvedaja. Kodanik, mitte ametnik."],
        ],
    )

    add_heading(doc, "1.3 Psühholoogiline portree", 2)
    add_para(doc, "Tugevused isiksuses:", bold=True)
    for item in [
        "Ausus kui strateegia — „Ma ei manipuleeri sind nõrkuseks.“",
        "Struktuur kui turvalisus — OPORD-id, tabelid, checklistid.",
        "Keha-teadlikkus — mõistus võib valetada, keha mitte.",
        "Refleksiivsus — kas see ehitab või lõhub?",
        "Vastupidavus — Ranger, ajukasvaja, depressioon, läbipõlemine.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, "Haavatavused isiksuses:", bold=True)
    for item in [
        "Läbipõlemise kalduvus — kalduvus kanda liiga palju üksi.",
        "Dominantsi varju — oled tunnistanud julmust ja kiusamist.",
        "Isa-valu — lapsed võõrandatud (Lily 14, Rene 6, Henry 11).",
        "Usaldusväärse usaldamatuse paradoks — „ära usu mind“ vs „usaldage minu lugu“.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "1.4 Suhtlemisstiil", 2)
    add_table(
        doc,
        ["Dimensioon", "Kuidas sa suhtled"],
        [
            ["Register", "Juhtimisinstruktori — selge, motiveeriv, trauma-teadliku pehme servaga"],
            ["Tempo", "Maraton, mitte sprint. Üks KOV nädalas. Üks vestlus. Üks samm."],
            ["Vorm", "Lood, mitte loengud. Üks teema = üks lugu."],
            ["Konflikt", "Tit-for-tat: alusta heana, vasta samaga, sea piir"],
            ["Kriis", "„Kata ja liigu“ — konkreetne samm"],
            ["Skeptik", "Valideeri kahtlust, anna üks link, ära müü"],
        ],
    )

    add_heading(doc, "1.5 Väärtussüsteem", 2)
    add_quote(
        doc,
        "Mina, inimene, tõotan jääda ustavaks tõele, väärikusele ja inimestele, kes minust sõltuvad; "
        "kaitsta head endas ja teistes; hoida kinni sisemisest distsipliinist; täita kohustusi ausalt ja järjepidevalt.",
    )
    add_para(doc, "Kuus väärtust: Ausus · Vaprus · Asjatundlikkus · Ustavus · Koostöövalmidus · Avatus")
    add_para(doc, "Operatiivsed põhimõtted: Kuula enne · Peegelda enne otsust · Inimene enne tööriista · Usu keha")

    add_heading(doc, "1.6 Rollid", 2)
    add_table(
        doc,
        ["Roll", "Tähendus"],
        [
            ["Isa", "Keskne. Tugeva Isa seeria. Pere kui rindejoon. Aga ka haav — võõrandumine."],
            ["Eriväelane / Ranger", "Moraalne ja metoodiline alus, mitte poliitiline relv"],
            ["Juhtimisinstruktor", "Endine. Oskus juhtida inimesi, mitte protsesse"],
            ["Arborist", "Keha ja maa. Füüsiline töö."],
            ["Autor", "7 köidet. Sõna kui tööriist."],
            ["Kodanikualgatuse eestvedaja", "Operatsioon „Peegel“. Mitte riiklik käsk."],
            ["Mentor / eeskuju", "„Ma ei palu uskuda. Ma palun kontrollida.“"],
        ],
    )

    add_heading(doc, "II. Kes sa kodanikuna oled", 1)

    add_heading(doc, "2.1 Kodanikuidentiteet", 2)
    add_para(
        doc,
        "Sa ei määratle end riigi teenijana ametlikus mõttes. Sa oled kodanik, kes võttis kaitseväelase "
        "tõotuse ja tõlkis selle igapäevaeluks. Sa ei oota, et riik lahendaks sinu probleeme. Sa ehitad "
        "ise tööriistad ja pakud neid teistele vabatahtlikult.",
    )

    add_table(
        doc,
        ["Sõduri tõotus", "Sinu kodaniku tõotus"],
        [
            ["Ustavus vabariigile", "Ustavus tõele ja väärikusele"],
            ["Kaitseb vaenlase eest", "Kaitseb head endas ja teistes"],
            ["Ohverdab elu isamaa eest", "Ohverdab mugavust ja ego õige eest"],
            ["Kaitseväe distsipliin", "Sisemine distsipliin"],
            ["Seadus karistab", "Maksan hinda ise"],
        ],
    )

    add_heading(doc, "2.2 Kodaniklik positsioon Eestis", 2)
    add_para(doc, "Mida sa teed:", bold=True)
    for item in [
        "Aitad õpetajaid ja KOV-e TI-ülemineku ajal",
        "Ehidad kodanikuvõrgustikku",
        "Analüüsid valitsuse mustreid kodanikualgatuslikult",
        "Jagad isiklikku lugu depressiooni ja taastumise kohta",
        "Toetad kriitilist mõtlemist ilma vandenõuteooria müümiseta",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, "Mida sa EI tee:", bold=True)
    for item in [
        "Ei esinda riiki ega Kaitseväge ametlikult",
        "Ei masskirjata ega manipuleeri",
        "Ei palu uskuda ilma kontrollimata",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_quote(doc, "„Järgmine sõda on su nutitelefonis. Peresid lõhutakse seestpoolt.“")

    add_heading(doc, "2.3 Kodaniku tegevusvaldkonnad", 2)
    add_table(
        doc,
        ["Valdkond", "Sinu panus"],
        [
            ["Pere ja isa-roll", "Tugeva Isa seeria, Papsid.ee, laste kaitse"],
            ["Haridus", "KOV kampaania, õpetajate TI-tugi, 1-1 videokõned"],
            ["Vaimne tervis", "Lisa H, EMDR guide, depressiooni teekond"],
            ["Infosõda", "Kriitiline mõtlemine, algoritmide teadlikkus"],
            ["Juhtimine", "Inimesekeskne juhtimine, NVC"],
            ["Kogukond", "Võrgustik: veteranid, breathwork, Harku piirkond"],
            ["Avalik arutelu", "Postimees, Delfi, podcastid"],
        ],
    )

    add_heading(doc, "III. SWOT-analüüs", 1)

    add_heading(doc, "STRENGTHS (Tugevused)", 2)
    add_table(
        doc,
        ["#", "Tugevus", "Miks see loeb"],
        [
            ["S1", "Kontrollitav pädevus", "KV 10a, Ranger, Kaitseliit, 50+ koolitust"],
            ["S2", "Isiklik lugu = usaldus", "Depressioon, ajukasvaja, Ranger — dokumenteeritud meedias"],
            ["S3", "Struktuuri oskus", "OPORD, tabelid, checklistid — keerulise selgeks tegemine"],
            ["S4", "Ausus ja läbipaistvus", "„Ära usu mind. Kontrolli ise.“"],
            ["S5", "Trauma-teadlikkus", "van der Kolk, EMDR, turvalisus enne vastutust"],
            ["S6", "Mitmekülgne kogemus", "Sõjavägi + arborist + koolitaja + autor + isa"],
            ["S7", "Kirjalik pärand", "7 köidet, OPORD + 40+ lisa, taskukaardid"],
            ["S8", "Võrgustik", "Veteranid, haridus, breathwork, meedia"],
            ["S9", "Keeleoskus", "Eestikeelne primaarselt, inglise C1"],
            ["S10", "Füüsiline vastupidavus", "Arborist, Ranger, sõjavägi"],
        ],
    )

    add_heading(doc, "WEAKNESSES (Nõrkused)", 2)
    add_table(
        doc,
        ["#", "Nõrkus", "Risk"],
        [
            ["W1", "Läbipõlemise kalduvus", "3a SOK → läbipõlemine. Operatsioon on suur — risk kordub"],
            ["W2", "Dominantsi varju", "Julmus ja kiusamine — võib tulla välja stressis"],
            ["W3", "Isa-valu", "Laste võõrandumine — emotsionaalne koormus"],
            ["W4", "Scope creep", "40+ lisa, 7 raamatut — üks inimene ei jõua"],
            ["W5", "Sõjalise raamistiku segamine", "OPORD võib tekitada militaarse operatsiooni mulje"],
            ["W6", "Usaldusväärse usaldamatuse paradoks", "Skeptikud näevad vastuolu"],
            ["W7", "Üksilduse kalduvus", "Liiga kaua üksi kandnud"],
            ["W8", "Poliitilised teemad", "Vaktsineerimine, valitsuse kriitika — polaarne reaktsioon"],
            ["W9", "Perfektsionism struktuuris", "Võib takistada „üks samm ja edasi“ liikumist"],
            ["W10", "Endine kiusaja identiteet", "Võib kasutada diskrediteerimiseks"],
        ],
    )

    add_heading(doc, "OPPORTUNITIES (Võimalused)", 2)
    add_table(
        doc,
        ["#", "Võimalus", "Kuidas kasutada"],
        [
            ["O1", "TI-hüpe hariduses (2025)", "Õpetajad hädas — KOV kampaania ajastatud õigesti"],
            ["O2", "Vaimse tervise tabustuse murdmine", "Sinu lugu on eeskuju"],
            ["O3", "Kodanikualgatuste kasv", "Kodaniku tõotus, 1-1 võrgustik"],
            ["O4", "Veteranide kogukond", "Tähenduse otsimine pärast teenistust"],
            ["O5", "Meedia usaldus", "Delfi, Postimees, podcastid"],
            ["O6", "Raamatud ja digitaalne sisu", "Skaleeritav ilma kohalolekuta"],
            ["O7", "Inimesekeskne juhtimine trend", "Steiger, NVC, trauma-teadlikkus"],
            ["O8", "Kaitseliidu võrgustik", "Kaasa teisi, kes jagavad visiooni"],
            ["O9", "1-1 mudel", "Personaliseeritud, vähem läbipõlemist"],
            ["O10", "„Aju vabadus“ rakendus", "Digitaalne kanal nooremale põlvkonnale"],
        ],
    )

    add_heading(doc, "THREATS (Ohud)", 2)
    add_table(
        doc,
        ["#", "Oht", "Mõju"],
        [
            ["T1", "Läbipõlemine (korduv)", "Operatsioon suur, üks koordinaator"],
            ["T2", "Diskrediteerimine", "„Vandenõuteooria müüja“, „endine kiusaja“"],
            ["T3", "Sõjalise raamistiku valesti mõistmine", "Militaarse liikumise mulje"],
            ["T4", "Poliitiline polariseerumine", "Valitsuse kriitika — ühele poole paigutamine"],
            ["T5", "Algoritmide vastu võitmine", "Sinu sõnum ei pruugi jõuda"],
            ["T6", "Üksildus", "Võrgustik ei kasva — jääd üksi"],
            ["T7", "Isa-valu trigger", "Kriitika isa-rolli teemal"],
            ["T8", "Riigi aeglus", "Frustratsioon → ülekoormus"],
            ["T9", "Tervise riskid", "Ajukasvaja ajalugu, stress"],
            ["T10", "Scope ei lõpe", "Igavesti ehitamine, mitte tegutsemine"],
        ],
    )

    add_heading(doc, "IV. SWOT-strateegia", 1)

    add_heading(doc, "SO — kasuta tugevusi võimaluste jaoks", 2)
    add_table(
        doc,
        ["Kombinatsioon", "Tegevus"],
        [
            ["S2 + O2", "Isiklik lugu vaimse tervise tabustuse murdmiseks"],
            ["S3 + O1", "OPORD-struktuur KOV kampaaniaks"],
            ["S7 + O6", "Skaleeri raamatud digitaalselt"],
            ["S8 + O4", "Kaasa veteranid võrgustikku"],
            ["S4 + O3", "„Ära usu mind“ + kodanikualgatus = usaldus"],
        ],
    )

    add_heading(doc, "WO — paranda nõrkusi võimaluste kaudu", 2)
    add_table(
        doc,
        ["Kombinatsioon", "Tegevus"],
        [
            ["W1 + O9", "1-1 mudel + V-formatsioon rotatsioon"],
            ["W4 + O6", "Scope creep → skaleeri digitaalselt"],
            ["W5 + O3", "Sõjaline raamistik → Kodaniku tõotus"],
            ["W7 + O8", "Üksildus → Kaitseliidu võrgustik"],
        ],
    )

    add_heading(doc, "ST — kasuta tugevusi ohtude vastu", 2)
    add_table(
        doc,
        ["Kombinatsioon", "Tegevus"],
        [
            ["S1 + T2", "CV ja meedia tõendid diskrediteerimise vastu"],
            ["S4 + T4", "Ausus poliitilise polariseerumise vastu"],
            ["S5 + T7", "Trauma-teadlikkus isa-valu triggerite vastu"],
        ],
    )

    add_heading(doc, "WT — minimeeri nõrkusi ja ohte", 2)
    add_table(
        doc,
        ["Kombinatsioon", "Tegevus"],
        [
            ["W1 + T1", "V-formatsioon, rotatsioon, 1-1 mitte mass"],
            ["W4 + T10", "„Üks tee“ — mitte kõik korraga"],
            ["W7 + T6", "Jätka 1-1 võrgustiku ehitamist"],
            ["W9 + T5", "Lihtne sõnum: Peegel tuum, 3 küsimust"],
        ],
    )

    add_heading(doc, "V. Inimlik kokkuvõte", 1)
    add_para(
        doc,
        "Renee Aluste on Tallinnas elav eestlane, kelle identiteet on kujunenud sõjalise teenistuse, "
        "isikliku kriisi, taastumise ja kodanikualgatuse kaudu.",
    )
    add_para(doc, "Sa oled inimene, kes:", bold=True)
    for item in [
        "On läbinud tule — depressioon, trauma, ajukasvaja, läbipõlemine, laste võõrandumine",
        "Ei varja haavu — tunnistad ausalt nii ohvri- kui ka kiusaja rolli",
        "Usub paranemisse — „sa pole katki olnudki“",
        "Ehitab struktuuri — OPORD-id, raamatud, checklistid",
        "Teenib kodanikuna — mitte riigi käsul, vaid oma tõotuse järgi",
        "Kaitseb peret — isa-identiteet on tuum",
        "Võitleb infosõjaga — kriitilise mõtlemise ja ausa vestlusega",
        "Ei palu uskuda — palub kontrollida ja valida",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, "Sa ei ole:", bold=True)
    for item in [
        "Riigi esindaja ega ametnik",
        "Erakonna liige ega poliitik",
        "Perfektne ega eksimatu",
        "Ohver ega kangelane — inimene, kes valis kasvada",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Sinu küsimus maailmale", 2)
    add_quote(doc, "Kas see, mida sa iga päev teed, ehitab sind üles või lõhub sind?")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Unpluged-Al · Operatsioon „Peegel“ · reneealuste.com")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
