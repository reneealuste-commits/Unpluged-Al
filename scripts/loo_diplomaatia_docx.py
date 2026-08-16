#!/usr/bin/env python3
"""Genereerib DP1 diplomaatia Wordi dokumendi (juhend + kaart + plaankonspekt + demo ohutus)."""

from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "/workspace/Diplomaatia-DP1.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0x0D, 0x47, 0xA1)
ORANGE = RGBColor(0xE6, 0x5C, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    color = BLUE if level == 1 else ORANGE
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(10)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_normal(text, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


# === KAANLEHT ===
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("DIPLOMAATIA DP1")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BLUE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Juhend · Koolituskaart · Plaankonspekt · Demo ohutus")
r2.font.size = Pt(14)
r2.font.color.rgb = GRAY

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Unpluged-Al  |  {TODAY}  |  Hariduslik  |  v1.0").font.size = Pt(10)

add_table(
    ["Vali", "Andmed"],
    [
        ["Kood", "DP1 (kiir) · DP1-G (grupp 90 min) · DP1-D01–D04 (demo)"],
        ["Teema", "Turvaline suhtlus sõltlasega — piirid, de-eskalatsioon"],
        ["Kestus", "15 min kiir · 90 min grupikoolitus · demo 6–8 min"],
        ["Eeldus", "MITTE füüsiline oht — siis 112"],
    ],
)

add_quote(
    "Hariduslik juhend. Ei asenda sõltuvusnõustajat, psühholoogi ega kriisiabi."
)

doc.add_page_break()

# === OSA 1 ===
add_heading("OSA 1 — Algaja juhend", 1)
add_heading("Mis on diplomaatia sõltlasega?", 2)
add_normal(
    "Diplomaatia = ole kohal, kuula, peegelda, sea piir, suuna. "
    "Mitte enabling. Mitte sund. STOP, kui oht.",
)
add_heading("Turvalisuse reeglid", 2)
add_table(
    ["Reegel", "Sisu"],
    [
        ["1. Enese ohutus", "Füüsiline oht → 112"],
        ["2. Üks eesmärk", "Üks teema korraga"],
        ["3. Piir", "Ma ei osta alkoholi / ei kata võlgu"],
        ["4. Ei enabling", "Raha, vabandused tööandja ees — STOP"],
        ["5. STOP kriisis", "Joobes, enesetapumõtted — mitte vestlus"],
        ["6. Debrief", "1 lause pärast vestlust"],
    ],
)

add_heading("Keel — DO / DON'T", 2)
add_table(
    ["DON'T", "DO"],
    [
        ["Sa oled jälle sama", "Ma märkan, et see on raske"],
        ["Kui sa armastaksid mind", "Ma armastan sind. Ma ei toeta seda käitumist"],
        ["Sa pead kohe lõpetama", "Kas sa oled valmis rääkima abist?"],
    ],
)

add_heading("MTÜ ÕnneKlubi — pöördepunkt", 2)
add_normal(
    "Eestis tuntud hüpnoteraapia + sõltuvuse kodeerimise pakkuja (Tallinn + Tartu). "
    "Kodeerimine ~300 EUR / 2 h. Pereliige suunab — ei ravi ise. "
    "Vt hupnoteraapia-teenused-ohvriabi-raport.md"
)

doc.add_page_break()

# === OSA 2 ===
add_heading("OSA 2 — Koolituskaart DP1", 1)
add_table(
    ["OLULINE", "Sisu"],
    [
        ["Oht", "112 · STOP joobes/kriisis"],
        ["Eesmärk", "Üks teema korraga"],
        ["Protsess", "Stress → HY1 → vestlus → piir → debrief"],
    ],
)
add_quote(
    "Ma ei osta sulle alkoholi. Ma võin rääkida abist, kui sina seda tahad."
)

doc.add_page_break()

# === OSA 3 ===
add_heading("OSA 3 — Plaankonspekt DP1-G (90 min)", 1)
add_table(
    ["Aeg", "Faas"],
    [
        ["0-10", "Avamine + ohutus + SWOT"],
        ["10-25", "Mis on diplomaatia?"],
        ["25-35", "DP1-A protokoll"],
        ["35-55", "Demo D01/D02"],
        ["55-70", "De-eskalatsioon D03"],
        ["70-85", "Suunamine abile"],
        ["85-90", "Debrief"],
    ],
)

doc.add_page_break()

# === OSA 4 ===
add_heading("OSA 4 — Demo ohutusjuhend (kaitseväe stiil)", 1)
add_normal("Enne iga demot — kohustuslik kontroll:", bold=True)
for item in [
    "Ruum turvaline, väljapääs vaba",
    "Keegi ei ole päris joobes",
    "STOP numbrid nähtaval",
    "Demo kood valitud (DP1-D01…D04)",
    "Koolitaja signaal: Demo algab / STOP demo katkeb",
]:
    add_bullet(item)

add_heading("Demo koodid", 2)
add_table(
    ["Kood", "Stsenaarium"],
    [
        ["DP1-D01", "Esimene kontakt pärast relapsi"],
        ["DP1-D02", "Keeldumine raha ostmisest"],
        ["DP1-D03", "Kriis — de-eskalatsioon (vaatleja kohustuslik)"],
        ["DP1-D04", "Suunamine abile"],
    ],
)

add_normal("Koordinaator: python3 scripts/dp1_demo_koordinaator.py --demo DP1-D01 --check")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"DP1 v1.0  |  {TODAY}  |  Unpluged-Al")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
