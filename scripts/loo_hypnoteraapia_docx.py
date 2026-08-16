#!/usr/bin/env python3
"""Genereerib HY1 hüpnoteraapia Wordi dokumendi (juhend + koolituskaart + plaankonspekt)."""

from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "/workspace/Hypnoteraapia-HY1.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

GREEN = RGBColor(0x1B, 0x5E, 0x20)
ORANGE = RGBColor(0xE6, 0x5C, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    color = GREEN if level == 1 else ORANGE
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GREEN
    run.font.size = Pt(10)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_normal(text, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


# === KAANLEHT ===
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HÜNOTERAAPIA HY1")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = GREEN

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Juhend · Koolituskaart · Plaankonspekt")
r2.font.size = Pt(14)
r2.font.color.rgb = GRAY

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Unpluged-Al  |  {TODAY}  |  Hariduslik  |  v1.2").font.size = Pt(10)

add_table(
    ["Vali", "Andmed"],
    [
        ["Kood", "HY1 (ise/paar) · HY1-G (grupp 90 min)"],
        ["Teema", "Turvaline hüpnoos — fookus, lõdvestus, eneseregulatsioon"],
        ["Kestus", "15-25 min kodus · 90 min grupikoolitus"],
        ["Eeldus", "MITTE trauma/PTSD/kriis — siis spetsialist"],
    ],
)

add_quote(
    "See dokument on hariduslik. Ei asenda psühholoogi, psühhiaatri ega "
    "sertifitseeritud hünoterapeudi abi."
)

doc.add_page_break()

# === OSA 1: JUHEND ===
add_heading("OSA 1 — Algaja juhend", 1)

add_heading("1. Mis on hünoteraapia?", 2)
add_normal(
    "Hünoteraapia kasutab hüpnoosi — fookustatud tähelepanu ja suuremat "
    "vastuvõtlikkust soovitud sisule — et aidata muuta mõtteid, tundeid, "
    "keha reaktsioone või käitumist.",
    size=11,
)
add_normal("See EI OLE:", bold=True)
for item in [
    "une peale sundimine",
    "tahte kaotamine",
    "kontrolli kaotamine",
    "kiire imeravi kõigele",
]:
    add_bullet(item)

add_normal("See ON:", bold=True)
for item in [
    "suunatud protsess ühe eesmärgiga",
    "turvalisuse ja nõusoleku peal",
    "sageli koos KVT, lõdvestus- või käitumisteraapiaga",
]:
    add_bullet(item)

add_heading("SWOT — lühikokkuvõte", 2)
add_table(
    ["", "Kokkuvõte"],
    [
        ["Tugevused", "Kiire abi stressile, uneale; keha kaudu ligipääs"],
        ["Nõrkused", "Kvaliteet kõikub; vale arusaam avalikkuses"],
        ["Võimalused", "Juhtimine, debrief, tervishoid"],
        ["Ohud", "Ülemüük; trauma vale käsitlus; asendada ravi"],
    ],
)

add_heading("2. Turvalisuse reeglid (KOHUSTUSLIK)", 2)
add_table(
    ["Reegel", "Sisu"],
    [
        ["1. Nõusolek", "Sa võid igal hetkel lõpetada"],
        ["2. Üks eesmärk", "Uni / rahunemine / fookus — mitte kõik korraga"],
        ["3. Teadlikkus", "Kuuled, tunned keha, saad silmad avada"],
        ["4. STOP trauma/kriis", "Raske minevik, dissotsiatsioon → mitte ise"],
        ["5. Ei asenda arsti", "Tõsine haigus vajab spetsialisti"],
        ["6. Protokoll", "1 lause: mis muutus? (debrief stiil)"],
    ],
)

add_normal("STOP — lõpeta kohe, kui:", bold=True)
for item in [
    "distress > 8/10 ja ei lange",
    "tunned end õhust väljas",
    "ülekoormav mälestus",
    "paanika tugevneb",
]:
    add_bullet(item)
add_normal(
    "STOP tegevus: Ava silmad. Jalad põrandale. 5-4-3-2-1. Vesi. Kõnni. Spetsialist vajadusel."
)

add_heading("3. Kodune protokoll (15-25 min)", 2)
add_table(
    ["Samm", "Faas", "Aeg"],
    [
        ["0", "Valmistumine — vaikne koht, vesi, üks eesmärk, stress 0-10", "2 min"],
        ["1", "Ankur — hingamine sisse 4, välja 6 × 3", "2 min"],
        ["2", "Keha skaneering — peast jalgadeni", "3 min"],
        ["3", "Turvakoht — kujutle turvalist kohta", "2 min"],
        ["4", "Suggestioon — HY1 lause 8-12 min", "8-12 min"],
        ["5", "Tagasitulek — loenda 5-1, ava silmad", "2 min"],
        ["6", "Integratsioon — mis muutus?", "1 min"],
    ],
)

add_quote(
    "Ma rahunen loomulikult. Mu keha teab, kuidas tulla rahule. "
    "Ma olen turvaliselt. Iga hingetõmme toob rohkem kergust."
)

add_heading("4. Millal spetsialist", 2)
for item in [
    "PTSD, raske trauma, dissotsiatsioon",
    "Depressioon, paanikahäired, krooniline unetus",
    "Tugev sõltuvus (suitsetamine, alkohol)",
    "Laste hüpnoos ilma lapsepsühholoogi taustata",
]:
    add_bullet(item)

doc.add_page_break()

# === OSA 2: KOOLITUSKAART ===
add_heading("OSA 2 — Koolituskaart HY1", 1)
add_normal("Prindi see osa eraldi — taskus või külmkapis.", bold=True, size=10)

add_table(
    ["Vali", "Sisu"],
    [
        ["OLULINE", "Kerge stress/uni/fookus — MITTE trauma/PTSD/kriis"],
        ["Kontroll", "Sa SAAD igal hetkel silmad avada"],
        ["STOP", "Distress > 8/10 või õhust väljas → lõpeta"],
    ],
)

add_heading("Protsess (kiire viide)", 2)
add_normal(
    "Ankur → Keha skaneering → Turvakoht → Suggestioon 8-12 min → "
    "Tagasi 5-1 → Vesi → Mis muutus?"
)

add_heading("Eesmärk (vali ÜKS)", 2)
add_table(
    ["Variant", "Valitud (märgi X)"],
    [
        ["Rahunemine", "[ ]"],
        ["Uni", "[ ]"],
        ["Fookus enne tegevust", "[ ]"],
        ["Debriefi ettevalmistus (5 min)", "[ ]"],
    ],
)

add_table(
    ["KEELATUD", "KOHUSTUSLIK"],
    [
        ["Trauma avamine ilma oskuseta", "Üks eesmärk korraga"],
        ["Sund / sa ei mäleta", "Nõusolek enne alustamist"],
        ["Diagnoosimine", "1 lause integratsioon/debrief"],
    ],
)

add_heading("Eneserefleksioon", 2)
add_table(
    ["Küsimus", "Jah", "Osaliselt", "Ei"],
    [
        ["Sain valida ühe selge eesmärgi?", "[ ]", "[ ]", "[ ]"],
        ["Sain igal hetkel lõpetada?", "[ ]", "[ ]", "[ ]"],
        ["Stress/distress langes?", "[ ]", "[ ]", "[ ]"],
        ["Kirjutasin mis muutus?", "[ ]", "[ ]", "[ ]"],
    ],
)

add_normal("Kuupäev: __________  Eesmärk: ________________________________")
add_normal("Stress enne: ___ /10    Stress pärast: ___ /10")

doc.add_page_break()

# === OSA 3: PLAANKONSPEKT ===
add_heading("OSA 3 — Plaankonspekt HY1-G (90 min)", 1)

add_heading("Eesmärgid", 2)
add_table(
    ["#", "Õpiväljund"],
    [
        ["1", "Selgitab mis hüpnoos on ja mis ei ole"],
        ["2", "Nimetab STOP reeglid"],
        ["3", "Teostab HY1 protokolli ise (15 min)"],
        ["4", "Juhib 5-min kiirvarianti"],
        ["5", "Teeb teadliku valiku spetsialisti poole"],
        ["6", "Kirjutab ühe integratsioonilause (debrief)"],
    ],
)

add_heading("Tunnikava", 2)
add_table(
    ["Aeg", "Faas", "Peamine tegevus"],
    [
        ["0-10", "Avamine", "Reeglid, STOP, üks sõna kehast, SWOT 2 min"],
        ["10-25", "Mis on hüpnoos?", "Definitsioon, mida EI ole, paarivestlus"],
        ["25-35", "Tõendus ja piirid", "Kus aitab / ei aita / spetsialist"],
        ["35-50", "Demo", "Grupi lühike HY1 — stress enne/pärast"],
        ["50-65", "Iseseisev", "12 min vaikne praktika, üks kerge eesmärk"],
        ["65-75", "Paariline", "5-min skript, vahetus"],
        ["75-85", "Debrief", "Üks lause ringis — mis muutus?"],
        ["85-90", "Lõpetus", "Kodutöö, HY1 kaart kaasa"],
    ],
)

add_heading("Koolitaja skript — võtmehetked", 2)
add_quote(
    "Täna ei muuda me teid kellegi teistsuguseks. Me harime fookust ja rahunemist. "
    "Te saate igal hetkel silmad avada. Üks eesmärk korraga."
)
add_quote(
    "Kui distress läheb 8 — STOP. Me ei ava raskeid minevikuid. "
    "See on eneseregulatsioon, mitte raviseanss."
)

add_heading("Kodutöö (7 päeva)", 2)
add_table(
    ["Päev", "Ülesanne"],
    [
        ["1-3", "1× HY1 (15 min) — logi stress enne/pärast"],
        ["4", "1× 5-min kiirvariant enne olulist tegevust"],
        ["5", "Loe SWOT uuesti läbi"],
        ["6-7", "1× paariline HY1 või debrief pärast sessiooni"],
    ],
)

add_heading("Tunnistus (valikuline)", 2)
add_normal('Mina, ___________________________, osalesin HY1-G koolitusel „Hünoteraapia algteadmised".')
add_normal("Kohustun kasutama STOP reegleid.")
add_normal("Kuupäev: ______________    Allkiri: ___________________________")

doc.add_page_break()

# === OSA 4: TEST NR 1 (Automaadi test vorming) ===
add_heading("OSA 4 — Hüpnoosi test nr. 1 (kaaslasega)", 1)
add_normal(
    "Iga mees saab seda õppida ja teha oma kaaslasele. "
    "Vorm: kaitseväe Automaadi test nr. 1 eeskujul.",
    bold=True,
)
add_heading("Õppetunni eesmärk", 2)
add_normal(
    "Kontrollida turvalise hüpnoosi juhtimise taset kaaslasega — "
    "15 min protokoll ilma OT vigadeta."
)
add_heading("OT vead (test mittesooritatud)", 2)
for item in [
    "Sundib silmi kinni / keelab lõpetamise",
    "Avab trauma ilma spetsialistita",
    "Jätkab kui distress > 8/10",
    "Diagnoosib või lubab meditsiinilist väidet",
    "Alustab ilma nõusolekuta",
]:
    add_bullet(item)

add_heading("Protseduur — peamised käsklused", 2)
add_table(
    ["#", "Käsklus"],
    [
        ["1", "Liikuge vaiksesse ruumi ja tehke ohutuskontroll"],
        ["2", "Lepi kokku üks eesmärk"],
        ["3", "Sul on alati kontroll. Ava silmad igal hetkel"],
        ["4", "Ankur — sisse 4, välja 6 × 3"],
        ["5", "Keha skaneering → Turvakoht → Suggestioon 8-12 min"],
        ["6", "Tagasi 5-1 → Vesi → Mis muutus?"],
    ],
)

add_heading("Hindamine", 2)
add_normal(
    "Sooritatud: 0 OT viga + min 14 punkti 18-st. "
    "Kaaslane kinnitab: Sain igal hetkel lõpetada."
)

doc.add_page_break()
add_heading("Allkiri — HY1 kasutamine", 2)
add_table(
    ["Roll", "Allkiri / kuupäev"],
    [
        ["Osaleja / praktik", "________________  __________"],
        ["Juht / koolitaja", "________________  __________"],
    ],
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"HY1 v1.2  |  {TODAY}  |  Unpluged-Al  |  "
    "hypnoteraapia-test-nr1-kaaslasega.md"
)
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
