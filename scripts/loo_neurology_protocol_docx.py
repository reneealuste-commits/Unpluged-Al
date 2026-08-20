#!/usr/bin/env python3
"""Genereerib NEUROLOGY PROTOCOL ostujuhendi Wordi dokumendina (Eesti tellimislingid)."""

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "/workspace/Neurology-Protocol-Ostujuhend.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0xB7, 0x1C, 0x1C)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)
LINK = RGBColor(0x05, 0x63, 0xC1)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    color = BLUE if level == 1 else DARK
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_normal(text, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_link_line(label, url, note=""):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ")
    add_hyperlink(p, url, url)
    if note:
        p.add_run(f"  —  {note}")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


# === KAANLEHT ===
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("NEUROLOGY PROTOCOL")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BLUE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Ajuvabaduse retsept — ostujuhend Eestist")
r2.font.size = Pt(14)
r2.font.color.rgb = GRAY

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Unpluged-Al  |  {TODAY}  |  v1.0").font.size = Pt(10)

add_table(
    ["Vali", "Andmed"],
    [
        ["Kood", "NP1 — Neurology Protocol"],
        ["Eesmärk", "Päevapõhine nootroopiline ja neurotoetav rutiin"],
        ["Ostustrateegia", "iFit (EE) + iHerb (rahvusvaheline) + Mycoland (seened)"],
        ["Hinnang starter pack", "~150–250 € (ilma punase valguse paneelita)"],
    ],
)

doc.add_page_break()

# === PROTSESS ===
add_heading("Päevaprotokoll", 1)
add_table(
    ["Aeg", "Toode", "Annus"],
    [
        ["Hommik", "Methylene Blue", "1 mg/kg"],
        ["Hommik", "Neuro Mag (Mg L-Threonate)", "3 × 144 mg"],
        ["Hommik", "Lion's Mane", "2 000 mg"],
        ["Hommik", "Turkey Tail", "2 000 mg"],
        ["Hommik", "Nikotiin", "muutuv, korduv"],
        ["Keskpäev", "Punase valguse teraapia (täiskeha)", "63 mW/cm² @ 20 min"],
        ["Keskpäev", "Peet", "3 000 mg"],
        ["Keskpäev", "Omega-3", "1 280 mg"],
        ["Keskpäev", "Pärm (nutritional yeast)", "1–3 spl"],
        ["Keskpäev", "Kurkumiin", "400 mg"],
        ["Õhtu", "Aeg looduses", "1 h"],
        ["Päikeseloojang", "Melatoniin / L-Glutathione", "200 mg / 250 mg"],
        ["Öö", "Meditatsioon / mindfulness", "35 min"],
        ["Varahommik", "Mikrodose (protokollis hägune)", "—"],
    ],
)

doc.add_page_break()

# === OSTUJUHT ===
add_heading("Lihtsaim ostustrateegia (3 kohta)", 1)
add_table(
    ["Pood", "Mida", "Miks"],
    [
        ["iFit.ee", "Magneesium, glutatioon, omega-3", "Eesti pood, kiire tarne"],
        ["iHerb.com", "Lion's Mane, Turkey Tail", "Suurim seenevalik, saadab Eestisse"],
        ["Mycoland.ee", "Kasvatuskomplektid", "Kohalik, seened ise kasvatamiseks"],
    ],
)

add_heading("HOMMIK — tellimislingid", 1)

add_heading("Methylene Blue (1 mg/kg)", 2)
add_link_line("VitaBlue (EU, Holland)", "https://vitablue.co/products/vitablue", "USP, tilgad, ~250 mcg/tilk")
add_link_line("LifeSolution.eu (Saksamaa)", "https://www.lifesolution.eu/en/products/methylenblau", "1% lahus, 100 ml")
add_link_line("ActiFolic Europe", "https://actifolic.com/product/buy-methylene-blue-europe/", "1% lahus, 30 ml")

add_heading("Neuro Mag — Magnesium L-Threonate (3 × 144 mg)", 2)
add_link_line("NOW Magtein @ iFit", "https://www.ifit.ee/et/a/now-foods-magtein-magnesium-l-threonate-90-vcaps-magneesium-l-tronaat", "2000 mg Magtein = 144 mg Mg / 3 kapslit")
add_link_line("Life Extension Neuro-Mag @ iFit", "https://www.ifit.ee/et/a/life-extension-neuro-mag-magnesium-l-threonate-90veg-caps-magneesium-treonaat")
add_link_line("OSAVI Magtein @ iFit", "https://www.ifit.ee/et/a/osavi-magtein-magnesium-l-threonate-90-vegan-caps")
add_link_line("Biotheka Magtein", "https://www.biotheka.ee/toode/magtein-magneesium-l-treonaat-90-kapslit-cytoplan-magtein-magnesium-l-threonate/")
add_link_line("Medpoint Magtein", "https://medpoint.ee/product/magtein-magnesium-l-threonate-90-kapslit/")

add_heading("Lion's Mane (2 000 mg)", 2)
add_link_line("iHerb — Lion's Mane kategooria", "https://www.iherb.com/c/lions-mane")
add_link_line("Real Mushrooms Lion's Mane 120 caps", "https://www.iherb.com/pr/real-mushrooms-lion-s-mane-mushroom-extract-powder-120-capsules/69422", "viljakeha ekstrakt, mitte mütseel")
add_link_line("NOW Foods Lion's Mane @ iHerb", "https://www.iherb.com/pr/now-foods-lion-s-mane-500-mg-60-veg-capsules/88819")

add_heading("Turkey Tail (2 000 mg)", 2)
add_link_line("iHerb — Turkey Tail kategooria", "https://www.iherb.com/c/turkey-tail")
add_link_line("Real Mushrooms Turkey Tail 90 caps", "https://www.iherb.com/pr/real-mushrooms-turkey-tail-mushroom-extract-powder-90-capsules/69428")
add_link_line("Host Defense Turkey Tail @ iHerb", "https://www.iherb.com/pr/host-defense-mushrooms-turkey-tail-120-vegetarian-capsules/16549")

add_heading("Nikotiin (muutuv)", 2)
add_normal(
    "Eestis on nikotiinitoodete e-müük piiratud (tubakaseadus). "
    "Osta füüsilisest poest: Circle K, Neste, Alexela, Rimi, Selver, Maxima — "
    "brändid VELO, ZYN, NOIS.",
    size=10,
)

doc.add_page_break()

add_heading("KESKPÄEV — tellimislingid", 1)

add_heading("Punase valguse teraapia (63 mW/cm², 20 min)", 2)
add_link_line("Kratomit.eu 1000W paneel", "https://www.kratomit.eu/led-red-and-infrared-panel-1000w-660-nm-850-nm-200-leds/", "660+850 nm, kuni 220 mW/cm²")
add_link_line("Dermfix RLF3000 (täiskeha)", "https://dermfix.com/products/rlf3000", "EU laos, 90 mW/cm² @ 15 cm")
add_link_line("Luxway Ruby 2.0 3600W", "https://www.luxway.eu/red-light-therapy-rlt/ruby-20-redlight-therapy-3600w.html", "720 LED, 660+850 nm")

add_heading("Peet (3 000 mg)", 2)
add_link_line("Vitabi punapeedi kapslid 6000 mg", "https://vitabi.ee/toode/punapeedi-kapslid-6000-mg-180-kapslit-6-kuu-jagu/", "1 kapsel = 6000 mg ekvivalent")
add_link_line("Biobutiik peedipulber", "https://biobutiik.ee/product/peedipulber/")
add_link_line("LIVIN punapeedi pulber", "https://www.livin.ee/p/urbanfood-punapeedi-pulber-urb1061", "1 spl ≈ 100 g värsket peeti")
add_link_line("It's Bio mahe peedipulber", "https://www.itsbio.ee/et/a/punapeedi-pulber-mahe-100g")

add_heading("Omega-3 (1 280 mg)", 2)
add_link_line("Nordic Naturals Ultimate Omega 1280 mg — Hind.ee", "https://www.hind.ee/s/nordic-naturals-ultimate-omega-3/", "võrdle hindu eri poodides")
add_link_line("iFit — Nordic Naturals", "https://www.ifit.ee/et/search?q=nordic+naturals+ultimate+omega")

add_heading("Pärm / Nutritional Yeast (1–3 spl)", 2)
add_link_line("Selver — BON VEGAN maitsepärm 125 g", "https://www.selver.ee/maitseparm-bon-vegan-125-g", "Eesti toode, B12")
add_link_line("LIVIN Engevita pärmihelbed", "https://www.livin.ee/p/engevita-parmihelbed-mk62")
add_link_line("It's Bio maitsepärm 200 g", "https://www.itsbio.ee/et/a/maitseparm-200g")
add_link_line("Bio4you maitsepärm", "https://bio4you.eu/et/bio-toit/biotoit-maitseained-ja-magustajad/maitsepaerm-125g")

add_heading("Kurkumiin (400 mg)", 2)
add_link_line("Ecosh Kurkumiin + Piperiin", "https://ecosh.ee/toode/kurkum-ekstrakt-95-kurkumiin-piperiin/", "4 kapslit ≈ 400 mg ekstrakti")
add_link_line("Nordic Ultimate Omega + Curcumin", "https://ecosupplements.ee/product/omega-curcumin-1200-mg-60-blode-kapsler/", "1200 mg omega-3 + 400 mg kurkumiin koos")

doc.add_page_break()

add_heading("ÕHTU JA ÖÖ — tellimislingid", 1)

add_heading("Melatoniin (200 mg protokollis)", 2)
add_link_line("Apotheka — melatoniin", "https://www.apotheka.ee/tooted/tervis/tervise-heaks/uni-ja-rahulik-meel/melatoniin")
add_link_line("Euroapteek — melatoniin", "https://www.euroapteek.ee/l/melatoniin")
add_link_line("ESI Melatonin 1,9 mg @ Apotheka", "https://www.apotheka.ee/melatonin-pura-esi-minitablett-1-9mg-n30-pmm0157016ee")
add_link_line("ICONFIT Melatoniin @ Euroapteek", "https://www.euroapteek.ee/l/melatoniin")

add_heading("L-Glutathione (250 mg)", 2)
add_link_line("NOW Glutathione 250 mg @ iFit", "https://www.ifit.ee/et/a/now-foods-glutathione-250mg-60-vcaps", "1 kapsel = 250 mg")
add_link_line("OSAVI Liposomal Glutathione 500 mg @ iFit", "https://www.ifit.ee/et/a/osavi-liposomal-glutathione-500-mg-60-vegan-caps", "2 kapslit = 500 mg")
add_link_line("OstroVit Glutathione @ iFit", "https://www.ifit.ee/et/a/ostrovit-glutathione-vege-90vcaps", "200 mg / kapsel")

add_heading("Meditatsioon (35 min)", 2)
add_link_line("Insight Timer (tasuta)", "https://insighttimer.com/")
add_link_line("Waking Up", "https://www.wakingup.com/")
add_link_line("Headspace", "https://www.headspace.com/")

doc.add_page_break()

# === SEENTE KASVATAMINE ===
add_heading("Seente ise kasvatamine", 1)
add_normal(
    "Odavam ja pikem plaan: kasvata Lion's Mane ja Turkey Tail ise. "
    "Värsked seened sobivad toiduks; kuivatatud türgi saba tee/tinktuuriks.",
    size=10,
)

add_heading("Lion's Mane — kasvatuskomplekt (algajale)", 2)
add_link_line("Mycoland Lion's Mane Growing Kit", "https://mycoland.ee/en/product/reishi-dowelswax-kopeeri-kopeeri-kopeeri-kopeeri-kopeeri-2/", "~18,90 €, tasuta tarne üle 60 €")
add_link_line("Mycoland Lion's mane dowels+wax (pakkidel)", "https://mycoland.ee/en/shop/", "otsi 'Lion's mane dowels+wax'")

add_normal("Sammud:", bold=True)
for step in [
    "Ava kott, lõika plastikusse X-kujuline lõige (~8 cm).",
    "Pihusta 2–3× päevas (komplektis on pihustuspudel).",
    "Temperatuur 12–21 °C, kaudne valgus, niiskus 85–95 %.",
    "Pinsid ilmuvad 7–14 päeva pärast; korista enne kollakaks muutumist.",
    "Pärast koristust leota blokki külmas vees 4–8 h → 2.–4. koristus.",
    "Saagikus: ~0,5–2 kg värsket seent komplekti kohta.",
]:
    add_bullet(step)

add_heading("Turkey Tail — variant A: siseruumides", 2)
add_link_line("Royal Flush Turkey Tail Kit (EU)", "https://royalflushmushrooms.com/products/turkey-tail-spray-n-grow-kit", "~17 €")
add_link_line("Näckrosgården Turkey Tail Growkit (Soome)", "https://de.nackrosgarden.com/products/turkey-tail-growkit", "~27 €, saadab EU-sse")

add_normal("Sammud: sama loogika — lõika kott, pihusta, hoia niiskelt. "
          "Kasvab plaatidena (fan-kujulised sabad). Saak 4–8 nädalat.", size=10)

add_heading("Turkey Tail — variant B: õues puupakkidel", 2)
add_link_line("Mycoland Libliktagela (Turkey Tail) dowels + vaha", "https://mycoland.ee/en/product/libliktagela-dowelswax/", "30 tüblit")
add_link_line("HomeGreen Turkey Tail 100 dowels", "https://www.homegreen.nl/en/mushroom-spawn/spawn-100-dowels/turkey-tail-trametes-versicolor-100-dowels", "~13 €, saadab EU-sse")

add_normal("Sammud:", bold=True)
for step in [
    "Võta värske lehtpuu (tamm, kask, saar), läbimõõt 15–50 cm.",
    "Puurista augud iga 10 cm tagant, löö tüblid sisse, kata vahaga.",
    "Hoia kottis soojas ~2 kuud (mütseseen valgeks).",
    "Mätta 2/3 mullasse, varjus. Esimene saak 6–12 kuu pärast.",
    "Kuivata, tee tee või tinktuur.",
]:
    add_bullet(step)

doc.add_page_break()

# === STARTER PACK ===
add_heading("Starter pack (~150–250 €)", 1)
add_table(
    ["Toode", "Koht", "Hind"],
    [
        ["Magtein 90 kapslit", "iFit", "~35–45 €"],
        ["Lion's Mane + Turkey Tail", "iHerb (Real Mushrooms)", "~60–80 €"],
        ["Omega-3 + kurkumiin", "Ecosh / Nordic", "~30–50 €"],
        ["Glutatioon + melatoniin", "iFit + Apotheka", "~25–35 €"],
        ["Maitsepärm + peedipulber", "Selver / LIVIN", "~10 €"],
        ["2× seenekomplekt", "Mycoland", "~38 €"],
    ],
)

add_heading("Praktilised märkused", 1)
for note in [
    "iHerb: tellimisel vali Eesti tarne; tellimus tavaliselt 5–14 päeva.",
    "iHerb sooduskoodid: otsi 'iHerb referral' — esimene tellimus sageli -10%.",
    "Mycoland: tasuta transport tellimustele üle 60 €.",
    "Punase valguse paneel: kontrolli mW/cm² mõõtmist kaugusel (protokoll: 63 mW/cm² @ 20 min).",
    "Metüleenisinine 1 mg/kg: arvuta kehakaalust (80 kg → 80 mg päevas).",
    "Seenekomplekt: alusta kasvatamist 7–10 päeva jooksul pärast kättesaamist.",
]:
    add_bullet(note)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"NP1 v1.0  |  {TODAY}  |  Unpluged-Al")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
