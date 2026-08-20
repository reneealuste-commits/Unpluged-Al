#!/usr/bin/env python3
"""NP1-LINN koolituskaart — kompaktne PDF + DOCX (prinditav, taskus)."""

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

sys.path.insert(0, str(Path(__file__).resolve().parent))
from loo_np1_kalender import generate_all  # noqa: E402

PDF_OUT = "/workspace/neurology-protocol-koolituskaart.pdf"
DOCX_OUT = "/workspace/Neurology-Protocol-Koolituskaart.docx"
TODAY = date.today().strftime("%d.%m.%Y")
ICS_ALL = "https://github.com/reneealuste-commits/Unpluged-Al/raw/main/np1-calendar/NP1-koik-meeldetuletused.ics"
CAL_PAGE = (
    "https://htmlpreview.github.io/?"
    "https://github.com/reneealuste-commits/Unpluged-Al/raw/main/np1-calendar/index.html"
)

RED = colors.HexColor("#B71C1C")
LIGHT = colors.HexColor("#FFEBEE")
DARK = colors.HexColor("#212121")
GRAY = colors.HexColor("#757575")


def build_pdf():
    doc = SimpleDocTemplate(
        PDF_OUT,
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title="NP1-LINN — Neurology Protocol koolituskaart",
        author="Unpluged-Al",
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=14,
        textColor=RED,
        fontName="Helvetica-Bold",
        spaceAfter=2,
    )
    sub = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=7.5, textColor=GRAY, spaceAfter=4)
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=9,
        textColor=RED,
        fontName="Helvetica-Bold",
        spaceBefore=3,
        spaceAfter=2,
    )
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=8, leading=10.5, textColor=DARK)
    link = ParagraphStyle("Link", parent=body, fontSize=7.5, leading=9, textColor=colors.HexColor("#1565C0"))
    tiny = ParagraphStyle("Tiny", parent=styles["Normal"], fontSize=6.5, textColor=GRAY, leading=8)

    cal = generate_all()
    story = []
    story.append(Paragraph("NP1-LINN — NEUROLOGY PROTOCOL · KOOLITUSKAART", title))
    story.append(Paragraph(f"Unpluged-Al · Linnas elav · {TODAY} · v1.0 · prindi — hoia taskus / külmkapi uksel", sub))
    story.append(HRFlowable(width="100%", thickness=0.8, color=RED))

    schedule = Table(
        [
            ["AEG", "PROTSESS", "ANNUSED / TEGEVUS"],
            ["07:00", "Hommik", "MB 1 mg/kg · Mag 3×144 mg · Lion's Mane 2 g · Turkey Tail 2 g · nikotiin"],
            ["12:30", "Keskpäev", "RLT 63 mW/cm² × 20 min · peet 3 g · O3 1280 mg · pärm 1–3 spl · kurkumiin 400 mg"],
            ["17:30", "Loodus", "1 h kõnd — Pääsküla raba · Nõmme · Kadriorg"],
            ["20:30", "Loojang", "Melatoniin 200 mg · L-glutathione 250 mg"],
            ["22:00", "Öö", "Meditatsioon 35 min"],
            ["05:30", "Varahommik", "Protokolli järgi"],
        ],
        colWidths=[18 * mm, 22 * mm, 142 * mm],
    )
    schedule.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("GRID", (0, 0), (-1, -1), 0.35, GRAY),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(Spacer(1, 4))
    story.append(schedule)

    story.append(Paragraph("OST — 3 KOHTA", h2))
    buy = Table(
        [
            ["POOD", "MIDA", "LINK"],
            ["iFit.ee", "Magtein, glutatioon, omega-3", "ifit.ee · Tartu mnt 52"],
            ["iHerb.com", "Lion's Mane, Turkey Tail", "iherb.com/c/lions-mane"],
            ["Mycoland.ee", "Seenekomplekt korterisse", "mycoland.ee · ~19 €"],
            ["Apotheka", "Melatoniin", "apotheka.ee"],
            ["Selver", "Maitsepärm", "selver.ee"],
            ["Circle K", "Nikotiin VELO/ZYN", "füüsiline pood"],
        ],
        colWidths=[28 * mm, 52 * mm, 102 * mm],
    )
    buy.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("GRID", (0, 0), (-1, -1), 0.35, GRAY),
            ]
        )
    )
    story.append(buy)

    story.append(Paragraph("KORTER — ühekordne setup", h2))
    for line in [
        "Köök: 7-päevane kapslite dosaator + maitsepärm + peedipulber",
        "Vannituba: punase valguse paneel ukse taga (20 min keskpäeval)",
        "Külmkapp: seenekomplekt enne avamist",
        "Akna äär: melatoniin + glutatioon (loojang)",
    ]:
        story.append(Paragraph(f"• {line}", body))

    story.append(Paragraph("SEENED KORTERIS", h2))
    story.append(
        Paragraph(
            "Lion's Mane köögis: X-lõige → pihusta 2× päevas → korista 14–21 päeva. "
            "Turkey Tail vannitoas või iHerb kapslid.",
            body,
        )
    )

    story.append(Paragraph("KALENDER — klõpsa, lisa meeldetuletus", h2))
    story.append(
        Paragraph(
            f'Apple/Outlook kõik korraga: <a href="{ICS_ALL}" color="#1565C0">NP1-koik-meeldetuletused.ics</a>',
            link,
        )
    )
    story.append(
        Paragraph(
            f'Mobiilinupud: <a href="{CAL_PAGE}" color="#1565C0">np1-calendar (Google)</a>',
            link,
        )
    )
    for t, url in cal["google"]:
        short = t.replace("NP1 ", "")
        story.append(
            Paragraph(f'• {short}: <a href="{url}" color="#1565C0">Lisa kalendrisse</a>', link)
        )

    story.append(Spacer(1, 6))
    story.append(
        Paragraph(
            "NP1-LINN v1.0 · Unpluged-Al · "
            "python3 scripts/loo_neurology_protocol_koolituskaart.py · "
            "neurology-protocol-koolituskaart.md",
            tiny,
        )
    )
    doc.build(story)
    print(f"Salvestatud: {PDF_OUT}")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "B71C1C")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_link_line(doc, label, url, note=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"{label}: ")
    run.font.size = Pt(8)
    add_hyperlink(p, url if len(url) < 55 else url[:52] + "...", url)
    if note:
        r2 = p.add_run(f"  {note}")
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_docx():
    cal = generate_all()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    RED_RGB = RGBColor(0xB7, 0x1C, 0x1C)
    GRAY_RGB = RGBColor(0x55, 0x55, 0x55)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("NP1-LINN — NEUROLOGY PROTOCOL")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RED_RGB

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(f"KOOLITUSKAART · Linnas elav · {TODAY}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY_RGB

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Prindi · voldi · hoia taskus või külmutuskapi uksel").font.size = Pt(9)

    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    rows = [
        ["AEG", "FAAS", "ANNUSED"],
        ["07:00", "Hommik", "MB 1 mg/kg · Mag 3×144 mg · Lion's Mane 2 g · Turkey Tail 2 g · nikotiin"],
        ["12:30", "Keskpäev", "RLT 20 min · peet 3 g · O3 1280 mg · pärm · kurkumiin 400 mg"],
        ["17:30", "Loodus", "1 h — Pääsküla · Nõmme · Kadriorg"],
        ["20:30", "Loojang", "Melatoniin 200 mg · glutatioon 250 mg"],
        ["22:00", "Öö", "Meditatsioon 35 min"],
        ["05:30", "Varahommik", "Protokolli järgi"],
    ]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if ri == 0:
                        run.bold = True

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run("OST (3 kohta)").bold = True
    h.runs[0].font.color.rgb = RED_RGB
    h.runs[0].font.size = Pt(9)
    for label, url, note in [
        ("iFit", "https://www.ifit.ee/et/p/kontakt", "Tartu mnt 52 — Mag, O3, glutatioon"),
        ("iHerb seened", "https://www.iherb.com/c/lions-mane", "Lion's Mane + Turkey Tail"),
        ("Mycoland", "https://mycoland.ee/en/shop/", "seenekomplekt ~19 €"),
        ("Apotheka", "https://www.apotheka.ee/tooted/tervis/tervise-heaks/uni-ja-rahulik-meel/melatoniin", "melatoniin"),
        ("Selver pärm", "https://www.selver.ee/maitseparm-bon-vegan-125-g", ""),
        ("VitaBlue MB", "https://vitablue.co/products/vitablue", "metüleenisinine EU"),
    ]:
        add_link_line(doc, label, url, note)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2.add_run("KALENDER — ühe klõpsuga").bold = True
    h2.runs[0].font.color.rgb = RED_RGB
    h2.runs[0].font.size = Pt(9)
    add_link_line(doc, "Kõik korraga (.ics)", ICS_ALL, "Apple / Outlook")
    add_link_line(doc, "Mobiilinupud", CAL_PAGE, "Google Calendar")
    for title, url in cal["google"]:
        add_link_line(doc, title, url, "Lisa kalendrisse")

    doc.add_paragraph()
    h3 = doc.add_paragraph()
    h3.add_run("Korter · seened · loodus").bold = True
    h3.runs[0].font.size = Pt(9)
    tips = [
        "Dosaator köögis · RLT vannitoas · seened külmkapis",
        "Lion's Mane: X-lõige, pihusta 2× päevas, korista ~14–21 p",
        "Nikotiin: Circle K / Neste (VELO, ZYN) — mitte e-pood",
        "1 h looduses: Pääsküla Ilmarise ring 3,4 km (rong Hiiu)",
    ]
    for tip in tips:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(tip)
        run.font.size = Pt(8)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(f"NP1-LINN v1.0 · Unpluged-Al · {TODAY}")
    fr.font.size = Pt(7)
    fr.font.color.rgb = GRAY_RGB

    doc.save(DOCX_OUT)
    print(f"Salvestatud: {DOCX_OUT}")


def build():
    generate_all()
    build_pdf()
    build_docx()


if __name__ == "__main__":
    build()
