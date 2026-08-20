#!/usr/bin/env python3
"""Genereerib NEUROLOGY PROTOCOL linnajuhendi Wordi dokumendina (korter, pakiautomaat, linnarajad)."""

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "/workspace/Neurology-Protocol-Linnajuhend.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0xB7, 0x1C, 0x1C)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    color = BLUE if level == 1 else DARK
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_normal(text, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_link_line(label, url, note=""):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ")
    add_hyperlink(p, url, url)
    if note:
        p.add_run(f"  —  {note}")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


# === KAANLEHT ===
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("NEUROLOGY PROTOCOL")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BLUE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Linnas elava inimese juhend")
r2.font.size = Pt(14)
r2.font.color.rgb = GRAY

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"NP1-LINN  |  Unpluged-Al  |  {TODAY}  |  v1.0").font.size = Pt(10)

add_normal(
    "Praktiline juhend korterielanikule: ostamine pakiautomaati, "
    "seened köögis, punane valgus väikeses ruumis, 1 h looduses linnas.",
    center=True,
    size=10,
)

add_table(
    ["Vali", "Andmed"],
    [
        ["Kood", "NP1-LINN — Neurology Protocol, linnavariant"],
        ["Sihtgrupp", "Korter, tööpäev, ühistransport, pakiautomaat"],
        ["Peamine strateegia", "1× iFit käiguga + iHerb pakiautomaati + Mycoland"],
        ["Eelarve starter", "~150–250 € (ilma valguspaneelita)"],
    ],
)

doc.add_page_break()

# === LINNA PÄEV ===
add_heading("Linnainimese päev — ajakava", 1)
add_table(
    ["Aeg", "Protokoll", "Linnalahendus"],
    [
        ["07:00", "Hommik: MB + Mag + seened + nikotiin", "Köögilaual dosaatorid; nikotiin taskus"],
        ["12:30", "Keskpäev: punane valgus + toidulisandid", "Kodus WFH või lõunapaus kodus; paneel vannitoas/ukse taga"],
        ["17:30", "1 h looduses", "Pääsküla raba / Nõmme / Kadriorg — rong või auto 20–30 min"],
        ["20:30", "Päikeseloojang: melatoniin + glutatioon", "Akna ääres, telefon ära"],
        ["22:00", "35 min meditatsioon", "Korter, kõrvaklapid, naabrite müra välja"],
        ["05:30", "Varahommik (protokoll)", "—"],
    ],
)

add_heading("Korteri setup — ühekordne", 1)
add_table(
    ["Tsoon", "Mida panna", "Miks"],
    [
        ["Köök / riiul", "Kapslite dosaator (7 päeva)", "Hommik kiire, ei mõtle"],
        ["Külmkapp", "Seenekomplekt (kuni avamiseni)", "Mütseseen vajab jahedust enne kasvatust"],
        ["Vannituba või magamistuba", "Punase valguse paneel + kinnitus", "20 min keskpäeval, privaatne"],
        ["Akna äär", "Melatoniin + glutatioon", "Loojanguga sünk"],
        ["Köögilaud", "Maitsepärm + peedipulber", "Sega lõunasse / smuutisse"],
    ],
)

doc.add_page_break()

# === OSTAMINE LINNAS ===
add_heading("Ostamine linnas — 3 sammuga", 1)
add_normal(
    "Eesmärk: minimaalselt poekäike. Üks käik linnas + kaks pakki pakiautomaati.",
    bold=True,
)

add_heading("Samm 1 — Üks käik korraga (sama päev)", 2)
add_table(
    ["Koht", "Aadress / ligipääs", "Mida võtta", "Link"],
    [
        [
            "iFit pood",
            "Tartu mnt 52, Tallinn (E–R 12–17, L 12–16)",
            "Magtein, glutatioon, omega-3",
            "ifit.ee",
        ],
        [
            "Apotheka",
            "Iga suurem keskus (Ülemiste, Kristiine, Rocca)",
            "Melatoniin",
            "apotheka.ee",
        ],
        [
            "Selver / Rimi",
            "Sinu naabrus",
            "Maitsepärm, värsked peedid",
            "selver.ee",
        ],
        [
            "Circle K / Neste",
            "Iga nurk",
            "VELO / ZYN nikotiin",
            "—",
        ],
    ],
)

add_link_line("iFit pood — Tartu mnt 52", "https://www.ifit.ee/et/p/kontakt", "5500+ toodet laos, kohe kaasa")
add_link_line("iFit pakiautomaat / Omniva", "https://www.ifit.ee/et/p/kohaletoimetamine", "Tasuta pakiautomaati al. 50 €")
add_link_line("Apotheka e-apteek", "https://www.apotheka.ee/tooted/tervis/tervise-heaks/uni-ja-rahulik-meel/melatoniin", "Pakiautomaat 1–2 päeva")
add_link_line("Selver — BON VEGAN maitsepärm", "https://www.selver.ee/maitseparm-bon-vegan-125-g")

add_heading("Samm 2 — Pakiautomaati (tellid korraga)", 2)
add_link_line("iHerb — Lion's Mane", "https://www.iherb.com/c/lions-mane", "Real Mushrooms, NOW Foods")
add_link_line("iHerb — Turkey Tail", "https://www.iherb.com/c/turkey-tail")
add_link_line("Real Mushrooms Lion's Mane 120 caps", "https://www.iherb.com/pr/real-mushrooms-lion-s-mane-mushroom-extract-powder-120-capsules/69422")
add_link_line("Real Mushrooms Turkey Tail 90 caps", "https://www.iherb.com/pr/real-mushrooms-turkey-tail-mushroom-extract-powder-90-capsules/69428")
add_link_line("Mycoland — 2× seenekomplekt", "https://mycoland.ee/en/shop/", "Lion's Mane + Turkey Tail, tasuta tarne üle 60 €")
add_link_line("Mycoland Lion's Mane kit", "https://mycoland.ee/en/product/reishi-dowelswax-kopeeri-kopeeri-kopeeri-kopeeri-kopeeri-2/", "~18,90 €")

add_heading("Samm 3 — EU post (harvem, 1× kvartal)", 2)
add_link_line("VitaBlue metüleenisinine", "https://vitablue.co/products/vitablue", "EU tarne 2–5 päeva")
add_link_line("LifeSolution metüleenisinine", "https://www.lifesolution.eu/en/products/methylenblau")
add_link_line("Kratomit punase valguse paneel", "https://www.kratomit.eu/led-red-and-infrared-panel-1000w-660-nm-850-nm-200-leds/", "Kompaktne, seinale")

doc.add_page_break()

# === TELLIMISLINGID PÄEVA JÄRGI ===
add_heading("Kõik tellimislingid — päeva järgi", 1)

add_heading("HOMMIK", 2)
add_link_line("NOW Magtein @ iFit", "https://www.ifit.ee/et/a/now-foods-magtein-magnesium-l-threonate-90-vcaps-magneesium-l-tronaat", "3 kapslit = 144 mg Mg")
add_link_line("Life Extension Neuro-Mag @ iFit", "https://www.ifit.ee/et/a/life-extension-neuro-mag-magnesium-l-threonate-90veg-caps-magneesium-treonaat")
add_link_line("Biotheka Magtein", "https://www.biotheka.ee/toode/magtein-magneesium-l-treonaat-90-kapslit-cytoplan-magtein-magnesium-l-threonate/")
add_link_line("VitaBlue MB", "https://vitablue.co/products/vitablue")
add_link_line("iHerb Lion's Mane", "https://www.iherb.com/c/lions-mane")
add_link_line("iHerb Turkey Tail", "https://www.iherb.com/c/turkey-tail")
add_normal("Nikotiin: Circle K, Neste, Rimi, Selver — VELO, ZYN, NOIS (e-pood keelatud).", size=10)

add_heading("KESKPÄEV", 2)
add_link_line("Kratomit RLT paneel (korterisse)", "https://www.kratomit.eu/led-red-and-infrared-panel-1000w-660-nm-850-nm-200-leds/", "Riputa uksele või seina")
add_link_line("Vitabi peedikapslid", "https://vitabi.ee/toode/punapeedi-kapslid-6000-mg-180-kapslit-6-kuu-jagu/")
add_link_line("LIVIN peedipulber", "https://www.livin.ee/p/urbanfood-punapeedi-pulber-urb1061")
add_link_line("Nordic Ultimate Omega — Hind.ee", "https://www.hind.ee/s/nordic-naturals-ultimate-omega-3/")
add_link_line("Ecosh kurkumiin", "https://ecosh.ee/toode/kurkum-ekstrakt-95-kurkumiin-piperiin/")
add_link_line("Omega + kurkumiin combo", "https://ecosupplements.ee/product/omega-curcumin-1200-mg-60-blode-kapsler/")
add_link_line("Selver maitsepärm", "https://www.selver.ee/maitseparm-bon-vegan-125-g")

add_heading("ÕHTU JA ÖÖ", 2)
add_link_line("Apotheka melatoniin", "https://www.apotheka.ee/tooted/tervis/tervise-heaks/uni-ja-rahulik-meel/melatoniin")
add_link_line("Euroapteek melatoniin", "https://www.euroapteek.ee/l/melatoniin")
add_link_line("NOW glutatioon 250 mg @ iFit", "https://www.ifit.ee/et/a/now-foods-glutathione-250mg-60-vcaps")
add_link_line("OSAVI liposomal glutatioon @ iFit", "https://www.ifit.ee/et/a/osavi-liposomal-glutathione-500-mg-60-vegan-caps")
add_link_line("Insight Timer (tasuta)", "https://insighttimer.com/")
add_link_line("Waking Up", "https://www.wakingup.com/")

doc.add_page_break()

# === PUNANE VALGUS KORTERIS ===
add_heading("Punane valgus korteris", 1)
add_normal(
    "Protokoll: 63 mW/cm², 20 min, täiskeha. Korteris piisab paneelist + peegeldavast seinast.",
    size=10,
)
for tip in [
    "Paneel vannitoa uksel või magamistoa seinale — 60–90 cm kaugus kehast.",
    "Seisa 15–20 min (riided võivad jääda peale; NIR läheb läbi kangast).",
    "Kompaktne 1000W paneel (~69×23 cm) mahub väiksesse korterisse.",
    "WFH: tee lõunapausil enne sööki; kontoritöö puhul tee kohe pärast tööpäeva.",
    "Naabrite müra: paneel on vaikne, ei sega kedagi.",
]:
    add_bullet(tip)

add_link_line("Kratomit 1000W paneel", "https://www.kratomit.eu/led-red-and-infrared-panel-1000w-660-nm-850-nm-200-leds/", "kuni 220 mW/cm²")
add_link_line("Dermfix RLF750 (kompaktne)", "https://dermfix.com/blogs/news/best-red-light-therapy-panel-europe-guide", "pool keha, väiksem")

# === 1H LOODUSES LINNAS ===
add_heading("1 tund looduses — linnarajad", 1)
add_normal("Protokoll: 1 h õues. Tallinnas — rong/metro/buss + jalg, ilma autota.", size=10)

add_heading("Tallinn (soovitused)", 2)
add_table(
    ["Koht", "Pikkus", "Kuidas minna", "Link"],
    [
        ["Pääsküla raba", "3,4–4 km (~1 h)", "Rong Hiiu, buss Kraavi tn", "tallinn.ee"],
        ["Nõmme — Glehni park", "~1 h ring", "Rong Nõmme", "puhkaeestis.ee"],
        ["Kadriorg", "~1 h", "Trammi 1/3", "—"],
        ["Pirita jõeorada", "~1 h", "Buss 1A, 5", "—"],
        ["Harku mets / Ülemiste", "~1 h", "Jalgsi kesklinnast", "—"],
    ],
)
add_link_line("Pääsküla raba — Tallinna koduleht", "https://www.tallinn.ee/et/paaskula", "Ilmarise ring 3,4 km, tasuta parkla")
add_link_line("Pääsküla raba matkajuhend", "https://nordicrent.ee/paaskula_raba_matkarada/")
add_link_line("Nõmme jalutuskäik", "https://puhkaeestis.ee/et/kuhu-minna/parandiga-metsapark-ja-vaikelinna-sarm-jalutuskaik-labi-nomme")

add_heading("Tartu / Pärnu / muud linnad", 2)
add_table(
    ["Linn", "1 h loodus", "Ligipääs"],
    [
        ["Tartu", "Emajõe roheline ring, Tähtvere mets", "Jalgsi kesklinnast"],
        ["Pärnu", "Rannametsa rada, Luitemaa (bussiga)", "Kesklinn + ratas/jalg"],
        ["Narva", "Narva jõeorada", "Jalgsi"],
        ["Üldine", "RMK lähim matkarada", "rmk.ee/mis-on-matkarada"],
    ],
)
add_link_line("RMK matkarajad", "https://www.rmk.ee/mis-on-matkarada")

add_normal(
    "Nipp: pane telefon reisirežiimi, 1 h kõndimist ilma podcastita — "
    "see on protokolli osa, mitte treening.",
    size=10,
)

doc.add_page_break()

# === SEENED KORTERIS ===
add_heading("Seened korteris — ainus realistlik variant", 1)
add_normal(
    "Linnakorteris ei kasvata puupakke. Kasvata mõlemad toas — kööginurk või vannituba.",
    bold=True,
)

add_heading("Lion's Mane köögis", 2)
add_link_line("Mycoland Lion's Mane kit", "https://mycoland.ee/en/product/reishi-dowelswax-kopeeri-kopeeri-kopeeri-kopeeri-kopeeri-2/", "~18,90 €")
for step in [
    "Pane komplekt köögilauale või riiulile — mitte otsepäikesele.",
    "Lõika X-lõige, pihusta 2× päevas (hommikul + õhtul enne magama).",
    "Kui korter kuiv (kütteperiood), kata lõdva kilega auguga.",
    "Korista, kui valge ja tihe; praad või supp — värskem kui kapsel.",
    "Leota pärast koristust 4–8 h → 2.–4. saak samast blokist.",
]:
    add_bullet(step)

add_heading("Turkey Tail vannitoas / panipaigas", 2)
add_link_line("Royal Flush Turkey Tail kit", "https://royalflushmushrooms.com/products/turkey-tail-spray-n-grow-kit", "~17 €, EU tarne")
add_link_line("Näckrosgården Turkey Tail", "https://de.nackrosgarden.com/products/turkey-tail-growkit", "~27 €")
add_link_line("Mycoland Turkey Tail dowels", "https://mycoland.ee/en/product/libliktagela-dowelswax/", "kui on rõdu/keldris ruum")
for step in [
    "Turkey Tail vajab niisket õhku — vannituba on ideaalne.",
    "Kasvab plaatidena; kuivata köögis ahjupannil 40°C või dehüdraatoris.",
    "Tee: 1–2 kuivatatud tükki + keev vesi, 15 min.",
    "Alternatiiv: jätka Turkey Tail kapsleid iHerbist, seened kasvata Lion's Mane'iks toiduks.",
]:
    add_bullet(step)

doc.add_page_break()

# === MUUD LINNAD ===
add_heading("Ostmine teistes Eesti linnades", 1)
add_table(
    ["Linn", "Füüsiline pood", "Pakiautomaat"],
    [
        ["Tallinn", "iFit Tartu mnt 52, Apotheka, Selver", "Kõik teenused"],
        ["Tartu", "Apotheka Lõunakeskus, Rimi, Selver", "iFit, iHerb, Mycoland"],
        ["Pärnu", "Apotheka Port Artur, Selver", "Sama"],
        ["Narva / Viljandi", "Apotheka, Maxima", "iFit + iHerb pakiautomaati"],
    ],
)
add_link_line("iFit e-pood (kogu Eesti)", "https://www.ifit.ee")
add_link_line("Apotheka", "https://www.apotheka.ee")
add_link_line("Euroapteek", "https://www.euroapteek.ee")

# === KUU EELARVE ===
add_heading("Linna starter pack", 1)
add_table(
    ["Toode", "Kust", "Hind", "Tarne"],
    [
        ["Magtein + glutatioon + omega-3", "iFit (üks käik või pakiautomaat)", "~80–100 €", "Kohe / 1–2 p"],
        ["Lion's Mane + Turkey Tail kapslid", "iHerb", "~60–80 €", "5–14 p"],
        ["2× seenekomplekt", "Mycoland", "~38 €", "3–5 p"],
        ["Maitsepärm + melatoniin", "Selver + Apotheka", "~15 €", "Sama päev"],
        ["Metüleenisinine", "VitaBlue", "~30–40 €", "3–7 p"],
        ["RLT paneel (valikuline)", "Kratomit", "~446 €", "1–2 näd"],
    ],
)

add_heading("Linna nipid", 1)
for tip in [
    "Telli iHerb + Mycoland samal nädalal — üks pakiautomaadi käik.",
    "iFit Tartu mnt 52: võta kõik korraga kätte, säästad postikulu.",
    "Seenekomplekt: alusta 7 päeva jooksul; hoia enne avamist külmkapis.",
    "Metüleenisinine 1 mg/kg: 75 kg → 75 mg; hoia köögikapis pimedas.",
    "Lõuna smuuti: peedipulber + maitsepärm + vesi — kiire keskpäevane portsjon.",
    "Meditatsioon 35 min: Insight Timer, režiim „Do Not Disturb“, naabrite müra kõrvaklappidega.",
    "Värsked peedid: Balti jaam (Tallinn), Nõmme turg, iga Selver — 1 peet ≈ 200 g.",
]:
    add_bullet(tip)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"NP1-LINN v1.0  |  {TODAY}  |  Unpluged-Al")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
