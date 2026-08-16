#!/usr/bin/env python3
"""Genereerib TLP + Ranger Boards perekonna plaankonspekti Wordi."""

from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "/workspace/Plaankonspekt-TLP-Ranger-Perekonna-Ehitamine.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0x0D, 0x47, 0xA1)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    color = BLUE if level == 1 else GRAY
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_normal(text, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


# Kaanleht
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PLAANKONSPEKT")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = BLUE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run("TLP + Ranger Boards stiil · 5-astmeline meeskonna ehitamine perekonnas").font.size = Pt(
    12
)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Renee Aluste · Unpluged-Al · {TODAY} · v1.0").font.size = Pt(10)

add_table(
    ["Vali", "Andmed"],
    [
        ["Alus", "TLP + Ranger School 20 Boards loogika"],
        ["Lõpptulemus", "Mees, naine, laps teavad rolli igas etapis"],
        ["Sihtgrupp", "Perejuhid, SOK, Kaitseliit, juhtimiskoolitus"],
    ],
)

doc.add_page_break()

add_heading("Põhimõte (nagu TLP-s)", 1)
for item in [
    "Kõigepealt saa missioon aru",
    "Anna varakult teada (warning order)",
    "Tee esialgne plaan, siis täpsusta",
    "Tunne oma inimesi (reconnaissance)",
    "Anna selge käsk (rollid + ülesanded)",
    "Jälgi ja korrigeeri (supervise & refine)",
]:
    add_bullet(item)
add_normal("1/3–2/3 reegel: juht planeerib 1/3, jätab 2/3 ettevalmistuseks.", bold=True)

ETAPID = [
    (
        "ETAPP 1 — Tutvustus (Reconnaissance)",
        "Luureinfo meeskonna kohta. Igaüks: haridus, hobid, tugevused, 1–2 lauset.",
        [
            ["MEES", "Juhib vooru", "Alustab ise. Kirjutab üles. Ei lase venida."],
            ["NAINE", "Aus osalemine", "Räägib selgelt. Ei vähenda ennast."],
            ["LAPS", "Enese nimetamine", "Ütleb mida oskab. Kuulab vanemaid."],
        ],
    ),
    (
        "ETAPP 2 — Eesmärgid ja ootused",
        "Missioon lühike. Ootused öeldud, mitte eeldatud.",
        [
            ["MEES", "Missiooni versioon", "Kirjutab üles. Ütleb lõpliku versiooni."],
            ["NAINE", "Reaalsus + emotsioon", "Ütleb otse mis töötab / mitte."],
            ["LAPS", "Ühine siht", "Kuulab. Vanem laps saab öelda."],
        ],
    ),
    (
        "ETAPP 3 — Rollid + toetus",
        "Põhiülesanne, toetus, asetäitja iga rolli juures.",
        [
            ["MEES", "Suund, turvatunne", "Määrab asetäitja. Ei delegeeri vastutust."],
            ["NAINE", "Kliima, rütm, lapsed", "Ütleb mida vajab. Ei võta suunda üle."],
            ["LAPS", "Vanusele vastav", "Teab ülesandeid. Austab rolle."],
        ],
    ),
    (
        "ETAPP 4 — Ülesannete jagamine",
        "Kes, mida, millal, mis on valmis.",
        [
            ["MEES", "Jagab ülesanded", "Küsib: Kas on selge?"],
            ["NAINE", "Võtab vastu", "Jah/ei + ettepanek. Ei ole ma arvan."],
            ["LAPS", "Täidab oma osa", "Teab mis on tema asi."],
        ],
    ),
    (
        "ETAPP 5 — Suhtlusplaan + jälgimine",
        "Check-in, konflikt, signaal, juhtimiskett.",
        [
            ["MEES", "Suhtlusplaan", "Ei kao konflikti ajal. Korrigeerib."],
            ["NAINE", "Info õigel ajal", "Ütleb otse. Ei hoia pingeid."],
            ["LAPS", "Konflikti lahendus", "Võib öelda kui raske."],
        ],
    ),
]

for title, desc, rows in ETAPID:
    add_heading(title, 1)
    add_normal(desc)
    add_table(["Kes", "Roll", "Nõutud tegevus"], rows)

doc.add_page_break()
add_heading("Rakendamine (14 päeva)", 1)
add_table(
    ["Päev", "Tegevus"],
    [
        ["1", "Tutvustuse voor — mees juhib"],
        ["2", "Missioon + ootused kirja"],
        ["3–4", "Rollid + toetus + asetäitjad"],
        ["5–7", "Ülesanded + check-in"],
        ["8–14", "Rütm. Päev 14: After Action"],
    ],
)

add_heading("Kuldreegel", 1)
add_normal(
    "Kõigepealt tundma õppida inimesi. Alles siis juhtida. "
    "Turvatunne = selge suund + rahulik järjekindlus.",
    bold=True,
)

add_heading("Juhi märkmed", 1)
for item in [
    "Extreme Ownership: selgus = juhi vastutus",
    "NVC: vajadused, mitte süüdistused",
    "HY1: mees juhib kaaslasele rahunemist (15 min)",
    "DP1: diplomaatia stressis",
    "Debrief: After Action päev 14",
]:
    add_bullet(item)

add_normal("")
add_normal("Koostas: Renee Aluste    Kuupäev: _______________")
add_normal("Mees: _______________    Naine: _______________")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"v1.0 · {TODAY} · Unpluged-Al · Selgus: Mees · Naine · Laps").font.size = Pt(9)

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
