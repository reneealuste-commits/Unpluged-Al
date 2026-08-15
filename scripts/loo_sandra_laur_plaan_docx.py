#!/usr/bin/env python3
"""Genereerib Wordi plaani: Sandra Laur pilot — Uks kool, uks protokoll."""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/workspace/Sandra-Laur-Pilot-Plaan.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

GREEN = RGBColor(0x1B, 0x5E, 0x20)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN
    return h


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_normal(text, bold=False, size=11, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_signature_block(title, name_placeholder):
    add_heading(title, level=2)
    add_table(
        ["Vali", "Sisu"],
        [
            ["Nimi", name_placeholder],
            ["Isikukood", "_________________________"],
            ["Kuupäev", TODAY],
            ["Allkiri", "_________________________"],
            ["Koht", "Tallinn / Tartu"],
        ],
    )


# === KAANLEHT ===
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PILOOTPLAAN")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = GREEN

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Üks kool — üks protokoll\nTartu haridus + Kaitseliit + aus juhtimine")
r2.font.size = Pt(14)
r2.font.color.rgb = GRAY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Versioon 1.0  |  {TODAY}  |  DigiDoc-valmis .docx").font.size = Pt(10)

doc.add_paragraph()
add_table(
    ["Vali", "Andmed"],
    [
        ["Projekt", "Pilot: Üks kool — üks protokoll"],
        ["Partner Tartus", "Sandra Laur, Tartu linnavolikogu, hariduskomisjoni esimees"],
        ["Kontakt", "sandra.laur@tartu.ee"],
        ["Toetaja", "[Sinu nimi], Tallinn — tsiviil-strateeg / projektide ehitaja"],
        ["Kestus", "1 semester (4–6 kuud)"],
        ["Staatus", "Ettepanek — allkirjastamiseks või heakskiitmiseks"],
    ],
)

doc.add_page_break()

# === 1. KOKKUVOTE ===
add_heading("1. Kokkuvõte", 1)
add_quote(
    "Üks Tartu gümnaasium + hariduskomisjon + malevkond — noorte juhtimine ja "
    "koostöö debrief-metoodikaga. Toetaja Tallinnist; Sandra Laur juhib Tartus."
)
add_normal(
    "See dokument on praktiline pilotplaan. See EI sisalda kanepi, ERMA ega "
    "Operation Mirror avalikku kampaaniat. Need teemad on eraldi ja ei kuulu "
    "selle piloti sisse.",
)

# === 2. MIKS ===
add_heading("2. Miks (alguseks)", 1)
add_bullet(
    "Eesti vajab juhte, kes tunnistavad vigu enne kui keegi teine — muidu "
    "meeskond vaikib (Remo Ojaste, kindral Riho Ühtegi)."
)
add_bullet(
    "Riigikaitse ja haridus ei tohi elada eraldi ministeeriumides — kool, "
    "kogukond ja malev peavad rääkima."
)
add_bullet(
    "Noored vajavad rolli ja vastutust, mitte ainult eksamit — talente tuleb "
    "hoida kodus (Sandra Lauri eesmärk)."
)
add_bullet(
    "Pilot on väike katse: enne suurt visiooni — üks tootav protokoll."
)

# === 3. MIDA PILOOT SISALDAB ===
add_heading("3. Mida pilot sisaldab", 1)
add_normal("Pilot koosneb neljast osast:", bold=True)
add_numbered = lambda t: doc.add_paragraph(t, style="List Number")
add_numbered("Üks Tartu gümnaasium nõustub 1 semestri programmiga.")
add_numbered("Tartu linnavolikogu hariduskomisjon (esimees Sandra Laur) annab poliitilise heakskiidu.")
add_numbered("Tartu malev / malevkond osaleb partnerina (2–4 korra semestris).")
add_numbered("Toetaja koostab plaani, protokollid ja koordineerib kohtumisi Tallinnist.")

add_heading("Programmi variandid (valitakse koos)", 2)
add_table(
    ["Variant", "Sisu", "Koormus koolile", "Soovitus"],
    [
        ["A — Debrief", "1x kuus 30 min; juht alustab; debrief kaart", "Madal", "Alustuseks"],
        ["B — Noored juhid", "4× töötuba: meeskond, vastutus, koostöö", "Keskmine", "Kui fookus noortel"],
        ["C — Kool + kogukond", "Kooli aed/projekt; malevkond aitab 3x", "Keskmine", "Kui fookus praktikal"],
    ],
)

# === 4. ROLLID ===
add_heading("4. Rollid ja vastutus", 1)
add_table(
    ["Roll", "Isik / üksus", "Kohustus", "Asukoht"],
    [
        ["Otsustaja Tartus", "Sandra Laur", "Valib kooli; komisjoni heakskiit; suhtleb kooliga", "Tartu"],
        ["Projektide ehitaja", "[Sinu nimi]", "Plaan, päevakord, protokollid, 2× kuus sparring", "Tallinn"],
        ["Kohapealne partner", "Tartu malevkond", "2–4 osalemist semestris", "Tartu"],
        ["Kool", "Direktor + 1–2 õpetajat", "Programm tunni/vabaaja raames", "Tartu"],
        ["Metoodika (vabatahtlik)", "Combat Ready / Remo Ojaste", "1× 2 h intro", "Tartu/Tallinn"],
    ],
)
add_normal(
    "Toetaja EI ole: volikogu liige, erakonna esindaja, maleva ülemus. "
    "Võrdne partnerlus — protokolliga.",
    bold=False,
)

# === 5. AJAKAVA ===
add_heading("5. Ajakava", 1)
add_table(
    ["Nadal", "Tegevus", "Koht"],
    [
        ["0", "45 min Zoom: kas miks klappib?", "Tallinn / Zoom"],
        ["1", "Plaani heakskiit (allkiri või e-kiri)", "—"],
        ["2", "Kohtumine: Laur + malev + kool (1,5 h)", "Tartu (toetaja kohal)"],
        ["3–20", "Programm koolis; 2x kuus Zoom Lauriga", "Tallinn"],
        ["4, 8, 12", "Debrief / töötuba koolis", "Tartu (toetaja 1× kuus)"],
        ["20", "Lõpp-debrief: jätkame või lõpetame", "Tartu / Zoom"],
    ],
)
add_normal("Reisid: 1× kuus Tartusse + 1 lisareis kickoff kuul.", bold=True)

# === 6. KPI ===
add_heading("6. Mõõtmised (edu kriteeriumid)", 1)
add_table(
    ["KPI", "Sihttase 6 kuu lõpus"],
    [
        ["Allkirjastatud kokkulepe (kool + komisjon)", "1 tk"],
        ["Läbiviidud debrief või töötuba", "min 3"],
        ["Kirjalikud protokollid", "min 4"],
        ["Ühisotsus: kas jätkame?", "Jah/ei dokumenteeritud"],
    ],
)
add_quote("Edu = üks tootav protokoll, mida saab korrata teises koolis. Mitte meedia ega loosungid.")

# === 7. PIIRID ===
add_heading("7. Piirid ja väljadused", 1)
limits = [
    "Ei kanepi, ei reguleeritud turu poliitikat, ei Operation Mirror avalikku kampaaniat.",
    "Ei Isamaa valimislogot projektil — projektipõhine, mitte parteiline.",
    "Kui tekib poliitiline tüli — projekt jääb kooli tasemele või peatub.",
    "Toetaja teeb tsiviilkihti; salastatus ja riigisaladus ei kuulu projekti.",
    "Kumbki pool võib 1 kuu etteteatamisega välja astuda.",
]
for lim in limits:
    add_bullet(lim)

# === 8. DEBRIEF ===
add_heading("8. Debrief metoodika (variant A)", 1)
add_normal("Iga debrief (30 min, min 3 osalejat):", bold=True)
for step in [
    "Mis läks HÄSTI?",
    "Mis läks VALESTI? — juht alustab.",
    "Mis oli MINU otsuse viga?",
    "Mida teeme TEISITI? — üks otsus + omanik + tähtaeg.",
    "Kuidas mõjutas see kooli / noori?",
]:
    add_numbered(step)
add_normal("Manus: debrief-kaart-malevapealik.pdf", size=10)

# === 9. SONUM ===
add_heading("9. Ettepanek Sandra Laurile", 1)
add_normal(
    "Sandra, panin kirja ühe konkreetse asja, kuidas saaksin sind toetada Tallinnist. "
    "Pakun piloti Üks kool — üks protokoll: sinu hariduskomisjon + üks gümnaasium + "
    "malevkond. Fookus: juhtimine, koostöö, debrief. Mina teen plaani ja protokollid; "
    "sina juhid Tartus. Suured majandusideed jätan kõrvale. Kas 45 min Zoom?",
    size=11,
)

doc.add_page_break()

# === 10. ALLKIRJAD (DigiDoc) ===
add_heading("10. Heakskiit ja allkirjad", 1)
add_normal(
    "Allpool olevad väljad on mõeldud digitaalseks allkirjastamiseks (DigiDoc, Smart-ID "
    "või ID-kaart). Dokument jõustub pärast mõlema poole nõustumist.",
    size=10,
)

add_signature_block(
    "10.1 Toetaja (projektide ehitaja, Tallinn)",
    "[Sinu nimi]",
)
add_signature_block(
    "10.2 Partner (Tartu linnavolikogu, hariduskomisjoni esimees)",
    "Sandra Laur",
)

add_heading("10.3 Kool (kui pilot kinnitatakse)", 2)
add_table(
    ["Vali", "Sisu"],
    [
        ["Kooli nimi", "_________________________"],
        ["Direktor", "_________________________"],
        ["Kuupäev", "_________________________"],
        ["Allkiri", "_________________________"],
    ],
)

add_heading("10.4 Valitud programmi variant", 2)
add_table(
    ["Variant", "Valitud (märgi X)"],
    [
        ["A — Debrief", "[ ]"],
        ["B — Noored juhid", "[ ]"],
        ["C — Kool + kogukond", "[ ]"],
    ],
)

# Footer note
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Dokument genereeritud {TODAY}  |  Pilot v1.0  |  "
    "Seotud: sandra-laur-visioon-jagamine.md, debrief-kaart-malevapealik.pdf"
)
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
