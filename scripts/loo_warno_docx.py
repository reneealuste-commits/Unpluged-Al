#!/usr/bin/env python3
"""Genereerib WARNO 001 Wordi dokumendi OPORD v1.1 pohjal."""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "/workspace/WARNO-001-Operation-Mirror.docx"
TODAY = date.today().strftime("%d.%m.%Y")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

GREEN = RGBColor(0x1B, 0x5E, 0x20)
ORANGE = RGBColor(0xE6, 0x5C, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN if level == 1 else ORANGE
    return h


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_normal(text, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


# === KAANLEHT ===
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("WARNO 001")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ORANGE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("HOIATKORRALDUS\nOperation Mirror + Olessanded alluksustele")
r2.font.size = Pt(14)
r2.font.color.rgb = GRAY

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Viide: OPORD v1.1  |  {TODAY}  |  SISE  |  Faas 0").font.size = Pt(10)

add_table(
    ["Vali", "Andmed"],
    [
        ["Operatsioon", "OPERATION MIRROR / OLESSANDED"],
        ["WARNO number", "001"],
        ["OPORD", "OPORD-sise.md v1.1"],
        ["Faas", "0 — pilot (Tartu, Sakala, Jogeva)"],
        ["Jargmine", "WARNO 002 after H+90 (~2026-12-15)"],
    ],
)

add_normal(
    "WARNO annab aega valmistuda. Ei asenda OPORD-i. Kanep/ERMA — taustal, mitte Faas 0 avalik sonum.",
    bold=True,
)

doc.add_page_break()

# 1 SITUATION
add_heading("1. Situation (luhike)", 1)
add_bullet("Juhtimine: vaikimise kultuur — debrief prioriteet #1")
add_bullet("Tsiviil-kaitse side nork — malevapealik vajab maastiku rolli")
add_bullet("ERMA/kanep: sise-arhiiv; partneritel korval Faas 0-s")

add_heading("Sobralikud joud — WARNO prioriteet", 2)
add_table(
    ["Joud", "Prioriteet", "Roll"],
    [
        ["Tartu — Sandra Laur", "Korge", "1 kool pilot"],
        ["Sakala — Sverre", "Keskmine", "KOV + kriis"],
        ["Jogeva malev", "Madal", "Kaardistus"],
        ["3 malevapealikku", "Korge", "Debrief rollout"],
    ],
)

# 2 MISSION
add_heading("2. Mission", 1)
add_normal(
    "2026 Q3-Q4: kavita 3 maleva pilot — debrief + KOV kohtumine + uks tsiviilprojekt "
    "maleva kohta — toestada Olessanded mudel enne riigilaia skaleerimist.",
    bold=True,
)
add_heading("Faas 0 loppkriteeriumid", 2)
add_table(
    ["#", "Kriteerium"],
    [
        ["1", "3 malevapealikku — min 1 debrief (Annex A)"],
        ["2", "3 kvartalikohtumist KOV-ga (protokoll)"],
        ["3", "1 tootav tsiviilprojekt — Tartu kool (prioriteet)"],
        ["4", "0 poliitilist skandaali — kanep valjas"],
    ],
)

# 3 EXECUTION
add_heading("3. Execution", 1)
add_heading("3.1 Ajajoon", 2)
add_table(
    ["Mark", "Kuupaev", "Tegevus"],
    [
        ["WARNO 001", "2026-08-15", "Valjastamine"],
        ["H-14", "2026-08-29", "Debrief kaart 3 malevale"],
        ["H-7", "2026-09-05", "Laur Zoom — pilot GO/NO-GO"],
        ["H-Hour", "2026-09-15", "Esimene kvartalikohtumine (Tartu)"],
        ["H+30", "2026-10-15", "1. debrief raportid"],
        ["H+90", "2026-12-15", "Faas 0 review — WARNO 002"],
    ],
)

add_heading("3.2 Eelulesanded", 2)
add_table(
    ["Uksus", "Peamine ulesanne", "Tahtaeg"],
    [
        ["Strateegia omanik", "Laur docx; 3 malevapealikku; debrief saatmine", "H-14"],
        ["Tartu (Laur)", "Zoom; kool; komisjoni heakskiit; kohtumine", "H-Hour+"],
        ["Sakala", "Malev + Sverre; Viljandi KOV", "H+30"],
        ["Jogeva", "Maastiku kaardistus 1 lk", "H+30"],
        ["Koik malevad", "1x debrief; 1x KOV; 1 lk raport", "H+30"],
    ],
)

add_heading("3.3 Koordineerimisreeglid", 2)
for rule in [
    "Debrief enne uusi programme — Annex A kohustuslik",
    "Protokollita kohtumist = toimumata",
    "Kanep/ERMA ei ole WARNO 001 sonum",
    "Laur = projekt; malevapealik = operatiivjuht",
]:
    add_bullet(rule)

# 4 ADMIN
add_heading("4. Administration and Logistics", 1)
add_table(
    ["Materjal", "Fail", "Kasutus"],
    [
        ["Debrief kaart", "debrief-kaart-malevapealik.pdf", "Kohe"],
        ["Tartu pilot", "Sandra-Laur-Pilot-Plaan.docx", "Laur"],
        ["OPORD", "OPORD-sise.md", "Taielik korraldus"],
        ["Malev brief", "olessanded-alluksustele.md", "Malevapealik"],
    ],
)

# 5 COMMAND
add_heading("5. Command and Signal", 1)
add_normal("Juhtimisstandard (kohustuslik): Extreme Ownership · Debrief · Ausus > positsioon", bold=True)
add_normal("Kaskude ahel: Strateegia omanik -> malevapealik (operatiiv); Laur (Tartu uks).", size=10)

# 6 90 PAEVA
add_heading("6. Plaan — 90 paeva", 1)
add_table(
    ["Paev", "Tegevus", "Tulemus"],
    [
        ["1-7", "Kontaktid; Laur docx", "3 kontakti"],
        ["8-14", "Debrief saatmine; Laur Zoom", "Pilot GO/NO-GO"],
        ["15-30", "Tartu kohtumine; Sakala intro", "Protokoll #1"],
        ["31-60", "Programm + debrief + Jogeva", "3 raportit"],
        ["61-90", "Faas 0 hindamine", "WARNO 002 draft"],
    ],
)

doc.add_page_break()
add_heading("7. Kinnitused", 1)
add_table(
    ["Roll", "Allkiri / kuupaev"],
    [
        ["Strateegia omanik — WARNO valjastaja", "________________  " + TODAY],
        ["Tartu partner (Sandra Laur)", "________________  __________"],
        ["Malevapealik (pilot)", "________________  __________"],
    ],
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"WARNO 001  |  {TODAY}  |  Seotud: OPORD-sise.md, WARNO-001-sise.md")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
