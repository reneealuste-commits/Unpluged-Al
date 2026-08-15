#!/usr/bin/env python3
"""Genereerib Wordi dokumendi: Sinu valik juhib turgu."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/workspace/Sinu-Valik-Juhib-Turgu.docx"

doc = Document()

# Marginaalid
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E7D32")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    return h


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    return p


def add_bullet(text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part and bold_part in text:
        before, after = text.split(bold_part, 1)
        if before:
            p.add_run(before)
        r = p.add_run(bold_part)
        r.bold = True
        if after:
            p.add_run(after)
    else:
        p.add_run(text)
    return p


def add_numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    return p


def add_normal(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════
# KAANLEHT
# ═══════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SINU VALIK JUHIB TURGU")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Kuidas iga inimene loob nõudlust,\nmis muudab Eesti majandust")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.add_run("Roheline majandus · Tarbijavõim · 5-aasta plaan").font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════
# SISUKORD
# ═══════════════════════════════════════════
add_heading("Sisukord", 1)
toc_items = [
    "1. Üks lause, mis kõike muudab",
    "2. Iga ost on hääl",
    "3. Nõudlus enne pakkumist",
    "4. Mida SINA saad teha — täna",
    "5. Hakkame ise kasvatama",
    "6. Kas AI aitab? (aus vastus)",
    "7. Parimad ettepanekud — haritud vaade",
    "8. 5-aasta tee: samm-sammult",
    "9. Lõppsõna: sina oled turu juht",
]
for item in toc_items:
    add_normal(item)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. ÜKS LAUSE
# ═══════════════════════════════════════════
add_heading("1. Üks lause, mis kõike muudab", 1)
add_quote("Sa ei oota poliitikut. Sa ei oota ettevõtjat. Sa hakkad ise — ja turg järgneb sinu järel.")

add_normal(
    "See dokument ei ole manifest ega propaganda. See on lihtne selgitus: "
    "kuidas igaühel meist on võim muuta seda, mida Eesti turul pakutakse — "
    "ilma erakonna liikmeks saamata, ilma Brüsseli luba küsimata, ilma miljonitega."
)

add_normal(
    "Põhimõte on vana ja tõestatud: kui piisavalt inimesi teeb sama valiku, "
    "tekib nõudlus. Kui nõudlus on suur, tuleb pakkumine. Kui pakkumine on kohalik, "
    "jääb raha Eestisse. Kui raha jääb Eestisse, saavad teised inimesed tööd, "
    "maksud lähevad koolidesse ja päikeseparkidesse."
)

add_normal("See on demokraatia igapäevases vormis. Mitte valimised iga 4 aasta tagant — vaid iga ostuga.", bold=True)

# ═══════════════════════════════════════════
# 2. IGA OST ON HÄÄL
# ═══════════════════════════════════════════
add_heading("2. Iga ost on hääl", 1)

add_normal("Kujuta ette, et sul on rahakotis 100 eurot. Sa otsustad, kuhu see läheb:")

add_table(
    ["Sinu valik", "Kuhu raha läheb", "Mida sa hääletad"],
    [
        ["Ostad välismaa veebipoest (cbditaly.ee)", "Itaaliasse / Hollandi", "Ma tahan, et teised rikas saaksid"],
        ["Ostad mustalt müüjalt", "Kellelegi, kellel pole arvet", "Ma ei taha makse ega kvaliteeti"],
        ["Ostad Eesti toodet (kui oleks)", "Eesti farmerile, töötajale, riigile", "Ma tahan kohalikku tööstust"],
        ["Kasvatad ise (tööstuskanep)", "Sinu enda taskusse + naabrile tööd", "Ma olen tootja, mitte ainult tarbija"],
    ],
)

add_normal("See ei ole moraalne etteheide. See on matemaatika.", bold=True)
add_normal(
    "Kui 10 000 inimest maksab kuus 15 eurot välismaale, on see 1,8 miljonit eurot aastas, "
    "mis lahkub Eestist. Kui sama 10 000 inimest ostab Eesti toodet, on see 1,8 miljonit, "
    "mis jääb siia — ja millest 20–30% läheb riigikassasse."
)

add_quote("Sa ei pea kedagi veenma. Sa pead ainult ise valima. Ja rääkima teistele, miks sa valisid.")

# ═══════════════════════════════════════════
# 3. NÕUDLUS ENNE PAKKUMIST
# ═══════════════════════════════════════════
add_heading("3. Nõudlus enne pakkumist", 1)

add_normal(
    "Turg ei tööta nii, et keegi ehitab kõigepealt tehase ja siis ootab kliente. "
    "Turg töötab nii:"
)

steps = [
    "INIMENE tahab midagi → ta küsib, otsib, küsib uuesti",
    "10 inimest küsib → üks ettevõtja kuuleb",
    "100 inimest küsib → ettevõtja avab poe",
    "1000 inimest küsib → tekib tööstus",
    "10 000 inimest küsib → tekib seadus ja reguleerimine",
]
for s in steps:
    add_bullet(s)

doc.add_paragraph()
add_normal("Me oleme praegu sammus 1–2 ja 3 vahel.", bold=True)
add_normal(
    "Eestis on juba tuhandeid inimesi, kes ostavad kanepitooteid — aga peaaegu kõik raha "
    "läheb välja. See tähendab: nõudlus ON OLEMAS. Pakkumine on vale — väljastpoolt, "
    "reguleerimata, ilma maksudeta."
)

add_normal("Meie ülesanne ei ole luua nõudlust nullist. Meie ülesanne on suunata olemasolev nõudlus Eestisse.", bold=True)

# ═══════════════════════════════════════════
# 4. MIDA SINA SAAD TEHA
# ═══════════════════════════════════════════
add_heading("4. Mida SINA saad teha — täna", 1)

add_heading("Tarbijana (kohe)", 2)
add_bullet('Küsi igas poest: "Kas teil on Eesti toodetud kanepitooted?" - isegi kui vastus on "ei", küsimus registreerub')
add_bullet("Osta tööstuskanepi tooteid, mis JUBA on olemas: kanepiseemneid, õli, kiust kotte, kreeme")
add_bullet("Ära osta mustalt, kui sul on valik — must turg ei ehita midagi")
add_bullet("Räägi 3 sõbrale, mida sa tead — see on tasuta turundus")

add_heading("Kogukonnana (1–3 kuud)", 2)
add_bullet('Loo Facebook/WhatsApp grupp: "Eesti kanepitooted - kust osta?"')
add_bullet('Kutsu kokku 10 inimest ja kirjuta ühine kiri KOV-ile: "Me soovime kohalikku toodangut"')
add_bullet("Külasta talu või puukooli — näita, et nõudlus on reaalne")

add_heading("Tootjana (6–12 kuud)", 2)
add_bullet("Alusta tööstuskanepiga — see on TÄNA legaalne (<0,3% THC)")
add_bullet("Taotle LEADER toetust puukooli või väikefarmi jaoks")
add_bullet("Liitu kasvatajate kooperatiiviga või loo see")
add_bullet("Osale kutsekooli kursusel — oskus on uus töökoht")

add_table(
    ["Roll", "Mida teed", "Mõju"],
    [
        ["Tarbija", "Küsid, valisid, räägid", "Nõudlus suureneb"],
        ["Kogukond", "Grupp, kiri, kohtumine", "Otsustajad kuulevad"],
        ["Tootja", "Kasvatad, müüd, ekspordid", "Pakkumine tekib"],
        ["Kõik koos", "5 aasta plaan", "Turg muutub"],
    ],
)

# ═══════════════════════════════════════════
# 5. HAKKAME ISE KASVATAMA
# ═══════════════════════════════════════════
add_heading("5. Hakkame ise kasvatama", 1)

add_normal(
    "See ei tähenda igale pööningule lampi. See tähendab nutikat, legaalset, "
    "kogukondlikku tootmist — nagu meie esivanemad tegid mett, õunu ja kartulit."
)

add_heading("Mis on TÄNA legaalne?", 2)
add_table(
    ["Tegevus", "Lubatud?", "Märkus"],
    [
        ["Tööstuskanepi kasvatus (<0,3% THC)", "JAH", "Eesti on EL tipus"],
        ["Kanepiseemnete müük toiduks", "JAH", "Poodides olemas"],
        ["Kiust riide/tootmine", "JAH", "Eksportpotentsiaal"],
        ["CBD kosmeetika tootmine", "JAH", "Kreemid, õlid, seebid"],
        ["Hampcrete ehitus", "JAH", "Uus ehitusmaterjal"],
        ["Täiskasvanute kanep (THC)", "EI (praegu)", "Aasta 3–4 reguleerimine"],
    ],
)

add_heading("Kuidas alustada — 5 sammu", 2)
add_numbered("Vali maatükk või liitu puukooliga — alusta väikesest (0,5–2 ha)")
add_numbered("Taotle PRIA-lt tööstuskanepi litsentsi — see on bürokraatia, aga tehtav")
add_numbered("Osta sertifitseeritud seemned — ERMA võrgust või otse")
add_numbered("Taotle LEADER toetust (kuni 50 000 €) — KOV aitab paberitega")
add_numbered("Müü seemneid, kiudu või õli — esimene tulu tuleb enne, kui growshop avab")

add_quote("Iga Eesti farmer, kes täna kanepit külvab, on turu juht. Ta ei oota luba — ta loob pakkumist.")

# ═══════════════════════════════════════════
# 6. KAS AI AITAB?
# ═══════════════════════════════════════════
add_heading("6. Kas AI aitab? (aus vastus)", 1)

add_normal(
    "Jah — aga mitte nii, nagu Hollywood kujutab. AI ei kasvata kanepit sinu eest. "
    "AI aitab sind teha targemaks, kiiremaks ja vähem üksi."
)

add_heading("Kus AI TÕESTI aitab", 2)
add_table(
    ["Valdkond", "Mida AI teeb", "Näide"],
    [
        ["Bürokraatia", "Aitab täita taotlusi", "LEADER, EAS, PRIA vormid"],
        ["Turuanalüüs", "Näitab nõudlust", "Kus ostetakse, mis hinnaga"],
        ["Kasvatus", "Optimeerib niiskust, valgust", "IoT sensorid + AI hoiatused"],
        ["Kvaliteet", "Tuvastab defektid", "Kaamera + AI sorteerib"],
        ["Turundus", "Kirjutab tekste, tõlgib", "Eesti brändi loomine"],
        ["Haridus", "Õpetab lastele ja farmeritele", "Interaktiivsed kursused"],
        ["Planeerimine", "5-aasta finantsmudel", "Excel + AI stsenaariumid"],
    ],
)

add_heading("Kus AI EI aita (ja ei peaks)", 2)
add_bullet("AI ei asenda inimest, kes maa peal töötab — käed, silmad, kogemus")
add_bullet("AI ei tee poliitilist otsust sinu eest — sa pead ise valima ja hääletama")
add_bullet("AI ei ehita usaldust — seda teeb aus toode ja inimlik teenindus")
add_bullet("AI ei tohiks otsustada, kes saab kanepit ja kes mitte — see on inimeste otsus")

add_normal("Parim AI kasutus selles projektis:", bold=True)
add_normal(
    "AI on nagu nutikas abiline, kes aitab farmeril, ettevõtjal ja tarbijal "
    "teha paremaid otsuseid — aga otsuse teeb ikka inimene. "
    "Robotid (Tesla Optimus jms) tulevad hiljem, kui turg on suurem ja masinad "
    "on liisinguga ostetud ning kliendid on need kinni maksnud."
)

# ═══════════════════════════════════════════
# 7. PARIMAD ETTEPANEKUD
# ═══════════════════════════════════════════
add_heading("7. Parimad ettepanekud — haritud vaade", 1)

add_normal(
    "Pärast majandusanalüüsi, EL regulatsiooni, Saksamaa kogemust ja Eesti "
    "võimaluste hindamist — siin on minu parimad soovitused:"
)

add_heading('Strateegia: "Nõudlus esmalt, pakkumine järgneb"', 2)

add_table(
    ["#", "Ettepanek", "Miks see töötab", "Aeg"],
    [
        ["1", 'Tarbijaliikumine: "Osta Eesti"', "Suunab olemasoleva nõudluse", "Kohe"],
        ["2", "10 puukooli piloot (LEADER)", "Loob pakkumise aluse", "6–12 kuud"],
        ["3", "ERMA AS loomine", "Riiklik kvaliteetne alternatiiv", "12 kuud"],
        ["4", "Kutsekooli eriala", "Tootjad tulevad, mitte imporditakse", "12–18 kuud"],
        ["5", "Growshop võrgustik", "Asendab välismaa poed", "18–24 kuud"],
        ["6", "Reguleeritud turg (THC)", "Maksud + ohutus", "36–48 kuud"],
        ["7", "AI + IoT kasvatuses", "Efektiivsus, mitte asendus", "Pidev"],
        ["8", "Eesti Tuleviku Fond", "Kasum → päike, robotid, basic income", "A5+"],
    ],
)

add_heading("Mida MITTE teha", 2)
add_bullet('Ära oota, kuni poliitikud "lubavad" — alusta tööstuskanepiga täna')
add_bullet('Ära ehita kõike korraga — 3 puukooli, mitte 300')
add_bullet('Ära valetage ("kanep ravib kõike") — ausus loob usaldust')
add_bullet("Ära ignoreeri noori — alla 25 on kanep ohtlikum, see on fakt")
add_bullet("Ära tee seda ilma kogukonnata — üksik wolf ei muuda turgu")

add_heading("Mis on minu #1 soovitus?", 2)
add_normal(
    'Alusta tarbijaliikumisest. See on kõige odavam, kiireim ja võimsaim samm. '
    'Kui 500 inimest kirjutab KOV-ile "me soovime Eesti kanepitoodet", '
    'siis see on tugevam signaal kui ükski lobigrupp Brüsselis.',
    bold=True,
)

# ═══════════════════════════════════════════
# 8. 5-AASTA TEE
# ═══════════════════════════════════════════
add_heading("8. 5-aasta tee: samm-sammult", 1)

add_table(
    ["Aasta", "Sina (tarbija)", "Kogukond", "Tootja", "Tulemus"],
    [
        ["2026", "Küsid, valisid tööstuskanepit", "Grupp 100 liiget", "3 puukooli", "Nõudlus nähtav"],
        ["2027", "Ostad Eesti CBD kreemi", "Kiri KOV-ile", "8 puukooli, 5 growshopi", "Pakkumine algab"],
        ["2028", "Toetad reguleerimist", "Koolide aiaprogramm", "15 puukooli, meditsiin", "Must turg väheneb"],
        ["2029", "Ostad kohalikku (10€/g)", "Urban farm linnas", "Eksport, 450 töökohta", "Raha Eestis"],
        ["2030", "Basic income pilot", "Loodus 500m kaugusel", "30 puukooli, vaba kassa", "Eesmärk saavutatud"],
    ],
)

add_normal(
    "5 aasta pärast oleme kindlasti eesmärgile lähemal kui täna — "
    "isegi kui kõik sammud ei õnnestu 100%. Iga samm loeb. Iga ost on hääl.",
)

# ═══════════════════════════════════════════
# 9. LÕPPSÕNA
# ═══════════════════════════════════════════
add_heading("9. Lõppsõna: sina oled turu juht", 1)

add_quote(
    "Sa ei pea ootama. Sa ei pea paluma. Sa pead valima — "
    "ja siis rääkima teistele, miks sa valisid. "
    "Küll pakkumine tuleb. Aga ainult siis, kui sina ja sinu naabrid "
    "näitate, et nõudlus on tõeline."
)

add_normal("Kolm asja, mida iga lugeja saab täna teha:", bold=True)
add_numbered("Osta üks Eesti tööstuskanepi toode (seemned, õli, kott)")
add_numbered("Räägi ühele inimesele sellest dokumendist")
add_numbered('Kirjuta üks lause KOV-ile või ettevõtjale: "Ma ostan, kui te toodate"')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("🌿 Sinu valik. Sinu turg. Sinu Eesti.")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

doc.add_paragraph()
add_normal("Seotud dokumendid:", bold=True)
add_bullet("kanepi-roheline-kuld-raport.md — visuaalne raport lastele ja täiskasvanutele")
add_bullet("Eesti-Kanepimajandus-5a-Finantsmudel.xlsx — kulud ja tulud")
add_bullet("riiklik-kanepiettevotte-mudel.md — ERMA AS strateegia")

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Koostatud august 2026 · Parem Tulevik").font.size = Pt(9)
footer.add_run("\nPõhimõte: ausus enne unistust. Unistus pärast plaani.").font.size = Pt(9)

doc.save(OUTPUT)
print(f"Salvestatud: {OUTPUT}")
